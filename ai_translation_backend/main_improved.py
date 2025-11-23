from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, RedirectResponse
from pydantic import BaseModel
import os
import io
import base64
import asyncio
import time
from typing import List, Optional, Dict, Tuple
from dotenv import load_dotenv
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt
import re
from concurrent.futures import ThreadPoolExecutor
import aiohttp
import json

load_dotenv()

# Allow OAuth scope changes (Google may return broader permissions than requested)
os.environ['OAUTHLIB_RELAX_TOKEN_SCOPE'] = '1'

app = FastAPI(title="Drive Document Translator API")

# CORS middleware - allow origins from environment or default to all
allowed_origins = os.getenv("ALLOWED_ORIGINS", "*").split(",")
if allowed_origins == ["*"]:
    # Allow all origins
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=False,  # Must be False when allow_origins is ["*"]
        allow_methods=["*"],
        allow_headers=["*"],
    )
else:
    # Specific origins with credentials support
    app.add_middleware(
        CORSMiddleware,
        allow_origins=[origin.strip() for origin in allowed_origins],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

# Configuration
MAX_RETRIES = 3
RETRY_DELAY = 2

# Enhanced smart batch configuration with better poetry detection
def get_smart_batch_size(text: str, para_style: str = None, para_alignment: int = None) -> int:
    """
    Enhanced batch size determination with better poetry and formatting detection.
    """
    # Count various formatting indicators
    line_breaks = text.count('\n')
    leading_spaces = len(text) - len(text.lstrip())
    trailing_spaces = len(text) - len(text.rstrip())
    has_indentation = leading_spaces > 0
    
    # Check for poetry/verse patterns
    lines = text.split('\n')
    short_lines = sum(1 for line in lines if 0 < len(line.strip()) < 50)
    empty_lines = sum(1 for line in lines if not line.strip())
    
    # Poetry/Verse detection (enhanced)
    is_poetry = (
        # Original checks
        '\\' in text or line_breaks > 3 or
        # New checks
        has_indentation or  # Any indentation suggests formatting
        (short_lines > 2 and len(lines) > 2) or  # Multiple short lines
        (empty_lines > 1 and len(lines) > 3) or  # Multiple empty lines
        (para_alignment and para_alignment == 1) or  # Center aligned
        re.search(r'^\s{2,}', text, re.MULTILINE)  # Lines starting with spaces
    )
    
    if is_poetry:
        # For poetry, go even smaller - paragraph by paragraph is safest
        return 1  # Process one paragraph at a time for perfect preservation
    
    # Dialogue-heavy content (enhanced detection)
    dialogue_indicators = (
        text.count('"') + text.count('"') + text.count('"') + 
        text.count('—') + text.count('–') + text.count('-"')
    )
    if dialogue_indicators > 4 or '"' in text:
        return 5  # Smaller batches for dialogue
    
    # List or enumerated content
    if re.search(r'^[\d\.\-\*\•]\s', text, re.MULTILINE):
        return 3  # Very small batches for lists
    
    # Long prose paragraphs (simple narrative text)
    if len(text) > 500 and '.' in text and not has_indentation:
        return 20  # Moderate batches for prose (reduced from 300)
    
    # Default for mixed content
    return 10  # Smaller default (reduced from 100)

# OAuth configuration
SCOPES = [
    'https://www.googleapis.com/auth/drive.file',  # Create and modify files that the app creates
    'https://www.googleapis.com/auth/drive.readonly',  # Read files
    'https://www.googleapis.com/auth/drive.metadata.readonly'  # Read metadata
]

# In-memory session storage (in production, use Redis or database)
sessions = {}

# In-memory progress tracking
progress_tracker = {}

# Request/Response Models
class TranslateRequest(BaseModel):
    fileData: str
    fileName: str
    language: str
    model: str
    apiKey: str
    progressId: Optional[str] = None

class TranslateResponse(BaseModel):
    translatedDocument: str
    logs: List[str]
    stats: dict

class ProgressResponse(BaseModel):
    completedBatches: int
    totalBatches: int
    error: bool = False

class CreateFolderRequest(BaseModel):
    state: str
    folderName: str
    parentFolderId: Optional[str] = None

class UploadFileRequest(BaseModel):
    state: str
    folderId: str
    fileName: str
    fileData: str
    mimeType: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

@app.get("/")
def read_root():
    return {"status": "Drive Document Translator API is running"}

@app.get("/api/drive/auth")
async def drive_auth():
    """Initiate OAuth flow"""
    try:
        flow = Flow.from_client_config(
            {
                "web": {
                    "client_id": os.getenv("GOOGLE_CLIENT_ID"),
                    "client_secret": os.getenv("GOOGLE_CLIENT_SECRET"),
                    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                    "token_uri": "https://oauth2.googleapis.com/token",
                    "redirect_uris": [f"{os.getenv('BACKEND_URL', 'http://localhost:8000')}/api/drive/callback"]
                }
            },
            scopes=SCOPES,
            redirect_uri=f"{os.getenv('GOOGLE_REDIRECT_URI', 'http://localhost:8000')}/api/drive/callback"
        )
        
        authorization_url, state = flow.authorization_url(
            access_type='offline',
            include_granted_scopes='true',
            prompt='consent'
        )
        
        # Store state for verification
        sessions[state] = {"flow": flow}
        
        return {"authUrl": authorization_url, "state": state}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Auth failed: {str(e)}")

@app.get("/api/drive/callback")
async def drive_callback(code: str, state: str):
    """Handle OAuth callback"""
    try:
        print(f"[CALLBACK] Received callback with state: {state}")
        print(f"[CALLBACK] Available sessions: {list(sessions.keys())}")
        
        if state not in sessions:
            print(f"[CALLBACK ERROR] State not found in sessions")
            raise HTTPException(status_code=400, detail="Invalid state")
        
        print(f"[CALLBACK] Fetching token with code...")
        flow = sessions[state]["flow"]
        flow.fetch_token(code=code)
        
        credentials = flow.credentials
        print(f"[CALLBACK] Credentials obtained successfully")
        
        sessions[state]["credentials"] = {
            "token": credentials.token,
            "refresh_token": credentials.refresh_token,
            "token_uri": credentials.token_uri,
            "client_id": credentials.client_id,
            "client_secret": credentials.client_secret,
            "scopes": credentials.scopes
        }
        
        print(f"[CALLBACK] Redirecting to frontend with state: {state}")
        # Redirect to frontend with state
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        return RedirectResponse(url=f"{frontend_url}/?authenticated=true&state={state}")
    except Exception as e:
        print(f"[CALLBACK ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        return RedirectResponse(url=f"{frontend_url}/?error=auth_failed&message={str(e)}")

@app.get("/api/drive/folders")
async def get_folders(state: str):
    """List folders from Google Drive"""
    try:
        if state not in sessions or "credentials" not in sessions[state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[state]["credentials"]
        credentials = Credentials(**creds_dict)
        
        service = build('drive', 'v3', credentials=credentials)
        
        results = service.files().list(
            q="mimeType='application/vnd.google-apps.folder' and trashed = false",
            pageSize=100,
            fields="files(id, name)"
        ).execute()
        
        return {"folders": results.get('files', [])}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch folders: {str(e)}")

@app.get("/api/drive/files")
async def get_files(state: str, folderId: Optional[str] = None, driveLink: Optional[str] = None):
    """List files from a folder or Drive link"""
    try:
        if state not in sessions or "credentials" not in sessions[state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[state]["credentials"]
        credentials = Credentials(**creds_dict)
        service = build('drive', 'v3', credentials=credentials)
        
        # Extract ID from link if provided
        if driveLink:
            folderId = extract_drive_id(driveLink)
        
        if not folderId:
            raise HTTPException(status_code=400, detail="Folder ID or Drive link required")
        
        # Get metadata to check if folder or file
        metadata = service.files().get(
            fileId=folderId,
            fields="id, name, mimeType"
        ).execute()
        
        is_folder = metadata['mimeType'] == 'application/vnd.google-apps.folder'
        
        if not is_folder:
            return {
                "type": "file",
                "file": metadata
            }
        
        # List files in folder
        supported_types = [
            'application/vnd.google-apps.document',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/pdf'
        ]
        
        mime_conditions = ' or '.join([f"mimeType='{t}'" for t in supported_types])
        query = f"'{folderId}' in parents and trashed = false and ({mime_conditions})"
        
        results = service.files().list(
            q=query,
            pageSize=100,
            fields="files(id, name, mimeType, size, modifiedTime)"
        ).execute()
        
        return {
            "type": "folder",
            "folderName": metadata['name'],
            "files": results.get('files', [])
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch files: {str(e)}")

@app.get("/api/drive/download")
async def download_file(state: str, fileId: str, mimeType: str):
    """Download a file from Google Drive"""
    try:
        if state not in sessions or "credentials" not in sessions[state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[state]["credentials"]
        credentials = Credentials(**creds_dict)
        service = build('drive', 'v3', credentials=credentials)
        
        # Google Workspace types that need export
        workspace_types = {
            'application/vnd.google-apps.document': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        }
        
        if mimeType in workspace_types:
            request = service.files().export_media(
                fileId=fileId,
                mimeType=workspace_types[mimeType]
            )
        else:
            request = service.files().get_media(fileId=fileId)
        
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        
        while not done:
            status, done = downloader.next_chunk()
        
        fh.seek(0)
        file_data = fh.read()
        
        # Convert to base64
        base64_data = base64.b64encode(file_data).decode('utf-8')
        
        return {
            "data": base64_data,
            "size": len(file_data)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to download file: {str(e)}")

@app.post("/api/translate", response_model=TranslateResponse)
async def translate_document(request: TranslateRequest):
    """Translate a document with async batch processing using Gemini"""
    try:
        print(f"[TRANSLATE] Starting async translation for {request.fileName}")
        
        # Decode base64 file data
        file_bytes = base64.b64decode(request.fileData)
        print(f"[TRANSLATE] File decoded, size: {len(file_bytes)} bytes")
        
        # Translate document using async function
        result = await translate_document_content_async(
            file_bytes,
            request.fileName,
            request.language,
            request.model,
            request.apiKey,
            request.progressId
        )
        
        print(f"[TRANSLATE] Translation complete")
        return result
    except Exception as e:
        print(f"[TRANSLATE] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        # Mark progress as failed
        if request.progressId and request.progressId in progress_tracker:
            progress_tracker[request.progressId]["error"] = True
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

@app.post("/api/translate/openrouter", response_model=TranslateResponse)
async def translate_document_openrouter(request: TranslateRequest):
    """Translate a document with async batch processing using OpenRouter"""
    try:
        print(f"[TRANSLATE OPENROUTER] Starting async translation for {request.fileName}")
        
        # Decode base64 file data
        file_bytes = base64.b64decode(request.fileData)
        print(f"[TRANSLATE OPENROUTER] File decoded, size: {len(file_bytes)} bytes")
        
        # Translate document using OpenRouter async function
        result = await translate_document_content_async_openrouter(
            file_bytes,
            request.fileName,
            request.language,
            request.model,
            request.apiKey,
            request.progressId
        )
        
        print(f"[TRANSLATE OPENROUTER] Translation complete")
        return result
    except Exception as e:
        print(f"[TRANSLATE OPENROUTER] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        # Mark progress as failed
        if request.progressId and request.progressId in progress_tracker:
            progress_tracker[request.progressId]["error"] = True
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

@app.get("/api/translate/progress/{progress_id}", response_model=ProgressResponse)
async def get_translation_progress(progress_id: str):
    """Get translation progress by ID"""
    if progress_id not in progress_tracker:
        raise HTTPException(status_code=404, detail=f"Progress ID not found: {progress_id}")
    
    progress = progress_tracker[progress_id]
    
    return ProgressResponse(
        completedBatches=progress.get("completedBatches", 0),
        totalBatches=progress.get("totalBatches", 0),
        error=progress.get("error", False)
    )

@app.post("/api/drive/logout")
async def logout(state: str):
    """Clear session"""
    if state in sessions:
        del sessions[state]
    return {"success": True}

@app.post("/api/drive/create-folder")
async def create_folder(request: CreateFolderRequest):
    """Create a new folder in Google Drive"""
    try:
        if request.state not in sessions or "credentials" not in sessions[request.state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[request.state]["credentials"]
        credentials = Credentials(**creds_dict)
        service = build('drive', 'v3', credentials=credentials)
        
        file_metadata = {
            'name': request.folderName,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        
        # If parent folder specified, add to parents
        if request.parentFolderId:
            file_metadata['parents'] = [request.parentFolderId]
        
        folder = service.files().create(
            body=file_metadata,
            fields='id, name, webViewLink'
        ).execute()
        
        print(f"[DRIVE] Created folder: {folder.get('name')} (ID: {folder.get('id')})")
        
        return {
            "folderId": folder.get('id'),
            "folderName": folder.get('name'),
            "webViewLink": folder.get('webViewLink')
        }
    except Exception as e:
        print(f"[DRIVE] Error creating folder: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create folder: {str(e)}")

@app.post("/api/drive/upload")
async def upload_to_drive(request: UploadFileRequest):
    """Upload a file to Google Drive folder"""
    try:
        if request.state not in sessions or "credentials" not in sessions[request.state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[request.state]["credentials"]
        credentials = Credentials(**creds_dict)
        service = build('drive', 'v3', credentials=credentials)
        
        # Decode base64 file data
        file_bytes = base64.b64decode(request.fileData)
        
        print(f"[DRIVE] Uploading {request.fileName} to folder {request.folderId}, size: {len(file_bytes)} bytes")
        
        # Create file metadata
        file_metadata = {
            'name': request.fileName,
            'parents': [request.folderId]
        }
        
        # Upload file using MediaIoBaseUpload
        from googleapiclient.http import MediaIoBaseUpload
        
        media = MediaIoBaseUpload(
            io.BytesIO(file_bytes),
            mimetype=request.mimeType,
            resumable=True
        )
        
        file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, name, webViewLink'
        ).execute()
        
        print(f"[DRIVE] File uploaded successfully: {file.get('name')} (ID: {file.get('id')})")
        
        return {
            "fileId": file.get('id'),
            "fileName": file.get('name'),
            "webViewLink": file.get('webViewLink')
        }
    except Exception as e:
        print(f"[DRIVE] Error uploading file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to upload file: {str(e)}")

# Helper functions
def extract_drive_id(link: str) -> str:
    """Extract file or folder ID from Google Drive link"""
    import re
    
    patterns = [
        r'/file/d/([a-zA-Z0-9_-]+)',
        r'/folders/([a-zA-Z0-9_-]+)',
        r'[?&]id=([a-zA-Z0-9_-]+)',
        r'/document/d/([a-zA-Z0-9_-]+)',
        r'/spreadsheets/d/([a-zA-Z0-9_-]+)',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, link)
        if match:
            return match.group(1)
    
    return link.strip()

def is_meaningful_text(text):
    """Check if text contains meaningful content"""
    cleaned = re.sub(r'[\W_]+', '', text)
    return bool(cleaned.strip())

def is_decorative_only(text):
    """Check if text is decorative only (symbols, single letters, etc.)"""
    stripped = text.strip()
    return not stripped or re.fullmatch(r"[^\w\s]+", stripped) or re.fullmatch(r"[A-Z]", stripped)

def sanitize_response(text: str) -> str:
    """Remove any <think>...</think> tokens and trim the response."""
    if not text:
        return text
    text = re.sub(r'<think>.*?</think>', '', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'</?think>', '', text, flags=re.IGNORECASE)
    # Do NOT strip here - preserve leading/trailing spaces
    return text

def preserve_paragraph_formatting(para) -> Dict[str, any]:
    """Extract and preserve paragraph formatting information."""
    formatting = {
        'alignment': para.alignment,
        'style_name': para.style.name if para.style else None,
        'left_indent': para.paragraph_format.left_indent,
        'right_indent': para.paragraph_format.right_indent,
        'first_line_indent': para.paragraph_format.first_line_indent,
        'space_before': para.paragraph_format.space_before,
        'space_after': para.paragraph_format.space_after,
        'line_spacing': para.paragraph_format.line_spacing,
        'keep_together': para.paragraph_format.keep_together,
        'keep_with_next': para.paragraph_format.keep_with_next,
        'page_break_before': para.paragraph_format.page_break_before,
    }
    
    # Preserve run-level formatting
    formatting['runs'] = []
    for run in para.runs:
        run_fmt = {
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'strike': run.strike,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
            'highlight_color': run.font.highlight_color,
            'text': run.text,  # Preserve exact text including spaces
        }
        formatting['runs'].append(run_fmt)
    
    return formatting

def apply_paragraph_formatting(para, formatting: Dict[str, any], translated_text: str):
    """Apply preserved formatting to translated paragraph."""
    # Apply paragraph-level formatting
    if formatting['alignment'] is not None:
        para.alignment = formatting['alignment']
    if formatting['style_name']:
        para.style = formatting['style_name']
    
    # Apply paragraph format properties
    para_fmt = para.paragraph_format
    if formatting['left_indent']:
        para_fmt.left_indent = formatting['left_indent']
    if formatting['right_indent']:
        para_fmt.right_indent = formatting['right_indent']
    if formatting['first_line_indent']:
        para_fmt.first_line_indent = formatting['first_line_indent']
    if formatting['space_before']:
        para_fmt.space_before = formatting['space_before']
    if formatting['space_after']:
        para_fmt.space_after = formatting['space_after']
    if formatting['line_spacing']:
        para_fmt.line_spacing = formatting['line_spacing']
    if formatting['keep_together'] is not None:
        para_fmt.keep_together = formatting['keep_together']
    if formatting['keep_with_next'] is not None:
        para_fmt.keep_with_next = formatting['keep_with_next']
    if formatting['page_break_before'] is not None:
        para_fmt.page_break_before = formatting['page_break_before']
    
    # For now, set the entire translated text as one run
    # TODO: Implement intelligent run preservation based on formatting patterns
    para.clear()
    run = para.add_run(translated_text)
    
    # Apply formatting from the first run (as a baseline)
    if formatting['runs']:
        first_run_fmt = formatting['runs'][0]
        if first_run_fmt['bold']:
            run.bold = True
        if first_run_fmt['italic']:
            run.italic = True
        if first_run_fmt['underline']:
            run.underline = True
        if first_run_fmt['font_name']:
            run.font.name = first_run_fmt['font_name']
        if first_run_fmt['font_size']:
            run.font.size = first_run_fmt['font_size']

def call_gemini_batch_api(client, prompt, model, logs=None):
    """
    Synchronous function to call Gemini API for batch processing with token tracking.
    This will be run in a thread executor.
    """
    for attempt in range(MAX_RETRIES):
        try:
            if logs is not None:
                logs.append(f"[BATCH API] Attempt {attempt + 1}/{MAX_RETRIES} - Model: {model}")
            
            # Generate content with JSON response format
            response = client.models.generate_content(
                model=model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.0,
                    response_mime_type="application/json"
                )
            )
            
            result_text = response.text.strip()
            
            # Get token usage from response
            usage = response.usage_metadata
            actual_input_tokens = usage.prompt_token_count
            output_tokens = usage.candidates_token_count
            total_tokens = usage.total_token_count
            
            if logs is not None:
                logs.append(f"[SUCCESS] Received response ({len(result_text)} chars)")
                logs.append(f"[TOKENS] Input: {actual_input_tokens}, Output: {output_tokens}, Total: {total_tokens}")
            
            return {
                'text': result_text,
                'input_tokens': actual_input_tokens,
                'output_tokens': output_tokens,
                'total_tokens': total_tokens
            }
            
        except Exception as e:
            if logs is not None:
                logs.append(f"[ERROR] Attempt {attempt + 1} failed: {str(e)}")
            
            if attempt < MAX_RETRIES - 1:
                time.sleep(RETRY_DELAY)
    
    if logs is not None:
        logs.append("[FAILED] All retry attempts exhausted")
    return None

async def call_gemini_batch_async(executor, client, prompt, model, logs=None):
    """Async wrapper for batch API call using ThreadPoolExecutor."""
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(executor, call_gemini_batch_api, client, prompt, model, logs)
    return result

async def call_openrouter_batch_api(session, prompt, model, openrouter_api_key, log_list=None):
    """
    Call OpenRouter API for batch processing with token tracking.
    """
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {openrouter_api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": os.getenv("OPENROUTER_REFERER", "https://yourdomain.com"),
        "X-Title": "EasyTranslate"
    }
    
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.0,
        "response_format": {"type": "json_object"}
    }
    
    for attempt in range(MAX_RETRIES):
        try:
            if log_list is not None:
                log_entry = f"[BATCH API] Attempt {attempt + 1}/{MAX_RETRIES} - Model: {model}"
                log_list.append(log_entry)
            
            async with session.post(url, headers=headers, json=payload) as resp:
                resp.raise_for_status()
                data = await resp.json()
                
                result = data["choices"][0]["message"]["content"].strip()
                
                # Get token usage from response
                usage = data.get("usage", {})
                input_tokens = usage.get("prompt_tokens", 0)
                output_tokens = usage.get("completion_tokens", 0)
                total_tokens = usage.get("total_tokens", input_tokens + output_tokens)
                
                if log_list is not None:
                    log_list.append(f"[SUCCESS] Received response ({len(result)} chars)")
                    log_list.append(f"[TOKENS] Input: {input_tokens}, Output: {output_tokens}, Total: {total_tokens}")
                
                return {
                    'text': result,
                    'input_tokens': input_tokens,
                    'output_tokens': output_tokens,
                    'total_tokens': total_tokens
                }
                
        except Exception as e:
            if log_list is not None:
                log_list.append(f"[ERROR] Attempt {attempt + 1} failed: {str(e)}")
            
            if attempt < MAX_RETRIES - 1:
                await asyncio.sleep(RETRY_DELAY)
    
    if log_list is not None:
        log_list.append("[FAILED] All retry attempts exhausted")
    return None

def parse_structured_response(response_text, expected_count, logs=None):
    """Parse structured JSON response into individual translations."""
    try:
        # Clean response text (remove markdown code blocks if present)
        cleaned_text = response_text.strip()
        if cleaned_text.startswith('```json'):
            cleaned_text = cleaned_text[7:]
        if cleaned_text.endswith('```'):
            cleaned_text = cleaned_text[:-3]
        cleaned_text = cleaned_text.strip()
        
        # Parse JSON
        parsed_response = json.loads(cleaned_text)
        
        if logs:
            logs.append(f"[JSON] Successfully parsed structured response")
        
        # Extract translations in order
        translations = []
        if 'translations' in parsed_response:
            # Sort by ID to ensure correct order
            translation_items = sorted(parsed_response['translations'], key=lambda x: x.get('id', 0))
            
            for item in translation_items:
                if 'translation' in item:
                    translations.append(item['translation'])
                    
            if logs:
                logs.append(f"[JSON] Extracted {len(translations)} translations from structured response")
        
        return translations
        
    except json.JSONDecodeError as e:
        if logs:
            logs.append(f"[JSON ERROR] Failed to parse JSON: {str(e)}")
            logs.append(f"[FALLBACK] Attempting fallback parsing...")
        
        # Fallback: split by double newlines
        lines = response_text.split('\n\n')
        alt_translations = [t.strip() for t in lines if t.strip() and not t.startswith('{')]
        return alt_translations[:expected_count]
    
    except Exception as e:
        if logs:
            logs.append(f"[PARSE ERROR] {str(e)}")
        return []

async def translate_document_content_async(file_bytes: bytes, file_name: str, language: str, model: str, api_key: str, progress_id: Optional[str] = None) -> TranslateResponse:
    """Enhanced translation with better formatting preservation"""
    
    print(f"[TRANSLATOR] Initializing Gemini client")
    client = genai.Client(api_key=api_key)
    
    # Load document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = doc.paragraphs
    
    print(f"[TRANSLATOR] Document has {len(paragraphs)} total paragraphs")
    
    translated_content = []
    logs = []
    total_input_tokens = 0
    total_output_tokens = 0
    
    logs.append(f"[START] Enhanced batch translation started for language: {language}")
    logs.append(f"[INFO] Source file: {file_name}")
    logs.append(f"[INFO] Document has {len(paragraphs)} total paragraphs")
    logs.append(f"[INFO] Using ENHANCED SMART BATCHING with paragraph-level processing for complex content")
    logs.append(f"[INFO] Using Gemini model: {model}")
    logs.append(f"[INFO] Processing in memory - no files saved to disk")
    
    # Prepare paragraph batches with ENHANCED smart batching logic
    paragraph_batches = []
    current_batch = []
    current_max_size = 10  # Start with smaller default
    
    # Track stats for optimization reporting
    total_paragraphs_to_translate = 0
    batch_size_distribution = {'poetry': 0, 'dialogue': 0, 'prose': 0, 'list': 0, 'default': 0}
    
    # Store formatting information for each paragraph
    paragraph_formatting = {}
    
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        # CRITICAL: Do NOT strip - preserve all whitespace
        original = para.text  # Keep exact formatting
        
        # Skip single uppercase letters followed by uppercase text
        if re.fullmatch(r"[A-Z]", original.strip()) and i + 1 < len(paragraphs) and paragraphs[i + 1].text.strip()[:1].isupper():
            p = para._element
            p.getparent().remove(p)
            paragraphs = doc.paragraphs
            logs.append(f"[SKIP] Removed single uppercase letter at para {i}")
            continue
        
        # Skip empty or decorative text
        if not original.strip() or not is_meaningful_text(original) or is_decorative_only(original):
            i += 1
            continue
        
        word_count = len(original.split())
        is_heading = para.style.name.lower().startswith("heading") or para.alignment == 1
        
        # Skip single-word non-heading paragraphs
        if word_count <= 1:
            if not original.strip().isupper() and not is_heading:
                i += 1
                continue
        
        # Preserve formatting information
        paragraph_formatting[i] = preserve_paragraph_formatting(para)
        
        # ENHANCED SMART BATCHING: Get optimal batch size for this content
        optimal_size = get_smart_batch_size(original, para.style.name, para.alignment)
        
        # Track content types for stats
        if optimal_size == 1:
            batch_size_distribution['poetry'] += 1
        elif optimal_size == 5:
            batch_size_distribution['dialogue'] += 1
        elif optimal_size == 3:
            batch_size_distribution['list'] += 1
        elif optimal_size == 20:
            batch_size_distribution['prose'] += 1
        else:
            batch_size_distribution['default'] += 1
        
        total_paragraphs_to_translate += 1
        
        # If content type changed significantly, start a new batch
        if current_batch and abs(optimal_size - current_max_size) > 5:
            paragraph_batches.append(current_batch)
            current_batch = []
            current_max_size = optimal_size
        
        # Add to current batch (with exact text, not stripped)
        current_batch.append((i, para, original))
        
        # Update max size to be the most restrictive (smallest) in the batch
        if len(current_batch) == 1:
            current_max_size = optimal_size
        else:
            current_max_size = min(current_max_size, optimal_size)
        
        # If batch is full or we're at the end, save it
        if len(current_batch) >= current_max_size or i == len(paragraphs) - 1:
            if current_batch:
                paragraph_batches.append(current_batch)
                current_batch = []
                current_max_size = 10  # Reset to default
        
        i += 1
    
    total_batches = len(paragraph_batches)
    
    # Calculate optimization stats
    estimated_fixed_batches = max(1, total_paragraphs_to_translate // 20)
    optimization_percentage = ((1 - total_batches / estimated_fixed_batches) * 100) if estimated_fixed_batches > 0 else 0
    
    logs.append(f"[ENHANCED BATCHING] Created {total_batches} optimized batches")
    logs.append(f"[CONTENT ANALYSIS] Poetry/Formatted: {batch_size_distribution['poetry']}, Dialogue: {batch_size_distribution['dialogue']}, Lists: {batch_size_distribution['list']}, Prose: {batch_size_distribution['prose']}, Default: {batch_size_distribution['default']}")
    logs.append(f"[OPTIMIZATION] Would have been ~{estimated_fixed_batches} calls with fixed size 20")
    logs.append(f"[EFFICIENCY] Optimized API usage while preserving formatting")
    logs.append(f"[PROCESSING] Starting parallel batch API requests (max 4 concurrent)...")
    
    # Initialize progress tracking
    if progress_id:
        progress_tracker[progress_id] = {
            "totalBatches": total_batches,
            "completedBatches": 0,
            "error": False
        }
    
    # Create thread pool executor for async processing
    executor = ThreadPoolExecutor(max_workers=4)
    
    async def process_batch_gemini(batch_idx, batch):
        """Process a single batch for Gemini - returns logs separately"""
        batch_logs = []  # Separate logs for this batch
        batch_paragraphs = [item[2] for item in batch]  # Extract exact text
        batch_size = len(batch_paragraphs)
        
        print(f"[TRANSLATOR] Processing batch {batch_idx + 1}/{total_batches} ({batch_size} paragraphs)")
        batch_logs.append(f"[BATCH {batch_idx + 1}/{total_batches}] Processing {batch_size} paragraphs...")
        
        # Determine if this is a poetry/formatted batch
        is_formatted_batch = batch_size <= 5
        
        # Create enhanced prompt for better formatting preservation
        prompt = create_enhanced_batch_prompt(batch_paragraphs, language, is_formatted_batch)
        
        # Call API asynchronously
        batch_result = await call_gemini_batch_async(executor, client, prompt, model, batch_logs)
        
        # Update progress immediately after batch completes
        if progress_id:
            progress_tracker[progress_id]["completedBatches"] = progress_tracker[progress_id]["completedBatches"] + 1
        
        return batch_idx, batch, batch_paragraphs, batch_result, batch_logs
    
    # Create all tasks for parallel processing
    tasks = [
        process_batch_gemini(batch_idx, batch)
        for batch_idx, batch in enumerate(paragraph_batches)
    ]
    
    # Execute all tasks in parallel (limited by ThreadPoolExecutor max_workers)
    results = await asyncio.gather(*tasks)
    
    # Process results in order and merge logs
    for batch_idx, batch, batch_paragraphs, batch_result, batch_logs in results:
        # Add this batch's logs in order
        logs.extend(batch_logs)
        
        if batch_result:
            total_input_tokens += batch_result['input_tokens']
            total_output_tokens += batch_result['output_tokens']
            
            # Parse structured response
            batch_translations = parse_structured_response(batch_result['text'], len(batch_paragraphs), logs)
            
            logs.append(f"[BATCH {batch_idx + 1}] Received {len(batch_translations)} translations")
            
            # Validate we got the expected number of translations
            if len(batch_translations) != len(batch_paragraphs):
                logs.append(f"[WARNING] Expected {len(batch_paragraphs)} translations, got {len(batch_translations)}")
                # Pad with empty strings if we're short
                while len(batch_translations) < len(batch_paragraphs):
                    batch_translations.append('[Translation missing]')
                # Trim if we got too many
                batch_translations = batch_translations[:len(batch_paragraphs)]
            
            # Apply translations to document (maintains order)
            for (para_idx, para, original), translation in zip(batch, batch_translations):
                if translation and translation.strip():
                    # Do NOT sanitize if we want to preserve all formatting
                    translation = sanitize_response(translation)
                    translated_content.append(translation)
                    
                    # Apply translation with formatting preservation
                    if para_idx in paragraph_formatting:
                        apply_paragraph_formatting(para, paragraph_formatting[para_idx], translation)
                    else:
                        # Fallback: simple text replacement
                        para.clear()
                        para.add_run(translation)
                else:
                    fallback_text = f"[Translation failed for paragraph {para_idx}]"
                    translated_content.append(fallback_text)
        else:
            # Batch failed - mark as error and stop translation
            logs.append(f"[BATCH ERROR] Batch {batch_idx + 1} failed completely")
            if progress_id:
                progress_tracker[progress_id]["error"] = True
            raise Exception(f"Translation failed at batch {batch_idx + 1}/{total_batches}. Please try again.")
    
    # Shutdown executor
    executor.shutdown(wait=False)
    
    # Save document to memory buffer
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    logs.append(f"[SAVE] Document saved to memory buffer")
    logs.append(f"[TOKENS] Final usage - Input: {total_input_tokens}, Output: {total_output_tokens}, Total: {total_input_tokens + total_output_tokens}")
    logs.append(f"[DONE] Translation complete!")
    
    # Convert to base64
    translated_base64 = base64.b64encode(output_buffer.read()).decode('utf-8')
    
    print("[TRANSLATE] Complete, returning response")
    
    return TranslateResponse(
        translatedDocument=translated_base64,
        logs=logs,
        stats={
            "totalInputTokens": total_input_tokens,
            "totalOutputTokens": total_output_tokens,
            "totalTokens": total_input_tokens + total_output_tokens,
            "paragraphCount": len(translated_content),
            "translatedText": "\n\n".join(translated_content)  # Add full text for preview
        }
    )

async def translate_document_content_async_openrouter(file_bytes: bytes, file_name: str, language: str, model: str, api_key: str, progress_id: Optional[str] = None) -> TranslateResponse:
    """Enhanced translation with better formatting preservation for OpenRouter"""
    
    print(f"[TRANSLATOR] Initializing OpenRouter client")
    
    # Load document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = doc.paragraphs
    
    print(f"[TRANSLATOR] Document has {len(paragraphs)} total paragraphs")
    
    translated_content = []
    logs = []
    total_input_tokens = 0
    total_output_tokens = 0
    
    logs.append(f"[START] Enhanced batch translation started for language: {language}")
    logs.append(f"[INFO] Source file: {file_name}")
    logs.append(f"[INFO] Document has {len(paragraphs)} total paragraphs")
    logs.append(f"[INFO] Using ENHANCED SMART BATCHING with paragraph-level processing for complex content")
    logs.append(f"[INFO] Using OpenRouter model: {model}")
    logs.append(f"[INFO] Processing in memory - no files saved to disk")
    
    # Prepare paragraph batches with ENHANCED smart batching logic
    paragraph_batches = []
    current_batch = []
    current_max_size = 10  # Start with smaller default
    
    # Track stats for optimization reporting
    total_paragraphs_to_translate = 0
    batch_size_distribution = {'poetry': 0, 'dialogue': 0, 'prose': 0, 'list': 0, 'default': 0}
    
    # Store formatting information for each paragraph
    paragraph_formatting = {}
    
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        # CRITICAL: Do NOT strip - preserve all whitespace
        original = para.text  # Keep exact formatting
        
        # Skip single uppercase letters followed by uppercase text
        if re.fullmatch(r"[A-Z]", original.strip()) and i + 1 < len(paragraphs) and paragraphs[i + 1].text.strip()[:1].isupper():
            p = para._element
            p.getparent().remove(p)
            paragraphs = doc.paragraphs
            logs.append(f"[SKIP] Removed single uppercase letter at para {i}")
            continue
        
        # Skip empty or decorative text
        if not original.strip() or not is_meaningful_text(original) or is_decorative_only(original):
            i += 1
            continue
        
        word_count = len(original.split())
        is_heading = para.style.name.lower().startswith("heading") or para.alignment == 1
        
        # Skip single-word non-heading paragraphs
        if word_count <= 1:
            if not original.strip().isupper() and not is_heading:
                i += 1
                continue
        
        # Preserve formatting information
        paragraph_formatting[i] = preserve_paragraph_formatting(para)
        
        # ENHANCED SMART BATCHING: Get optimal batch size for this content
        optimal_size = get_smart_batch_size(original, para.style.name, para.alignment)
        
        # Track content types for stats
        if optimal_size == 1:
            batch_size_distribution['poetry'] += 1
        elif optimal_size == 5:
            batch_size_distribution['dialogue'] += 1
        elif optimal_size == 3:
            batch_size_distribution['list'] += 1
        elif optimal_size == 20:
            batch_size_distribution['prose'] += 1
        else:
            batch_size_distribution['default'] += 1
        
        total_paragraphs_to_translate += 1
        
        # If content type changed significantly, start a new batch
        if current_batch and abs(optimal_size - current_max_size) > 5:
            paragraph_batches.append(current_batch)
            current_batch = []
            current_max_size = optimal_size
        
        # Add to current batch (with exact text, not stripped)
        current_batch.append((i, para, original))
        
        # Update max size to be the most restrictive (smallest) in the batch
        if len(current_batch) == 1:
            current_max_size = optimal_size
        else:
            current_max_size = min(current_max_size, optimal_size)
        
        # If batch is full or we're at the end, save it
        if len(current_batch) >= current_max_size or i == len(paragraphs) - 1:
            if current_batch:
                paragraph_batches.append(current_batch)
                current_batch = []
                current_max_size = 10  # Reset to default
        
        i += 1
    
    total_batches = len(paragraph_batches)
    
    # Calculate optimization stats
    estimated_fixed_batches = max(1, total_paragraphs_to_translate // 20)
    optimization_percentage = ((1 - total_batches / estimated_fixed_batches) * 100) if estimated_fixed_batches > 0 else 0
    
    logs.append(f"[ENHANCED BATCHING] Created {total_batches} optimized batches")
    logs.append(f"[CONTENT ANALYSIS] Poetry/Formatted: {batch_size_distribution['poetry']}, Dialogue: {batch_size_distribution['dialogue']}, Lists: {batch_size_distribution['list']}, Prose: {batch_size_distribution['prose']}, Default: {batch_size_distribution['default']}")
    logs.append(f"[OPTIMIZATION] Would have been ~{estimated_fixed_batches} calls with fixed size 20")
    logs.append(f"[EFFICIENCY] Optimized API usage while preserving formatting")
    logs.append(f"[PROCESSING] Starting parallel batch API requests (max 4 concurrent)...")
    
    # Initialize progress tracking
    if progress_id:
        progress_tracker[progress_id] = {
            "totalBatches": total_batches,
            "completedBatches": 0,
            "error": False
        }
    
    # Create semaphore to limit concurrent requests
    semaphore = asyncio.Semaphore(4)
    
    async def process_batch_with_semaphore(batch_idx, batch, session):
        """Process a single batch with semaphore control - returns logs separately"""
        batch_logs = []  # Separate logs for this batch
        
        async with semaphore:
            batch_paragraphs = [item[2] for item in batch]  # Extract exact text
            batch_size = len(batch_paragraphs)
            
            print(f"[TRANSLATOR] Processing batch {batch_idx + 1}/{total_batches} ({batch_size} paragraphs)")
            batch_logs.append(f"[BATCH {batch_idx + 1}/{total_batches}] Processing {batch_size} paragraphs...")
            
            # Determine if this is a poetry/formatted batch
            is_formatted_batch = batch_size <= 5
            
            # Create enhanced prompt for better formatting preservation
            prompt = create_enhanced_batch_prompt(batch_paragraphs, language, is_formatted_batch)
            
            # Call API asynchronously
            batch_result = await call_openrouter_batch_api(session, prompt, model, api_key, batch_logs)
            
            # Update progress immediately after batch completes
            if progress_id:
                progress_tracker[progress_id]["completedBatches"] = progress_tracker[progress_id]["completedBatches"] + 1
            
            return batch_idx, batch, batch_paragraphs, batch_result, batch_logs
    
    # Create aiohttp session for async requests
    async with aiohttp.ClientSession() as session:
        # Create all tasks for parallel processing
        tasks = [
            process_batch_with_semaphore(batch_idx, batch, session)
            for batch_idx, batch in enumerate(paragraph_batches)
        ]
        
        # Execute all tasks in parallel (limited by semaphore)
        results = await asyncio.gather(*tasks)
        
        # Process results in order and merge logs
        for batch_idx, batch, batch_paragraphs, batch_result, batch_logs in results:
            # Add this batch's logs in order
            logs.extend(batch_logs)
            
            if batch_result:
                total_input_tokens += batch_result['input_tokens']
                total_output_tokens += batch_result['output_tokens']
                
                # Parse structured response
                batch_translations = parse_structured_response(batch_result['text'], len(batch_paragraphs), logs)
                
                logs.append(f"[BATCH {batch_idx + 1}] Received {len(batch_translations)} translations")
                
                # Validate we got the expected number of translations
                if len(batch_translations) != len(batch_paragraphs):
                    logs.append(f"[WARNING] Expected {len(batch_paragraphs)} translations, got {len(batch_translations)}")
                    # Pad with empty strings if we're short
                    while len(batch_translations) < len(batch_paragraphs):
                        batch_translations.append('[Translation missing]')
                    # Trim if we got too many
                    batch_translations = batch_translations[:len(batch_paragraphs)]
                
                # Apply translations to document (maintains order)
                for (para_idx, para, original), translation in zip(batch, batch_translations):
                    if translation and translation.strip():
                        # Do NOT sanitize if we want to preserve all formatting
                        translation = sanitize_response(translation)
                        translated_content.append(translation)
                        
                        # Apply translation with formatting preservation
                        if para_idx in paragraph_formatting:
                            apply_paragraph_formatting(para, paragraph_formatting[para_idx], translation)
                        else:
                            # Fallback: simple text replacement
                            para.clear()
                            para.add_run(translation)
                    else:
                        fallback_text = f"[Translation failed for paragraph {para_idx}]"
                        translated_content.append(fallback_text)
            else:
                # Batch failed - mark as error and stop translation
                logs.append(f"[BATCH ERROR] Batch {batch_idx + 1} failed completely")
                if progress_id:
                    progress_tracker[progress_id]["error"] = True
                raise Exception(f"Translation failed at batch {batch_idx + 1}/{total_batches}. Please try again.")
    
    # Save document to memory buffer
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    logs.append(f"[SAVE] Document saved to memory buffer")
    logs.append(f"[TOKENS] Final usage - Input: {total_input_tokens}, Output: {total_output_tokens}, Total: {total_input_tokens + total_output_tokens}")
    logs.append(f"[DONE] Translation complete!")
    
    # Convert to base64
    translated_base64 = base64.b64encode(output_buffer.read()).decode('utf-8')
    
    print("[TRANSLATE] Complete, returning response")
    
    return TranslateResponse(
        translatedDocument=translated_base64,
        logs=logs,
        stats={
            "totalInputTokens": total_input_tokens,
            "totalOutputTokens": total_output_tokens,
            "totalTokens": total_input_tokens + total_output_tokens,
            "paragraphCount": len(translated_content),
            "translatedText": "\n\n".join(translated_content)  # Add full text for preview
        }
    )

def create_enhanced_batch_prompt(paragraphs: List[str], language: str, is_formatted: bool = False) -> str:
    """Enhanced prompt with stronger formatting preservation instructions"""
    
    if is_formatted:
        # Special prompt for poetry/formatted content
        FORMATTED_PROMPT_TEMPLATE = """
You are a professional translator specializing in literary and formatted texts. You must translate {count} passage(s) from a manuscript into {language}.

ABSOLUTELY CRITICAL - FORMATTING PRESERVATION RULES:

1. EXACT CHARACTER-BY-CHARACTER PRESERVATION:
   - Count and preserve EVERY SPACE character (including leading spaces)
   - Count and preserve EVERY NEWLINE character (\n)
   - If a line starts with 4 spaces, the translation MUST start with 4 spaces
   - If there are 2 newlines between lines, there MUST be 2 newlines in translation
   - NEVER add or remove ANY whitespace characters

2. POETRY/FORMATTED TEXT SPECIFIC:
   - This appears to be poetry or specially formatted text
   - Line breaks are ARTISTIC CHOICES - preserve them exactly
   - Indentation creates visual rhythm - preserve every space
   - Short lines are intentional - keep them short in translation
   - DO NOT reflow text to "improve" readability

3. TRANSLATION APPROACH FOR FORMATTED TEXT:
   - Translate meaning while maintaining exact structure
   - If a line has 5 words in original, try to keep similar length
   - Preserve the visual shape of the text on the page
   - Think of it as translating a visual poem

4. SPACING VERIFICATION:
   - After translating, count the spaces at the start of each line
   - Verify newline counts between paragraphs match exactly
   - The translated text should have the SAME visual layout

OUTPUT FORMAT:
Return ONLY a valid JSON object:

{{
  "translations": [
    {{
      "id": 1,
      "translation": "EXACT formatting with ALL spaces and newlines preserved"
    }}
  ]
}}

CRITICAL: The visual layout of your translation should be IDENTICAL to the original when printed.

Original Passage(s) to Translate:
{passages}

Translate into {language} while preserving the EXACT formatting shown above.
"""
    else:
        # Standard prompt with strong formatting emphasis
        STANDARD_PROMPT_TEMPLATE = """
You are a professional translator tasked with translating {count} passages from a manuscript into {language}.

FORMATTING PRESERVATION RULES:

1. PRESERVE ALL WHITESPACE:
   - Keep ALL leading spaces (indentation)
   - Keep ALL trailing spaces
   - Keep ALL newlines and line breaks exactly as shown
   - Keep paragraph spacing identical

2. STRUCTURAL PRESERVATION:
   - Do NOT reformat or "improve" the layout
   - Do NOT change line breaks or paragraph breaks
   - If text appears oddly spaced, keep that odd spacing

3. TRANSLATION REQUIREMENTS:
   - Translate completely without shortening
   - Maintain natural, fluent {language}
   - Stay faithful to the meaning

OUTPUT FORMAT:
Return ONLY a valid JSON object:

{{
  "translations": [
    {{
      "id": 1,
      "translation": "translated text with exact formatting preserved"
    }},
    {{
      "id": 2,
      "translation": "translated text with exact formatting preserved"
    }}
  ]
}}

Original Passages to Translate:
{passages}

Remember: Translate the CONTENT but preserve the STRUCTURE exactly.
"""
    
    passages_text = ""
    for i, para in enumerate(paragraphs, 1):
        # Use repr() to show exact whitespace
        passages_text += f'\nPassage {i} (ID: {i}):\n"""\n{para}\n"""\n'
    
    template = FORMATTED_PROMPT_TEMPLATE if is_formatted else STANDARD_PROMPT_TEMPLATE
    
    return template.format(
        count=len(paragraphs),
        language=language,
        passages=passages_text
    )

if __name__ == "__main__":
    import uvicorn
    # Use reload so code changes are picked up automatically in development
    uvicorn.run(
        "main_improved:app",
        host="0.0.0.0",
        port=7860,
        reload=True,
    )
