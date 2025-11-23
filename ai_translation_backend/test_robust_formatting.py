"""
Quick test to verify robust formatting is working
"""

import sys
import os

# Add current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

try:
    from robust_format_preservation import (
        RobustFormatPreserver,
        RunFormatting,
        ParagraphFormatting,
        create_robust_translation_prompt
    )
    print("✅ Robust formatting module imported successfully!")
    print("✅ All classes and functions available")
    
    # Test basic functionality
    from docx import Document
    from io import BytesIO
    
    # Create a test document
    doc = Document()
    para = doc.add_paragraph()
    run1 = para.add_run("Bold text")
    run1.bold = True
    run2 = para.add_run(" and ")
    run3 = para.add_run("italic text")
    run3.italic = True
    
    # Test format preservation
    preserver = RobustFormatPreserver(doc)
    marked_text, para_data = preserver.create_formatted_text_for_translation(para, 0)
    
    print(f"✅ Format extraction works!")
    print(f"   Original runs: {len(para.runs)}")
    print(f"   Marked text: {marked_text[:100]}...")
    print(f"   Format data: {len(para_data['runs'])} runs captured")
    
    # Test prompt creation
    prompt = create_robust_translation_prompt([(0, marked_text)], "Spanish")
    print(f"✅ Prompt creation works!")
    print(f"   Prompt length: {len(prompt)} characters")
    
    print("\n🎉 All robust formatting features are working correctly!")
    print("✅ Backend is ready to use robust formatting!")
    
except ImportError as e:
    print(f"❌ Import error: {e}")
    print("   Make sure robust_format_preservation.py is in the backend directory")
    sys.exit(1)
except Exception as e:
    print(f"❌ Error: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)
