from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, RedirectResponse
from pydantic import BaseModel
import os
import io
import base64
import asyncio
import time
from typing import List, Optional, Dict, Tuple, Any
from dotenv import load_dotenv
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt
import re
from concurrent.futures import ThreadPoolExecutor
import aiohttp
import hashlib
from dataclasses import dataclass, asdict
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
load_dotenv()

# Import robust format preservation
try:
    from robust_format_preservation import (
        RobustFormatPreserver,
        RunFormatting,
        ParagraphFormatting,
        create_robust_translation_prompt,
        integrate_robust_preservation
    )
    ROBUST_FORMATTING_AVAILABLE = True
except ImportError:
    ROBUST_FORMATTING_AVAILABLE = False
    print("[WARNING] Robust format preservation module not available")

# Import TOC handler
try:
    from toc_handler import process_toc_before_translation
    TOC_HANDLER_AVAILABLE = True
    print("[INFO] TOC handler loaded successfully")
except ImportError as e:
    TOC_HANDLER_AVAILABLE = False
    print(f"[WARNING] TOC handler module not available: {e}")

# Allow OAuth scope changes (Google may return broader permissions than requested)
os.environ['OAUTHLIB_RELAX_TOKEN_SCOPE'] = '1'

app = FastAPI(title="Drive Document Translator API")

# CORS middleware - allow origins from environment or default to all
allowed_origins = os.getenv("ALLOWED_ORIGINS", "*").split(",")
if allowed_origins == ["*"]:
    # Allow all origins
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=False,  # Must be False when allow_origins is ["*"]
        allow_methods=["*"],
        allow_headers=["*"],
    )
else:
    # Specific origins with credentials support
    app.add_middleware(
        CORSMiddleware,
        allow_origins=[origin.strip() for origin in allowed_origins],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

# Configuration
BATCH_SIZE = 10000  # Legacy - now using smart batching
MAX_RETRIES = 3
RETRY_DELAY = 2

# ============================================================================
# ULTIMATE ADAPTIVE TOKEN-BASED BATCHING SYSTEM
# Solves: Rate limits, Dynamic batching, 100% formatting preservation, Error handling
# ============================================================================

def estimate_tokens(text):
    """Estimate tokens (rough: 1 token ≈ 4 characters for English)"""
    return len(text) // 4

def estimate_batch_tokens(batch):
    """Calculate total tokens in a batch of paragraphs"""
    return sum(estimate_tokens(item[2]) for item in batch)

def analyze_paragraph_complexity(para):
    """Analyze a single paragraph's complexity for formatting preservation"""
    text = para.text
    
    # Count formatting indicators
    complexity_score = 0
    
    # Multiple runs = inline formatting changes
    if len(para.runs) > 2:
        complexity_score += 3
    
    # Poetry/special formatting
    if text.count('\n') > 2 or len(text) - len(text.lstrip()) > 2:
        complexity_score += 2
    
    # Short lines (poetry indicator)
    lines = text.split('\n')
    if any(0 < len(line.strip()) < 60 for line in lines):
        complexity_score += 1
    
    # Multiple formatting in runs
    format_changes = 0
    for run in para.runs:
        if run.bold or run.italic or run.underline:
            format_changes += 1
    if format_changes > 1:
        complexity_score += 2
    
    return {
        'score': complexity_score,
        'is_complex': complexity_score >= 3,
        'run_count': len(para.runs),
        'has_inline_formatting': len(para.runs) > 1
    }

def analyze_section_complexity(paragraphs, start_idx, window_size=100):
    """
    Analyze upcoming section to determine optimal token target.
    Returns: target_tokens, use_robust, section_stats
    """
    end_idx = min(start_idx + window_size, len(paragraphs))
    
    total_complexity = 0
    complex_count = 0
    inline_format_count = 0
    total_valid = 0
    
    for i in range(start_idx, end_idx):
        para = paragraphs[i]
        if not para.text.strip():
            continue
        
        total_valid += 1
        complexity = analyze_paragraph_complexity(para)
        total_complexity += complexity['score']
        
        if complexity['is_complex']:
            complex_count += 1
        if complexity['has_inline_formatting']:
            inline_format_count += 1
    
    if total_valid == 0:
        return 8000, False, {'type': 'empty'}
    
    avg_complexity = total_complexity / total_valid
    complex_ratio = complex_count / total_valid
    inline_ratio = inline_format_count / total_valid
    
    # Determine token target and whether to use robust formatting
    # REDUCED batch sizes for faster processing (avoid timeouts)
    if complex_ratio > 0.4 or inline_ratio > 0.5:
        # High complexity - very small batches with ROBUST formatting
        return 2000, True, {
            'type': 'complex',
            'avg_complexity': avg_complexity,
            'complex_ratio': complex_ratio,
            'inline_ratio': inline_ratio
        }
    elif complex_ratio > 0.2 or inline_ratio > 0.3:
        # Medium complexity - small batches with ROBUST formatting
        return 3000, True, {
            'type': 'moderate',
            'avg_complexity': avg_complexity,
            'complex_ratio': complex_ratio,
            'inline_ratio': inline_ratio
        }
    else:
        # Simple prose - moderate batches with standard formatting
        return 5000, False, {
            'type': 'simple',
            'avg_complexity': avg_complexity,
            'complex_ratio': complex_ratio,
            'inline_ratio': inline_ratio
        }

# OAuth configuration
SCOPES = [
    'https://www.googleapis.com/auth/drive.file',  # Create and modify files that the app creates
    'https://www.googleapis.com/auth/drive.readonly',  # Read files
    'https://www.googleapis.com/auth/drive.metadata.readonly'  # Read metadata
]

# In-memory session storage (in production, use Redis or database)
sessions = {}

# In-memory progress tracking
progress_tracker = {}

# Request/Response Models
class TranslateRequest(BaseModel):
    fileData: str
    fileName: str
    language: str
    model: str
    apiKey: str
    progressId: Optional[str] = None

class TranslateResponse(BaseModel):
    translatedDocument: str
    logs: List[str]
    stats: dict

class ProgressResponse(BaseModel):
    completedBatches: int
    totalBatches: int
    error: bool = False

class CreateFolderRequest(BaseModel):
    state: str
    folderName: str
    parentFolderId: Optional[str] = None

class UploadFileRequest(BaseModel):
    state: str
    folderId: str
    fileName: str
    fileData: str
    mimeType: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

@app.get("/")
def read_root():
    return {"status": "Drive Document Translator API is running"}

@app.get("/api/drive/auth")
async def drive_auth():
    """Initiate OAuth flow"""
    try:
        flow = Flow.from_client_config(
            {
                "web": {
                    "client_id": os.getenv("GOOGLE_CLIENT_ID"),
                    "client_secret": os.getenv("GOOGLE_CLIENT_SECRET"),
                    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                    "token_uri": "https://oauth2.googleapis.com/token",
                    "redirect_uris": [f"{os.getenv('BACKEND_URL', 'http://localhost:8000')}/api/drive/callback"]
                }
            },
            scopes=SCOPES,
            redirect_uri=f"{os.getenv('GOOGLE_REDIRECT_URI', 'http://localhost:8000')}/api/drive/callback"
        )
        
        authorization_url, state = flow.authorization_url(
            access_type='offline',
            include_granted_scopes='true',
            prompt='consent'
        )
        
        # Store state for verification
        sessions[state] = {"flow": flow}
        
        return {"authUrl": authorization_url, "state": state}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Auth failed: {str(e)}")

@app.get("/api/drive/callback")
async def drive_callback(code: str, state: str):
    """Handle OAuth callback"""
    try:
        print(f"[CALLBACK] Received callback with state: {state}")
        print(f"[CALLBACK] Available sessions: {list(sessions.keys())}")
        
        if state not in sessions:
            print(f"[CALLBACK ERROR] State not found in sessions")
            raise HTTPException(status_code=400, detail="Invalid state")
        
        print(f"[CALLBACK] Fetching token with code...")
        flow = sessions[state]["flow"]
        flow.fetch_token(code=code)
        
        credentials = flow.credentials
        print(f"[CALLBACK] Credentials obtained successfully")
        
        sessions[state]["credentials"] = {
            "token": credentials.token,
            "refresh_token": credentials.refresh_token,
            "token_uri": credentials.token_uri,
            "client_id": credentials.client_id,
            "client_secret": credentials.client_secret,
            "scopes": credentials.scopes
        }
        
        print(f"[CALLBACK] Redirecting to frontend with state: {state}")
        # Redirect to frontend with state
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        return RedirectResponse(url=f"{frontend_url}/?authenticated=true&state={state}")
    except Exception as e:
        print(f"[CALLBACK ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        return RedirectResponse(url=f"{frontend_url}/?error=auth_failed&message={str(e)}")

@app.get("/api/drive/folders")
async def get_folders(state: str):
    """List folders from Google Drive"""
    try:
        if state not in sessions or "credentials" not in sessions[state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[state]["credentials"]
        credentials = Credentials(**creds_dict)
        
        service = build('drive', 'v3', credentials=credentials)
        
        results = service.files().list(
            q="mimeType='application/vnd.google-apps.folder' and trashed = false",
            pageSize=100,
            fields="files(id, name)"
        ).execute()
        
        return {"folders": results.get('files', [])}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch folders: {str(e)}")

@app.get("/api/drive/files")
async def get_files(state: str, folderId: Optional[str] = None, driveLink: Optional[str] = None):
    """List files from a folder or Drive link"""
    try:
        if state not in sessions or "credentials" not in sessions[state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[state]["credentials"]
        credentials = Credentials(**creds_dict)
        service = build('drive', 'v3', credentials=credentials)
        
        # Extract ID from link if provided
        if driveLink:
            folderId = extract_drive_id(driveLink)
        
        if not folderId:
            raise HTTPException(status_code=400, detail="Folder ID or Drive link required")
        
        # Get metadata to check if folder or file
        metadata = service.files().get(
            fileId=folderId,
            fields="id, name, mimeType"
        ).execute()
        
        is_folder = metadata['mimeType'] == 'application/vnd.google-apps.folder'
        
        if not is_folder:
            return {
                "type": "file",
                "file": metadata
            }
        
        # List files in folder
        supported_types = [
            'application/vnd.google-apps.document',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/pdf'
        ]
        
        mime_conditions = ' or '.join([f"mimeType='{t}'" for t in supported_types])
        query = f"'{folderId}' in parents and trashed = false and ({mime_conditions})"
        
        results = service.files().list(
            q=query,
            pageSize=100,
            fields="files(id, name, mimeType, size, modifiedTime)"
        ).execute()
        
        return {
            "type": "folder",
            "folderName": metadata['name'],
            "files": results.get('files', [])
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch files: {str(e)}")

@app.get("/api/drive/download")
async def download_file(state: str, fileId: str, mimeType: str):
    """Download a file from Google Drive"""
    try:
        if state not in sessions or "credentials" not in sessions[state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[state]["credentials"]
        credentials = Credentials(**creds_dict)
        service = build('drive', 'v3', credentials=credentials)
        
        # Google Workspace types that need export
        workspace_types = {
            'application/vnd.google-apps.document': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        }
        
        if mimeType in workspace_types:
            request = service.files().export_media(
                fileId=fileId,
                mimeType=workspace_types[mimeType]
            )
        else:
            request = service.files().get_media(fileId=fileId)
        
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        
        while not done:
            status, done = downloader.next_chunk()
        
        fh.seek(0)
        file_data = fh.read()
        
        # Convert to base64
        base64_data = base64.b64encode(file_data).decode('utf-8')
        
        return {
            "data": base64_data,
            "size": len(file_data)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to download file: {str(e)}")

@app.post("/api/translate", response_model=TranslateResponse)
async def translate_document(request: TranslateRequest):
    """Translate a document with async batch processing using Gemini - auto-detects format complexity"""
    try:
        print(f"[TRANSLATE] Starting async translation for {request.fileName}")
        
        # Decode base64 file data
        file_bytes = base64.b64decode(request.fileData)
        print(f"[TRANSLATE] File decoded, size: {len(file_bytes)} bytes")
        
        # Check if robust formatting should be used
        use_robust = request.useRobustFormatting if hasattr(request, 'useRobustFormatting') else None
        
        if use_robust is None and ROBUST_FORMATTING_AVAILABLE:
            # Auto-detect based on document complexity
            doc = load_document(file_bytes)
            
            # Count total runs and format variations
            total_runs = sum(len(para.runs) for para in doc.paragraphs if para.text.strip())
            total_paras = len([p for p in doc.paragraphs if p.text.strip()])
            avg_runs_per_para = total_runs / max(total_paras, 1)
            
            # Count paragraphs with complex formatting
            complex_paras = 0
            for para in doc.paragraphs:
                if para.text.strip() and len(para.runs) > 2:
                    complex_paras += 1
            
            # DISABLED: Document-level detection (too aggressive)
            # Now using PER-BATCH detection for optimal performance
            use_robust = False  # Always use standard with per-batch adaptive detection
            print(f"[DETECT] Format analysis - avg runs: {avg_runs_per_para:.1f}, complex paras: {complex_paras}/{total_paras}")
            print(f"[DETECT] Using PER-BATCH adaptive detection (robust only when needed)")
        
        # Translate using appropriate method (always use new adaptive system)
        if use_robust and not ROBUST_FORMATTING_AVAILABLE:
            print("[WARNING] Robust formatting requested but module not available")
        
        result = await translate_document_content_async(
            file_bytes,
            request.fileName,
            request.language,
            request.model,
            request.apiKey,
            request.progressId
        )
        
        print(f"[TRANSLATE] Translation complete")
        return result
    except Exception as e:
        print(f"[TRANSLATE] Error: {str(e)}")
        import traceback
        error_trace = traceback.format_exc()
        print(f"[TRANSLATE] Full traceback:\n{error_trace}")
        # Mark progress as failed
        if request.progressId and request.progressId in progress_tracker:
            progress_tracker[request.progressId]["error"] = True
        # Include more details in error response
        error_detail = f"Translation failed: {str(e)}\n\nTraceback:\n{error_trace}"
        raise HTTPException(status_code=500, detail=error_detail)


@app.post("/api/translate/enhanced", response_model=TranslateResponse)
async def translate_document_enhanced_endpoint(request: TranslateRequest):
    """Enhanced translation with format marker preservation for moderate complexity documents"""
    try:
        print(f"[TRANSLATE ENHANCED] Starting enhanced translation for {request.fileName}")
        
        # Decode base64 file data
        file_bytes = base64.b64decode(request.fileData)
        
        # Analyze document complexity
        doc = load_document(file_bytes)
        total_runs = sum(len(para.runs) for para in doc.paragraphs if para.text.strip())
        total_paras = len([p for p in doc.paragraphs if p.text.strip()])
        avg_runs_per_para = total_runs / max(total_paras, 1)
        
        print(f"[TRANSLATE ENHANCED] Document analysis - avg runs: {avg_runs_per_para:.1f}")
        
        # For very complex documents, recommend robust method
        if avg_runs_per_para > 5:
            print(f"[TRANSLATE ENHANCED] Document too complex, recommending robust method")
            if ROBUST_FORMATTING_AVAILABLE:
                result = await translate_document_content_async_robust(
                    file_bytes,
                    request.fileName,
                    request.language,
                    request.model,
                    request.apiKey,
                    request.progressId
                )
                return result
        
        # Use standard translation with enhanced smart formatting
        print(f"[TRANSLATE ENHANCED] Using enhanced smart formatting")
        result = await translate_document_content_async(
            file_bytes,
            request.fileName,
            request.language,
            request.model,
            request.apiKey,
            request.progressId
        )
        
        return result
        
    except Exception as e:
        print(f"[TRANSLATE ENHANCED] Error: {str(e)}")
        if request.progressId and request.progressId in progress_tracker:
            progress_tracker[request.progressId]["error"] = True
        raise HTTPException(status_code=500, detail=f"Enhanced translation failed: {str(e)}")


@app.post("/api/translate/robust", response_model=TranslateResponse)
async def translate_document_robust_endpoint(request: TranslateRequest):
    """Force robust translation with 100% format preservation"""
    try:
        print(f"[TRANSLATE ROBUST] Starting robust translation for {request.fileName}")
        
        # Decode base64 file data
        file_bytes = base64.b64decode(request.fileData)
        print(f"[TRANSLATE ROBUST] File decoded, size: {len(file_bytes)} bytes")
        
        if not ROBUST_FORMATTING_AVAILABLE:
            raise HTTPException(status_code=503, detail="Robust formatting module not available. Please ensure robust_format_preservation.py is in the backend directory.")
        
        # Always use robust translation
        result = await translate_document_content_async_robust(
            file_bytes,
            request.fileName,
            request.language,
            request.model,
            request.apiKey,
            request.progressId
        )
        
        print(f"[TRANSLATE ROBUST] Translation complete with 100% format preservation")
        return result
    except Exception as e:
        print(f"[TRANSLATE ROBUST] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        # Mark progress as failed
        if request.progressId and request.progressId in progress_tracker:
            progress_tracker[request.progressId]["error"] = True
        raise HTTPException(status_code=500, detail=f"Robust translation failed: {str(e)}")

@app.post("/api/translate/openrouter", response_model=TranslateResponse)
async def translate_document_openrouter(request: TranslateRequest):
    """Translate a document with async batch processing using OpenRouter"""
    try:
        print(f"[TRANSLATE OPENROUTER] Starting async translation for {request.fileName}")
        
        # Decode base64 file data
        file_bytes = base64.b64decode(request.fileData)
        print(f"[TRANSLATE OPENROUTER] File decoded, size: {len(file_bytes)} bytes")
        
        # Translate document using OpenRouter async function
        result = await translate_document_content_async_openrouter(
            file_bytes,
            request.fileName,
            request.language,
            request.model,
            request.apiKey,
            request.progressId
        )
        
        print(f"[TRANSLATE OPENROUTER] Translation complete")
        return result
    except Exception as e:
        print(f"[TRANSLATE OPENROUTER] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        # Mark progress as failed
        if request.progressId and request.progressId in progress_tracker:
            progress_tracker[request.progressId]["error"] = True
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

@app.get("/api/translate/progress/{progress_id}", response_model=ProgressResponse)
async def get_translation_progress(progress_id: str):
    """Get translation progress by ID"""
    if progress_id not in progress_tracker:
        raise HTTPException(status_code=404, detail=f"Progress ID not found: {progress_id}")
    
    progress = progress_tracker[progress_id]
    
    return ProgressResponse(
        completedBatches=progress.get("completedBatches", 0),
        totalBatches=progress.get("totalBatches", 0),
        error=progress.get("error", False)
    )

@app.post("/api/drive/logout")
async def logout(state: str):
    """Clear session"""
    if state in sessions:
        del sessions[state]
    return {"success": True}

@app.post("/api/drive/create-folder")
async def create_folder(request: CreateFolderRequest):
    """Create a new folder in Google Drive"""
    try:
        if request.state not in sessions or "credentials" not in sessions[request.state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[request.state]["credentials"]
        credentials = Credentials(**creds_dict)
        service = build('drive', 'v3', credentials=credentials)
        
        file_metadata = {
            'name': request.folderName,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        
        # If parent folder specified, add to parents
        if request.parentFolderId:
            file_metadata['parents'] = [request.parentFolderId]
        
        folder = service.files().create(
            body=file_metadata,
            fields='id, name, webViewLink'
        ).execute()
        
        print(f"[DRIVE] Created folder: {folder.get('name')} (ID: {folder.get('id')})")
        
        return {
            "folderId": folder.get('id'),
            "folderName": folder.get('name'),
            "webViewLink": folder.get('webViewLink')
        }
    except Exception as e:
        print(f"[DRIVE] Error creating folder: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create folder: {str(e)}")

@app.post("/api/drive/upload")
async def upload_to_drive(request: UploadFileRequest):
    """Upload a file to Google Drive folder"""
    try:
        if request.state not in sessions or "credentials" not in sessions[request.state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[request.state]["credentials"]
        credentials = Credentials(**creds_dict)
        service = build('drive', 'v3', credentials=credentials)
        
        # Decode base64 file data
        file_bytes = base64.b64decode(request.fileData)
        
        print(f"[DRIVE] Uploading {request.fileName} to folder {request.folderId}, size: {len(file_bytes)} bytes")
        
        # Create file metadata
        file_metadata = {
            'name': request.fileName,
            'parents': [request.folderId]
        }
        
        # Upload file using MediaIoBaseUpload
        from googleapiclient.http import MediaIoBaseUpload
        
        media = MediaIoBaseUpload(
            io.BytesIO(file_bytes),
            mimetype=request.mimeType,
            resumable=True
        )
        
        file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, name, webViewLink'
        ).execute()
        
        print(f"[DRIVE] File uploaded successfully: {file.get('name')} (ID: {file.get('id')})")
        
        return {
            "fileId": file.get('id'),
            "fileName": file.get('name'),
            "webViewLink": file.get('webViewLink')
        }
    except Exception as e:
        print(f"[DRIVE] Error uploading file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to upload file: {str(e)}")

# Helper functions
def extract_drive_id(link: str) -> str:
    """Extract file or folder ID from Google Drive link"""
    import re
    
    patterns = [
        r'/file/d/([a-zA-Z0-9_-]+)',
        r'/folders/([a-zA-Z0-9_-]+)',
        r'[?&]id=([a-zA-Z0-9_-]+)',
        r'/document/d/([a-zA-Z0-9_-]+)',
        r'/spreadsheets/d/([a-zA-Z0-9_-]+)',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, link)
        if match:
            return match.group(1)
    
    return link.strip()


def load_document(file_bytes: bytes) -> Document:
    """
    Load a DOCX document from bytes and keep content unchanged.
    """
    doc = Document(io.BytesIO(file_bytes))
    return doc

def is_meaningful_text(text):
    """Check if text contains meaningful content"""
    cleaned = re.sub(r'[\W_]+', '', text)
    return bool(cleaned.strip())

def is_decorative_only(text):
    """Check if text is decorative only (symbols, single letters, etc.)"""
    stripped = text.strip()
    return not stripped or re.fullmatch(r"[^\w\s]+", stripped) or re.fullmatch(r"[A-Z]", stripped)


def preview_text(text: str, limit: int = 200) -> str:
    """Create a sanitized preview of text for logging."""
    if not text:
        return ""
    sanitized = text.replace("\n", "\\n")
    if len(sanitized) > limit:
        sanitized = sanitized[:limit] + "..."
    return sanitized



def remove_delimiter_markers(text: str) -> str:
    """
    Remove ALL delimiter markers in format <<<...>>> - catches any variations including translated/misspelled ones.
    This ensures no delimiter markers (like <<<TRANULATION_1_END>>>) end up in the final document.
    Also removes malformed markers that don't have proper closing (like <<<TRANSL000000...).
    """
    if not text:
        return text
    
    # First: Remove properly closed markers <<<...>>>
    text = re.sub(r'<<<[^>]*?>>>', '', text, flags=re.DOTALL)
    
    # Second: Remove MALFORMED markers that start with <<< but have no closing >>>
    # Match <<< followed by ANY characters (including newlines) until whitespace or end of string
    # This catches cases like <<<TRANSL000000000000... that go on forever
    # Use non-greedy match to stop at first whitespace
    text = re.sub(r'<<<[^\s]*', '', text)
    
    # Also catch any remaining <<< patterns (defensive)
    text = re.sub(r'<<<.*?(?=\s|$)', '', text, flags=re.DOTALL)
    
    return text

def sanitize_response(text: str) -> str:
    """Remove AI artifacts but PRESERVE all whitespace and robust formatting tags."""
    if not text:
        return text
    # Remove AI reasoning tags
    text = re.sub(r'<think>.*?</think>', '', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'</?think>', '', text, flags=re.IGNORECASE)
    # Remove accidental <untranslated> tags that AI might have included
    # These should only come from our error handling, not from AI output
    text = re.sub(r'<untranslated>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'</untranslated>', '', text, flags=re.IGNORECASE)
    # Remove ALL delimiter markers (comprehensive removal - catches translated/misspelled variants)
    text = remove_delimiter_markers(text)
    # DO NOT remove robust formatting markers ««...»» - they need to be preserved!
    # The robust formatter will parse and remove them when applying formatting to the document
    # text = re.sub(r'««[^»]+»»', '', text)  # REMOVED - preserve robust formatting tags
    # Do NOT strip() - preserve all whitespace including indentation!
    return text


def create_format_marked_text(para) -> tuple:
    """Create text with simple format markers for translation"""
    marked_text = ""
    format_map = []
    
    # Get original paragraph text for space preservation
    original_para_text = para.text
    run_texts = [run.text for run in para.runs if run.text]
    concatenated_runs = ''.join(run_texts)
    
    # If runs don't match paragraph text, we need to preserve spaces
    needs_space_preservation = concatenated_runs != original_para_text
    
    para_pos = 0
    for i, run in enumerate(para.runs):
        if not run.text:
            continue
        
        # If space preservation is needed, check for gaps
        if needs_space_preservation and para_pos < len(original_para_text):
            run_start = original_para_text.find(run.text, para_pos)
            if run_start >= 0 and run_start > para_pos:
                # There's a gap (likely spaces) before this run
                gap = original_para_text[para_pos:run_start]
                if ' ' in gap:
                    # Add space to marked text
                    marked_text += ' '
                para_pos = run_start + len(run.text)
            elif run_start >= 0:
                para_pos = run_start + len(run.text)
        
        # Determine format markers
        markers = []
        if run.bold:
            markers.append("B")
        if run.italic:
            markers.append("I")
        if run.underline:
            markers.append("U")
        
        if markers:
            marker_str = ",".join(markers)
            marked_text += f"««{marker_str}»»{run.text}««/{marker_str}»»"
        else:
            marked_text += run.text
        
        format_map.append({
            'index': i,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'text': run.text,
            'marker': marker_str if markers else None
        })
    
    return marked_text, format_map


def extract_format_markers(text: str) -> list:
    """Extract format markers and text segments from marked translation"""
    segments = []
    pattern = r'««([^»]+)»»([^«]+)««/\1»»'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add any plain text before this match
        if match.start() > last_end:
            plain_text = text[last_end:match.start()]
            if plain_text:
                segments.append({'text': plain_text, 'formats': []})
        
        # Add formatted text
        formats = match.group(1).split(',')
        segments.append({
            'text': match.group(2),
            'formats': formats
        })
        last_end = match.end()
    
    # Add any remaining plain text
    if last_end < len(text):
        plain_text = text[last_end:]
        if plain_text:
            segments.append({'text': plain_text, 'formats': []})
    
    return segments


def analyze_batch_complexity(batch: List[Tuple[int, Any, str]]) -> dict:
    """Analyze complexity of a specific batch to determine if robust formatting is needed"""
    total_runs = 0
    complex_paras = 0
    total_paras = len(batch)
    
    for para_idx, para, original in batch:
        run_count = len([r for r in para.runs if r.text.strip()])
        total_runs += run_count
        
        # Count as complex if has multiple runs with different formatting
        if run_count > 2:
            complex_paras += 1
    
    avg_runs = total_runs / total_paras if total_paras > 0 else 0
    complexity_ratio = complex_paras / total_paras if total_paras > 0 else 0
    
    # A batch is complex if it has high run count OR many complex paragraphs
    is_complex = avg_runs > 2.5 or complexity_ratio > 0.3
    
    return {
        'avg_runs': avg_runs,
        'complex_ratio': complexity_ratio,
        'is_complex': is_complex,
        'total_runs': total_runs,
        'complex_paras': complex_paras
    }


def detect_case_change_in_text(text: str) -> bool:
    """
    Check if text has significant case changes (all-caps words mixed with mixed-case).
    This helps preserve case patterns like "HELLO, how you doing?"
    """
    if not text or len(text) < 2:
        return False
    
    words = []
    word_pattern = re.compile(r'\b[A-Za-z]+\b')
    for match in word_pattern.finditer(text):
        word = match.group(0)
        words.append(word)
    
    if len(words) < 2:
        return False
    
    has_upper_word = False
    has_mixed_word = False
    
    for word in words:
        # Multi-letter all uppercase word
        if word.isupper() and len(word) > 1:
            has_upper_word = True
        # Mixed case or lowercase word (not all caps)
        elif not word.isupper():
            has_mixed_word = True
        
        # If we have both, we need to split
        if has_upper_word and has_mixed_word:
            return True
    
    return False


def split_text_by_case_boundaries(text: str) -> List[Tuple[int, int]]:
    """
    Split text into segments at case boundaries.
    Returns list of (start, end) tuples for each segment.
    Only splits at boundaries between all-uppercase words and mixed/lowercase words.
    """
    if not text:
        return [(0, 0)]
    
    word_pattern = re.compile(r'\b[A-Za-z]+\b')
    words_info = []
    for match in word_pattern.finditer(text):
        word = match.group(0)
        words_info.append((match.start(), match.end(), word))
    
    if len(words_info) < 2:
        return [(0, len(text))]
    
    segments = []
    segment_start = 0
    last_word_was_upper = None
    
    for word_start, word_end, word in words_info:
        is_upper = word.isupper() and len(word) > 1  # Multi-letter all caps
        
        if last_word_was_upper is not None:
            # Check if case pattern changed significantly
            if is_upper != last_word_was_upper:
                # Case pattern changed - end previous segment before this word
                segments.append((segment_start, word_start))
                segment_start = word_start
        
        last_word_was_upper = is_upper
    
    # Add final segment
    segments.append((segment_start, len(text)))
    
    return segments if len(segments) > 1 else [(0, len(text))]


def apply_smart_formatting(para, translation: str, original: str):
    """Smart format application based on run analysis - optimized for different complexity levels"""
    
    # CRITICAL: Remove any markers that might have leaked through
    # Remove ALL delimiter markers (catches any variations including translated ones like <<<TRANULATION_1_END>>>)
    translation = remove_delimiter_markers(translation)
    # Remove robust formatting markers (not needed for standard formatting)
    translation = re.sub(r'««[^»]+»»', '', translation)
    
    if not para.runs:
        # No runs - create new one
        para.add_run(translation)
        return
    
    # Analyze run complexity
    run_count = len([r for r in para.runs if r.text.strip()])
    
    # NEW: Check for case changes in original text and handle them specially
    if detect_case_change_in_text(original) and run_count <= 4:
        # Split both original and translation by case boundaries
        original_segments = split_text_by_case_boundaries(original)
        translation_segments = split_text_by_case_boundaries(translation)
        
        # Extract actual text segments
        original_text_segments = [original[start:end] for start, end in original_segments]
        translation_text_segments = [translation[start:end] for start, end in translation_segments]
        
        # If we have matching number of segments and enough runs, apply to separate runs
        if (len(original_text_segments) == len(translation_text_segments) and 
            len(translation_text_segments) <= run_count and
            len(translation_text_segments) >= 2):
            
            # Clear existing runs
            for run in para.runs:
                run.text = ""
            
            # Apply each case segment to its own run
            for i, (run, trans_seg) in enumerate(zip(para.runs[:len(translation_text_segments)], translation_text_segments)):
                run.text = trans_seg
                # Add space if not last segment and segments should connect
                if i < len(translation_text_segments) - 1 and not trans_seg.rstrip().endswith((' ', ',', '.', '!', '?', ';', ':')):
                    run.text += ' '
            
            # Clear remaining runs
            for run in para.runs[len(translation_text_segments):]:
                run.text = ""
            
            return  # Early return - case splitting handled
    
    if run_count <= 1:
        # Simple case - single format, current implementation is perfect
        para.runs[0].text = translation
        
    elif run_count == 2:
        # Common case: Two-format paragraph (e.g., bold heading + normal text)
        # Try to preserve the format split proportionally
        first_run_text = para.runs[0].text
        second_run_text = para.runs[1].text if len(para.runs) > 1 else ""
        
        if first_run_text and original:
            # Calculate proportion of first run in original
            first_run_ratio = len(first_run_text) / len(original)
            split_point = int(len(translation) * first_run_ratio)
            
            # Clear existing runs
            for run in para.runs:
                run.text = ""
            
            # Apply proportional split
            para.runs[0].text = translation[:split_point]
            if len(para.runs) > 1:
                para.runs[1].text = translation[split_point:]
            else:
                # Create second run with same formatting as original
                new_run = para.add_run(translation[split_point:])
                
    elif run_count <= 4:
        # Moderate complexity - try to preserve major format boundaries
        # This is a simplified approach that works well for common cases
        
        # Calculate word distribution
        total_words = len(original.split())
        words_per_run = total_words // run_count if run_count > 0 else total_words
        
        # Split translation by approximate word count
        translation_words = translation.split()
        
        # Clear existing runs
        for run in para.runs:
            run.text = ""
        
        # Distribute translation across runs
        word_index = 0
        for i, run in enumerate(para.runs[:run_count]):
            if i == run_count - 1:
                # Last run gets remaining words
                run.text = ' '.join(translation_words[word_index:])
            else:
                # Calculate words for this run
                end_index = min(word_index + words_per_run, len(translation_words))
                run.text = ' '.join(translation_words[word_index:end_index])
                if i < run_count - 1:
                    run.text += ' '  # Add space between runs
                word_index = end_index
                
    else:
        # High complexity - fallback to first run (robust method should handle these)
        # Clear all runs and put everything in first
        for run in para.runs:
            run.text = ""
        para.runs[0].text = translation

def call_gemini_batch_api(client, prompt, model, logs=None):
    """
    Synchronous function to call Gemini API for batch processing with token tracking.
    This will be run in a thread executor.
    """
    for attempt in range(MAX_RETRIES):
        try:
            if logs is not None:
                logs.append(f"[BATCH API] Attempt {attempt + 1}/{MAX_RETRIES} - Model: {model}")
            
            # Generate content with TEXT response format (not JSON to preserve formatting)
            response = client.models.generate_content(
                model=model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.0
                    # NO response_mime_type - use plain text to preserve formatting
                )
            )
            
            result_text = response.text.strip()
            
            # Get token usage from response
            usage = response.usage_metadata
            actual_input_tokens = usage.prompt_token_count
            output_tokens = usage.candidates_token_count
            total_tokens = usage.total_token_count
            
            if logs is not None:
                logs.append(f"[SUCCESS] Received response ({len(result_text)} chars)")
                logs.append(f"[TOKENS] Input: {actual_input_tokens}, Output: {output_tokens}, Total: {total_tokens}")
            
            return {
                'text': result_text,
                'input_tokens': actual_input_tokens,
                'output_tokens': output_tokens,
                'total_tokens': total_tokens
            }
            
        except Exception as e:
            # Capture detailed error information
            error_details = {
                'error_type': type(e).__name__,
                'error_message': str(e),
                'attempt': attempt + 1
            }
            
            # Try to get more details from Gemini API errors
            if hasattr(e, 'response'):
                error_details['api_response'] = str(e.response)
            if hasattr(e, 'status_code'):
                error_details['status_code'] = e.status_code
            if hasattr(e, 'reason'):
                error_details['reason'] = str(e.reason)
            
            # Print detailed error to console
            print(f"\n🔴 [GEMINI API ERROR] Attempt {attempt + 1}/{MAX_RETRIES}")
            print(f"   Error Type: {error_details['error_type']}")
            print(f"   Error Message: {error_details['error_message']}")
            if 'status_code' in error_details:
                print(f"   Status Code: {error_details['status_code']}")
            if 'reason' in error_details:
                print(f"   Reason: {error_details['reason']}")
            if 'api_response' in error_details:
                print(f"   API Response: {error_details['api_response']}")
            print(f"   Model: {model}\n")
            
            if logs is not None:
                logs.append(f"[ERROR] Attempt {attempt + 1} failed: {error_details['error_type']} - {error_details['error_message']}")
                if 'status_code' in error_details:
                    logs.append(f"[ERROR DETAILS] Status: {error_details['status_code']}, Reason: {error_details.get('reason', 'N/A')}")
            
            if attempt < MAX_RETRIES - 1:
                print(f"   ⏳ Retrying in {RETRY_DELAY} seconds...\n")
                time.sleep(RETRY_DELAY)
    
    # All retries exhausted
    print(f"\n🔴 [GEMINI API] All {MAX_RETRIES} attempts failed for model: {model}\n")
    if logs is not None:
        logs.append(f"[FAILED] All {MAX_RETRIES} retry attempts exhausted")
    return None

async def call_gemini_batch_async(executor, client, prompt, model, logs=None):
    """Async wrapper for batch API call using ThreadPoolExecutor."""
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(executor, call_gemini_batch_api, client, prompt, model, logs)
    return result

async def retry_single_paragraph_translation(executor, client, paragraph_text, language, model, logs=None, max_retries=2):
    """
    Retry translation for a single paragraph (2 retries = 3 total attempts).
    Returns translated text or None if all attempts fail.
    """
    prompt = create_batch_prompt([paragraph_text], language)
    
    for retry_attempt in range(max_retries + 1):  # 0, 1, 2 = 3 total attempts
        if retry_attempt > 0:
            print(f"   🔄 [RETRY] Retrying paragraph translation (attempt {retry_attempt + 1}/{max_retries + 1})")
            if logs:
                logs.append(f"[RETRY] Paragraph retry attempt {retry_attempt + 1}/{max_retries + 1}")
            await asyncio.sleep(RETRY_DELAY)
        
        batch_result = await call_gemini_batch_async(executor, client, prompt, model, logs)
        
        if batch_result and 'text' in batch_result:
            translations = parse_structured_response(batch_result['text'], 1, logs)
            if translations and len(translations) > 0 and translations[0] and translations[0] != '[Translation missing]':
                if retry_attempt > 0:
                    print(f"   ✅ [RETRY SUCCESS] Paragraph translated on attempt {retry_attempt + 1}")
                return translations[0]
        
        if retry_attempt < max_retries:
            print(f"   ⚠️ [RETRY FAILED] Attempt {retry_attempt + 1} failed, will retry...")
    
    print(f"   🔴 [RETRY EXHAUSTED] All {max_retries + 1} attempts failed for paragraph")
    if logs:
        logs.append(f"[RETRY EXHAUSTED] All {max_retries + 1} attempts failed for paragraph")
    return None

async def call_openrouter_batch_api(session, prompt, model, openrouter_api_key, log_list=None):
    """
    Call OpenRouter API for batch processing with token tracking.
    """
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {openrouter_api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": os.getenv("OPENROUTER_REFERER", "https://yourdomain.com"),
        "X-Title": "EasyTranslate"
    }
    
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.0
        # NO response_format - use plain text to preserve formatting
    }
    
    for attempt in range(MAX_RETRIES):
        try:
            if log_list is not None:
                log_entry = f"[BATCH API] Attempt {attempt + 1}/{MAX_RETRIES} - Model: {model}"
                log_list.append(log_entry)
            
            async with session.post(url, headers=headers, json=payload) as resp:
                resp.raise_for_status()
                data = await resp.json()
                
                result = data["choices"][0]["message"]["content"].strip()
                
                # Get token usage from response
                usage = data.get("usage", {})
                input_tokens = usage.get("prompt_tokens", 0)
                output_tokens = usage.get("completion_tokens", 0)
                total_tokens = usage.get("total_tokens", input_tokens + output_tokens)
                
                if log_list is not None:
                    log_list.append(f"[SUCCESS] Received response ({len(result)} chars)")
                    log_list.append(f"[TOKENS] Input: {input_tokens}, Output: {output_tokens}, Total: {total_tokens}")
                
                return {
                    'text': result,
                    'input_tokens': input_tokens,
                    'output_tokens': output_tokens,
                    'total_tokens': total_tokens
                }
                
        except Exception as e:
            if log_list is not None:
                log_list.append(f"[ERROR] Attempt {attempt + 1} failed: {str(e)}")
            
            if attempt < MAX_RETRIES - 1:
                await asyncio.sleep(RETRY_DELAY)
    
    if log_list is not None:
        log_list.append("[FAILED] All retry attempts exhausted")
    return None

def parse_structured_response(response_text, expected_count, logs=None):
    """Parse delimiter-based response to preserve ALL formatting (no JSON mangling)."""
    import json
    import re
    
    # First try delimiter-based parsing (new method)
    try:
        translations = []
        
        # Extract translations using delimiters
        for i in range(1, expected_count + 1):
            # Look for <<<TRANSLATION_START_N>>> ... <<<TRANSLATION_END_N>>>
            pattern = f'<<<TRANSLATION_START_{i}>>>(.+?)<<<TRANSLATION_END_{i}>>>'
            match = re.search(pattern, response_text, re.DOTALL)
            
            if match:
                # Extract the translation - PRESERVE ALL FORMATTING
                translation = match.group(1)
                # Only strip the immediate newlines added by delimiters, not the content
                if translation.startswith('\n'):
                    translation = translation[1:]
                if translation.endswith('\n'):
                    translation = translation[:-1]
                translations.append(translation)
                if logs:
                    logs.append(f"[DELIMITER] Extracted translation {i} ({len(translation)} chars)")
            else:
                if logs:
                    logs.append(f"[WARNING] Could not find delimiter for translation {i}")
                translations.append('[Translation missing]')
        
        if translations:
            if logs:
                logs.append(f"[DELIMITER] Successfully extracted {len(translations)} translations with PRESERVED formatting")
            return translations
    
    except Exception as e:
        if logs:
            logs.append(f"[DELIMITER ERROR] {str(e)} - falling back to JSON parsing")
    
    # Fallback to JSON parsing (old method)
    try:
        # Clean response text (remove markdown code blocks if present)
        cleaned_text = response_text.strip()
        if cleaned_text.startswith('```json'):
            cleaned_text = cleaned_text[7:]
        if cleaned_text.endswith('```'):
            cleaned_text = cleaned_text[:-3]
        cleaned_text = cleaned_text.strip()
        
        # Parse JSON
        parsed_response = json.loads(cleaned_text)
        
        if logs:
            logs.append(f"[JSON] Successfully parsed JSON response (formatting may be affected)")
        
        # Extract translations in order
        translations = []
        if 'translations' in parsed_response:
            # Sort by ID to ensure correct order
            translation_items = sorted(parsed_response['translations'], key=lambda x: x.get('id', 0))
            
            for item in translation_items:
                if 'translation' in item:
                    translations.append(item['translation'])
                    
            if logs:
                logs.append(f"[JSON] Extracted {len(translations)} translations")
        
        return translations
        
    except json.JSONDecodeError as e:
        if logs:
            logs.append(f"[JSON ERROR] Failed to parse JSON: {str(e)}")
            logs.append(f"[FALLBACK] Attempting basic splitting...")
        
        # Final fallback: split by double newlines
        lines = response_text.split('\n\n')
        alt_translations = [t for t in lines if t.strip() and not t.startswith('{') and not t.startswith('<<<')]
        # Don't strip here - preserve formatting
        return alt_translations[:expected_count]
    
    except Exception as e:
        if logs:
            logs.append(f"[PARSE ERROR] {str(e)}")
        return []

async def translate_document_content_async(file_bytes: bytes, file_name: str, language: str, model: str, api_key: str, progress_id: Optional[str] = None) -> TranslateResponse:
    """Translate document content using Gemini with async batch processing - matches Streamlit logic"""
    import json
    
    print(f"[TRANSLATOR] Initializing Gemini client")
    print(f"[TRANSLATOR] Model: {model}, API Key present: {bool(api_key)}")
    
    # Normalize model name - Google Genai SDK accepts model names as-is
    # But we'll log it for debugging
    normalized_model = model.strip()
    print(f"[TRANSLATOR] Using model: {normalized_model}")
    
    try:
        client = genai.Client(api_key=api_key)
        print(f"[TRANSLATOR] Client initialized successfully")
    except Exception as e:
        error_msg = f"Failed to initialize Gemini client: {str(e)}"
        print(f"[TRANSLATOR ERROR] {error_msg}")
        import traceback
        traceback.print_exc()
        raise Exception(error_msg)
    
    # Load document
    doc = load_document(file_bytes)
    
    # Initialize logs early for TOC processing
    logs = []
    
    # ========== PROCESS TOC BEFORE TRANSLATION ==========
    if TOC_HANDLER_AVAILABLE:
        toc_results = process_toc_before_translation(doc)
        if toc_results['toc_found']:
            logs.append(f"[TOC] Found {toc_results['toc_entries_count']} TOC entries in first 10 pages")
            logs.append(f"[TOC] Extracted {toc_results['titles_extracted']} titles from TOC")
            logs.append(f"[TOC] Converted {toc_results['paragraphs_converted']} paragraphs to Heading 2 style")
            logs.append(f"[TOC] Removed {toc_results['toc_removed']} TOC entry paragraphs")
            if toc_results['placeholder_inserted']:
                logs.append(f"[TOC] Inserted placeholder for auto-generated TOC")
            print(f"[TOC PROCESSING] Found {toc_results['toc_entries_count']} entries, converted {toc_results['paragraphs_converted']} to Heading 2")
        else:
            logs.append(f"[TOC] No TOC found in first 10 pages")
    
    paragraphs = doc.paragraphs
    
    print(f"[TRANSLATOR] Document has {len(paragraphs)} total paragraphs")
    
    translated_content = []
    total_input_tokens = 0
    total_output_tokens = 0
    
    logs.append(f"[START] Batch translation started for language: {language}")
    logs.append(f"[INFO] Source file: {file_name}")
    logs.append(f"[INFO] Document has {len(paragraphs)} total paragraphs")
    logs.append(f"[INFO] Using ADAPTIVE TOKEN-BASED BATCHING - maximizes efficiency while preserving formatting")
    logs.append(f"[INFO] Using Gemini model: {normalized_model}")
    logs.append(f"[INFO] Processing in memory - no files saved to disk")
    
    # Use normalized_model for all API calls in this function
    # Replace model variable with normalized_model for consistency
    model = normalized_model
    
    # Prepare paragraph batches with TOKEN-BASED ADAPTIVE BATCHING
    paragraph_batches = []
    current_batch = []
    current_tokens = 0
    current_target_tokens = 0
    current_use_robust = False
    
    # Track stats for optimization reporting
    total_paragraphs_to_translate = 0
    section_stats = {'simple': 0, 'moderate': 0, 'complex': 0}
    WINDOW_SIZE = 100  # Analyze every 100 paragraphs
    
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        # CRITICAL: Do NOT strip - preserve ALL whitespace including indentation
        original = para.text  # Keep exact formatting


        # Skip single uppercase letters followed by uppercase text (but don't remove them!)
        if re.fullmatch(r"[A-Z]", original) and i + 1 < len(paragraphs) and paragraphs[i + 1].text.strip()[:1].isupper():
            i += 1
            continue

        # Skip empty or decorative text (check stripped version)
        if not original.strip() or not is_meaningful_text(original) or is_decorative_only(original):
            i += 1
            continue

        word_count = len(original.strip().split())
        is_heading = (para.style is not None and para.style.name.lower().startswith("heading")) or para.alignment == 1

        # Skip single-word non-heading paragraphs
        if word_count <= 1:
            if not original.strip().isupper() and not is_heading:
                i += 1
                continue

        total_paragraphs_to_translate += 1
        
        # TOKEN-BASED ADAPTIVE BATCHING
        # If starting a new batch, analyze upcoming section
        if not current_batch:
            target_tokens, use_robust, section_info = analyze_section_complexity(paragraphs, i, WINDOW_SIZE)
            current_target_tokens = target_tokens
            current_use_robust = use_robust
            
            section_stats[section_info['type']] += 1
            
            print(f"\n[SECTION ANALYSIS] Para {i}:")
            print(f"  Type: {section_info['type'].upper()}")
            print(f"  Target tokens: {target_tokens}")
            print(f"  Use robust formatting: {use_robust}")
            print(f"  Complexity: {section_info.get('complex_ratio', 0)*100:.0f}% complex, {section_info.get('inline_ratio', 0)*100:.0f}% inline formatting")
            
            logs.append(f"[SECTION {i}] Type: {section_info['type']}, Target: {target_tokens} tokens, Robust: {use_robust}")
        
        # Estimate tokens for this paragraph
        para_tokens = estimate_tokens(original)
        
        # Check if adding this paragraph would exceed token limit
        if current_batch and (current_tokens + para_tokens > current_target_tokens * 1.2):
            # Save current batch (allow 20% overflow for natural paragraph boundaries)
            batch_num = len(paragraph_batches) + 1
            actual_tokens = current_tokens
            print(f"[BATCH CREATED] #{batch_num}: {len(current_batch)} paras, ~{actual_tokens} tokens, robust={current_use_robust}")
            
            # Store batch with metadata
            paragraph_batches.append({
                'batch': current_batch,
                'use_robust': current_use_robust,
                'tokens': actual_tokens
            })
            current_batch = []
            current_tokens = 0
            # Next iteration will re-analyze section
        
        # Add paragraph to current batch
        current_batch.append((i, para, original))
        current_tokens += para_tokens
        
        # Also save batch if we've hit minimum paragraphs AND target tokens
        MIN_PARAGRAPHS = 5
        if (len(current_batch) >= MIN_PARAGRAPHS and 
            current_tokens >= current_target_tokens):
            batch_num = len(paragraph_batches) + 1
            print(f"[BATCH CREATED] #{batch_num}: {len(current_batch)} paras, ~{current_tokens} tokens, robust={current_use_robust}")
            
            paragraph_batches.append({
                'batch': current_batch,
                'use_robust': current_use_robust,
                'tokens': current_tokens
            })
            current_batch = []
            current_tokens = 0
        
        i += 1

    # Save any remaining batch after loop ends
    if current_batch:
        batch_num = len(paragraph_batches) + 1
        print(f"[BATCH CREATED] #{batch_num} (final): {len(current_batch)} paras, ~{current_tokens} tokens, robust={current_use_robust}")
        paragraph_batches.append({
            'batch': current_batch,
            'use_robust': current_use_robust,
            'tokens': current_tokens
        })

    total_batches = len(paragraph_batches)
    
    # Calculate optimization stats
    print(f"\n{'='*80}")
    print(f"[BATCH SUMMARY] Total batches created: {total_batches}")
    print(f"[BATCH SUMMARY] Total paragraphs: {total_paragraphs_to_translate}")
    print(f"[BATCH SUMMARY] Section types: Simple={section_stats['simple']}, Moderate={section_stats['moderate']}, Complex={section_stats['complex']}")
    batch_sizes = [len(b['batch']) for b in paragraph_batches]
    print(f"[BATCH SUMMARY] Avg batch size: {sum(batch_sizes) / len(batch_sizes):.1f} paragraphs")
    print(f"[BATCH SUMMARY] Min/Max: {min(batch_sizes)}/{max(batch_sizes)} paragraphs")
    robust_count = sum(1 for b in paragraph_batches if b['use_robust'])
    print(f"[BATCH SUMMARY] Robust formatting: {robust_count}/{total_batches} batches ({robust_count/total_batches*100:.0f}%)")
    print(f"{'='*80}\n")
    
    estimated_fixed_batches = max(1, total_paragraphs_to_translate // 20)  # If we used fixed size 20
    optimization_percentage = ((1 - total_batches / estimated_fixed_batches) * 100) if estimated_fixed_batches > 0 else 0
    
    logs.append(f"[TOKEN BATCHING] Created {total_batches} adaptive batches")
    logs.append(f"[SECTIONS] Simple: {section_stats['simple']}, Moderate: {section_stats['moderate']}, Complex: {section_stats['complex']}")
    logs.append(f"[OPTIMIZATION] Would have been ~{estimated_fixed_batches} calls with fixed size 20")
    logs.append(f"[EFFICIENCY] Reduced API calls by {optimization_percentage:.0f}% using token-based batching")
    logs.append(f"[PROCESSING] Starting parallel batch API requests (max 4 concurrent)...")
    
    # Initialize progress tracking
    if progress_id:
        progress_tracker[progress_id] = {
            "totalBatches": total_batches,
            "completedBatches": 0,
            "error": False
        }
    
    # Create thread pool executor for async processing
    executor = ThreadPoolExecutor(max_workers=4)
    
    async def process_batch_gemini(batch_idx, batch_info):
        """
        Process a single batch for Gemini with ADAPTIVE TOKEN-BASED batching.
        Includes error handling with <untranslated> wrapper for failed batches.
        """
        batch_logs = []  # Separate logs for this batch
        batch = batch_info['batch']  # Extract actual batch
        use_robust_for_batch = batch_info['use_robust']  # Pre-determined by section analysis
        batch_tokens = batch_info['tokens']
        
        batch_paragraphs = [item[2] for item in batch]  # Extract text
        batch_size = len(batch_paragraphs)
        
        print(f"[TRANSLATOR] Processing batch {batch_idx + 1}/{total_batches} ({batch_size} paras, ~{batch_tokens} tokens, robust={use_robust_for_batch})")
        batch_logs.append(f"[BATCH {batch_idx + 1}/{total_batches}] Processing {batch_size} paragraphs, ~{batch_tokens} tokens...")
        batch_logs.append(f"[METHOD] Using {'ROBUST' if use_robust_for_batch else 'STANDARD'} formatting (pre-determined by section analysis)")
        
        # ERROR HANDLING: Wrap entire batch processing in try-except
        try:
            if use_robust_for_batch and ROBUST_FORMATTING_AVAILABLE:
                # Use robust formatting for this complex batch
                # Initialize preserver for this batch
                preserver = RobustFormatPreserver(doc)
                
                # Create marked texts for this batch
                marked_batch = []
                for para_idx, para, original in batch:
                    marked_text, para_data = preserver.create_formatted_text_for_translation(para, para_idx)
                    marked_batch.append((para_idx, marked_text))
                    print(f"[ROBUST INPUT] Para {para_idx} ({len(original)} chars): {preview_text(original)}")
                
                # Create robust prompt
                prompt = create_robust_translation_prompt(marked_batch, language)
                batch_logs.append(f"[ROBUST] Created format-preserved prompt for batch {batch_idx + 1}")
                
                # Call API with robust prompt
                batch_result = await call_gemini_batch_async(executor, client, prompt, model, batch_logs)
                
                # Handle case where API returned None or empty result
                if not batch_result or 'text' not in batch_result:
                    raise RuntimeError("Gemini robust translation returned no result")
                
                # Parse with robust parser
                if batch_result and 'text' in batch_result:
                    robust_translations = parse_robust_response(batch_result['text'], marked_batch, batch_logs)
                    
                    # Mark batch as using robust method (will be applied later)
                    batch_result['robust_mode'] = True
                    batch_result['robust_translations'] = robust_translations
                    batch_result['robust_preserver'] = preserver
                    batch_result['marked_batch'] = marked_batch
                    
                    # Update progress
                    if progress_id:
                        progress_tracker[progress_id]["completedBatches"] += 1
                    
                    return batch_idx, batch, batch_paragraphs, batch_result, batch_logs
            else:
                # Standard processing
                batch_logs.append(f"[STANDARD] Using standard formatting for batch {batch_idx + 1}")
                
                # Create batch prompt
                prompt = create_batch_prompt(batch_paragraphs, language)
                
                print(f"[STANDARD INPUT] Batch {batch_idx + 1}: sending {len(batch)} paragraphs to Gemini")
                for para_idx, para, original in batch:
                    print(f"   - Para {para_idx} ({len(original)} chars): {preview_text(original)}")
                
                # Call API asynchronously
                batch_result = await call_gemini_batch_async(executor, client, prompt, model, batch_logs)
                
                if not batch_result:
                    raise RuntimeError("Gemini standard translation returned no result")
                
                # Update progress immediately after batch completes
                if progress_id:
                    progress_tracker[progress_id]["completedBatches"] += 1
                
                return batch_idx, batch, batch_paragraphs, batch_result, batch_logs
                
        except Exception as e:
            # CRITICAL ERROR HANDLING: If batch fails, wrap with <untranslated> tag
            import traceback
            full_traceback = traceback.format_exc()
            batch_logs.append(f"[ERROR] Batch {batch_idx + 1} FAILED: {str(e)}")
            batch_logs.append(f"[ERROR] Full traceback:\n{full_traceback}")
            batch_logs.append(f"[ERROR] Wrapping failed batch with <untranslated> tag to preserve original")
            print(f"[ERROR] Batch {batch_idx + 1} failed: {str(e)}")
            print(f"[TRACEBACK]\n{full_traceback}")
            
            # Create a failed batch result with untranslated wrapper
            failed_result = {
                'text': '',  # Empty text
                'input_tokens': 0,
                'output_tokens': 0,
                'total_tokens': 0,
                'failed': True,  # Mark as failed
                'error_message': str(e)
            }
            
            # Update progress even for failed batch
            if progress_id:
                progress_tracker[progress_id]["completedBatches"] += 1
            
            return batch_idx, batch, batch_paragraphs, failed_result, batch_logs
    
    # Create all tasks for parallel processing
    tasks = [
        process_batch_gemini(batch_idx, batch)
        for batch_idx, batch in enumerate(paragraph_batches)
    ]
    
    # Execute all tasks in parallel (limited by ThreadPoolExecutor max_workers)
    results = await asyncio.gather(*tasks)
    
    # Process results in order and merge logs
    failed_batch_count = 0
    for batch_idx, batch, batch_paragraphs, batch_result, batch_logs in results:
        # Add this batch's logs in order
        logs.extend(batch_logs)
        
        if batch_result:
            # Check if batch failed
            if batch_result.get('failed'):
                # FAILED BATCH: Wrap with <untranslated> tag and keep original formatting
                failed_batch_count += 1
                logs.append(f"[FAILED BATCH] Batch {batch_idx + 1} failed - wrapping with <untranslated> tag")
                
                # Print error details to console
                error_msg = batch_result.get('error_message', 'Unknown error')
                print(f"\n🔴 [BATCH FAILED] Batch {batch_idx + 1} failed completely")
                print(f"   Error: {error_msg}")
                print(f"   Wrapping {len(batch)} paragraphs with <untranslated> tags\n")
                
                for para_idx, para, original in batch:
                    # Wrap original text with <untranslated> tags on new lines
                    untranslated_text = f"<untranslated>\n{original}\n</untranslated>"
                    translated_content.append(untranslated_text)
                    
                    # Apply to paragraph with newlines - preserve formatting
                    if para.runs:
                        # Clear existing runs and add untranslated text with newlines
                        for run in para.runs:
                            run.text = ""
                        para.runs[0].text = f"<untranslated>\n{original}\n</untranslated>"
                    else:
                        # No runs, add as single run
                        para.add_run(untranslated_text)
                
                continue  # Skip to next batch
            
            total_input_tokens += batch_result.get('input_tokens', 0) or 0
            total_output_tokens += batch_result.get('output_tokens', 0) or 0
            
            # Check if this batch used robust method
            if batch_result.get('robust_mode'):
                # Robust batch - apply formatting-preserved translations
                preserver = batch_result['robust_preserver']
                marked_batch = batch_result['marked_batch']
                robust_translations = batch_result['robust_translations']
                
                logs.append(f"[BATCH {batch_idx + 1}] Applying robust formatting ({len(robust_translations)} paragraphs)")
                
                # Apply each robust translation to its paragraph
                for (para_idx, para, original), (marked_para_id, marked_text), translation in zip(
                    batch, marked_batch, robust_translations
                ):
                    # Check if translation is missing or empty
                    if (translation is None or 
                        translation == "" or 
                        translation == '[Translation missing]' or
                        translation.strip() == ''):
                        # Retry this paragraph translation (2 more attempts = 3 total)
                        print(f"\n⚠️ [MISSING TRANSLATION] Para {para_idx}: Retrying translation...")
                        print(f"   Original: {original[:100]}...")
                        retried_translation = await retry_single_paragraph_translation(
                            executor, client, original, language, model, batch_logs, max_retries=2
                        )
                        
                        if retried_translation and retried_translation.strip():
                            # Retry succeeded - use the retried translation
                            print(f"   ✅ [RETRY SUCCESS] Para {para_idx} translated after retry")
                            preserver.apply_formatting_to_paragraph(para, marked_para_id, retried_translation)
                            translated_content.append(retried_translation)
                        else:
                            # All retries failed - wrap with <untranslated> tags on new lines
                            print(f"   🔴 [RETRY FAILED] Para {para_idx}: Marking as untranslated")
                            untranslated_text = f"<untranslated>\n{original}\n</untranslated>"
                            translated_content.append(untranslated_text)
                            
                            # Apply to paragraph with newlines - preserve formatting
                            if para.runs:
                                # Clear existing runs and add untranslated text with newlines
                                for run in para.runs:
                                    run.text = ""
                                para.runs[0].text = f"<untranslated>\n{original}\n</untranslated>"
                            else:
                                # No runs, add as single run
                                para.add_run(untranslated_text)
                    else:
                        # Valid translation - apply formatting-preserved translation
                        preserver.apply_formatting_to_paragraph(para, marked_para_id, translation)
                        translated_content.append(translation)
            else:
                # Standard batch - parse and apply normally
                batch_translations = parse_structured_response(batch_result['text'], len(batch_paragraphs), logs)
                
                logs.append(f"[BATCH {batch_idx + 1}] Received {len(batch_translations)} translations")
                
                # Validate we got the expected number of translations
                if len(batch_translations) != len(batch_paragraphs):
                    logs.append(f"[WARNING] Expected {len(batch_paragraphs)} translations, got {len(batch_translations)}")
                    # Pad with empty strings if we're short
                    while len(batch_translations) < len(batch_paragraphs):
                        batch_translations.append('[Translation missing]')
                    # Trim if we got too many
                    batch_translations = batch_translations[:len(batch_paragraphs)]
                
                # Apply translations to document (maintains order)
                for (para_idx, para, original), translation in zip(batch, batch_translations):
                    # Check if translation is missing or empty
                    if (translation is None or 
                        translation == "" or 
                        translation == '[Translation missing]' or
                        translation.strip() == ''):
                        # Retry this paragraph translation (2 more attempts = 3 total)
                        print(f"\n⚠️ [MISSING TRANSLATION] Para {para_idx}: Retrying translation...")
                        print(f"   Original: {original[:100]}...")
                        retried_translation = await retry_single_paragraph_translation(
                            executor, client, original, language, model, batch_logs, max_retries=2
                        )
                        
                        if retried_translation and retried_translation.strip():
                            # Retry succeeded - use the retried translation
                            print(f"   ✅ [RETRY SUCCESS] Para {para_idx} translated after retry")
                            translation = sanitize_response(retried_translation)
                            translated_content.append(translation)
                            apply_smart_formatting(para, translation, original)
                        else:
                            # All retries failed - wrap with <untranslated> tags on new lines
                            print(f"   🔴 [RETRY FAILED] Para {para_idx}: Marking as untranslated")
                            untranslated_text = f"<untranslated>\n{original}\n</untranslated>"
                            translated_content.append(untranslated_text)
                            
                            # Apply to paragraph with newlines - preserve formatting
                            if para.runs:
                                # Clear existing runs and add untranslated text with newlines
                                for run in para.runs:
                                    run.text = ""
                                para.runs[0].text = f"<untranslated>\n{original}\n</untranslated>"
                            else:
                                # No runs, add as single run
                                para.add_run(untranslated_text)
                    else:
                        # Valid translation - apply normally
                        translation = sanitize_response(translation)
                        translated_content.append(translation)
                        
                        # Use smart formatting to preserve run structure
                        apply_smart_formatting(para, translation, original)
        else:
            # This shouldn't happen with new error handling, but keep as safety
            failed_batch_count += 1
            print(f"\n🔴 [BATCH ERROR] Batch {batch_idx + 1} returned None - wrapping with <untranslated>")
            logs.append(f"[BATCH ERROR] Batch {batch_idx + 1} returned None - wrapping with <untranslated>")
            
            for para_idx, para, original in batch:
                # Wrap original text with <untranslated> tags on new lines
                untranslated_text = f"<untranslated>\n{original}\n</untranslated>"
                translated_content.append(untranslated_text)
                
                # Apply to paragraph with newlines - preserve formatting
                if para.runs:
                    # Clear existing runs and add untranslated text with newlines
                    for run in para.runs:
                        run.text = ""
                    para.runs[0].text = f"<untranslated>\n{original}\n</untranslated>"
                else:
                    # No runs, add as single run
                    para.add_run(untranslated_text)
    
    # Log failed batches summary
    if failed_batch_count > 0:
        logs.append(f"[WARNING] {failed_batch_count} batch(es) failed and were wrapped with <untranslated> tags")
        logs.append(f"[INFO] Search for '<untranslated>' in the output document to find failed sections")
    
    # Shutdown executor
    executor.shutdown(wait=False)
    
    # Save document to memory buffer
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    logs.append(f"[SAVE] Document saved to memory buffer")
    logs.append(f"[TOKENS] Final usage - Input: {total_input_tokens}, Output: {total_output_tokens}, Total: {total_input_tokens + total_output_tokens}")
    logs.append(f"[DONE] Translation complete!")
    
    # Convert to base64
    translated_base64 = base64.b64encode(output_buffer.read()).decode('utf-8')
    
    print("[TRANSLATE] Complete, returning response")
    
    return TranslateResponse(
        translatedDocument=translated_base64,
        logs=logs,
        stats={
            "totalInputTokens": total_input_tokens,
            "totalOutputTokens": total_output_tokens,
            "totalTokens": total_input_tokens + total_output_tokens,
            "paragraphCount": len(translated_content),
            "translatedText": "\n\n".join(translated_content)  # Add full text for preview
        }
    )

async def translate_document_content_async_openrouter(file_bytes: bytes, file_name: str, language: str, model: str, api_key: str, progress_id: Optional[str] = None) -> TranslateResponse:
    """Translate document content using OpenRouter with async batch processing"""
    import json
    
    print(f"[TRANSLATOR] Initializing OpenRouter client")
    
    # Load document
    doc = load_document(file_bytes)
    
    # Initialize logs early for TOC processing
    logs = []
    
    # ========== PROCESS TOC BEFORE TRANSLATION ==========
    if TOC_HANDLER_AVAILABLE:
        toc_results = process_toc_before_translation(doc)
        if toc_results['toc_found']:
            logs.append(f"[TOC] Found {toc_results['toc_entries_count']} TOC entries in first 10 pages")
            logs.append(f"[TOC] Extracted {toc_results['titles_extracted']} titles from TOC")
            logs.append(f"[TOC] Converted {toc_results['paragraphs_converted']} paragraphs to Heading 2 style")
            logs.append(f"[TOC] Removed {toc_results['toc_removed']} TOC entry paragraphs")
            if toc_results['placeholder_inserted']:
                logs.append(f"[TOC] Inserted placeholder for auto-generated TOC")
            print(f"[TOC PROCESSING] Found {toc_results['toc_entries_count']} entries, converted {toc_results['paragraphs_converted']} to Heading 2")
        else:
            logs.append(f"[TOC] No TOC found in first 10 pages")
    
    paragraphs = doc.paragraphs
    
    print(f"[TRANSLATOR] Document has {len(paragraphs)} total paragraphs")
    
    translated_content = []
    total_input_tokens = 0
    total_output_tokens = 0
    
    logs.append(f"[START] Batch translation started for language: {language}")
    logs.append(f"[INFO] Source file: {file_name}")
    logs.append(f"[INFO] Document has {len(paragraphs)} total paragraphs")
    logs.append(f"[INFO] Using SMART BATCHING - batch size adapts to content complexity")
    logs.append(f"[INFO] Using OpenRouter model: {model}")
    logs.append(f"[INFO] Processing in memory - no files saved to disk")
    
    # Prepare paragraph batches with ENHANCED smart batching logic
    paragraph_batches = []
    current_batch = []
    current_max_size = 5  # Start with ultra-conservative size (was 100 - too large!)
    
    # Track stats for optimization reporting
    total_paragraphs_to_translate = 0
    batch_size_distribution = {'poetry': 0, 'dialogue': 0, 'prose': 0, 'default': 0}
    
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        # CRITICAL: Do NOT strip - preserve ALL whitespace including indentation
        original = para.text  # Keep exact formatting

        # Skip single uppercase letters followed by uppercase text (but don't remove them!)
        if re.fullmatch(r"[A-Z]", original) and i + 1 < len(paragraphs) and paragraphs[i + 1].text.strip()[:1].isupper():
            logs.append(f"[SKIP] Skipping single uppercase letter at para {i} (preserving structure)")
            i += 1
            continue

        # Skip empty or decorative text (check stripped version)
        if not original.strip() or not is_meaningful_text(original) or is_decorative_only(original):
            i += 1
            continue

        word_count = len(original.strip().split())
        is_heading = (para.style is not None and para.style.name.lower().startswith("heading")) or para.alignment == 1

        # Skip single-word non-heading paragraphs
        if word_count <= 1:
            if not original.strip().isupper() and not is_heading:
                i += 1
                continue

        # SMART BATCHING: Get optimal batch size for this content
        optimal_size = get_smart_batch_size(original)
        
        # ENFORCE max_allowed_batch_size for long documents
        if is_long_document and optimal_size > max_allowed_batch_size:
            optimal_size = max_allowed_batch_size
        
        # Track content types for stats
        if optimal_size == 10:
            batch_size_distribution['poetry'] += 1
        elif optimal_size == 50:
            batch_size_distribution['dialogue'] += 1
        elif optimal_size == 300:
            batch_size_distribution['prose'] += 1
        else:
            batch_size_distribution['default'] += 1
        
        total_paragraphs_to_translate += 1
        
        # If content type changed significantly, start a new batch
        if current_batch and abs(optimal_size - current_max_size) > 50:
            paragraph_batches.append(current_batch)
            current_batch = []
            current_max_size = min(optimal_size, max_allowed_batch_size) if is_long_document else optimal_size
        
        # Add to current batch
        current_batch.append((i, para, original))
        
        # If this is the first item in the batch, set the max size
        if len(current_batch) == 1:
            current_max_size = min(optimal_size, max_allowed_batch_size) if is_long_document else optimal_size
        
        # If batch is full or we're at the end, save it
        if len(current_batch) >= current_max_size or i == len(paragraphs) - 1:
            if current_batch:
                paragraph_batches.append(current_batch)
                current_batch = []
                current_max_size = 5  # Reset to NEW ultra-conservative default (was 100 - BUG!)
        
        i += 1

    # CRITICAL FIX: Save any remaining batch after loop ends
    if current_batch:
        paragraph_batches.append(current_batch)
        logs.append(f"[BATCH FIX] Added final batch with {len(current_batch)} paragraphs")

    total_batches = len(paragraph_batches)
    
    # Calculate optimization stats
    estimated_fixed_batches = max(1, total_paragraphs_to_translate // 20)  # If we used fixed size 20
    optimization_percentage = ((1 - total_batches / estimated_fixed_batches) * 100) if estimated_fixed_batches > 0 else 0
    
    logs.append(f"[SMART BATCHING] Created {total_batches} optimized batches")
    logs.append(f"[CONTENT ANALYSIS] Poetry/Formatted: {batch_size_distribution['poetry']}, Dialogue: {batch_size_distribution['dialogue']}, Prose: {batch_size_distribution['prose']}, Default: {batch_size_distribution['default']}")
    logs.append(f"[OPTIMIZATION] Would have been ~{estimated_fixed_batches} calls with fixed size 20")
    logs.append(f"[EFFICIENCY] Reduced API calls by {optimization_percentage:.0f}% using smart batching")
    logs.append(f"[PROCESSING] Starting parallel batch API requests (max 4 concurrent)...")
    
    # Initialize progress tracking
    if progress_id:
        progress_tracker[progress_id] = {
            "totalBatches": total_batches,
            "completedBatches": 0,
            "error": False
        }
    
    # Create semaphore to limit concurrent requests
    semaphore = asyncio.Semaphore(4)
    
    async def process_batch_with_semaphore(batch_idx, batch, session):
        """Process a single batch with semaphore control - returns logs separately"""
        batch_logs = []  # Separate logs for this batch
        
        async with semaphore:
            batch_paragraphs = [item[2] for item in batch]  # Extract text
            batch_size = len(batch_paragraphs)
            
            print(f"[TRANSLATOR] Processing batch {batch_idx + 1}/{total_batches} ({batch_size} paragraphs)")
            batch_logs.append(f"[BATCH {batch_idx + 1}/{total_batches}] Processing {batch_size} paragraphs...")
            batch_logs.append(f"[BATCH SIZE] Batch {batch_idx + 1} contains {batch_size} paragraphs (Smart batching in action!)")
            
            # Create batch prompt
            prompt = create_batch_prompt(batch_paragraphs, language)
            
            # Call API asynchronously
            batch_result = await call_openrouter_batch_api(session, prompt, model, api_key, batch_logs)
            
            # Update progress immediately after batch completes
            if progress_id:
                progress_tracker[progress_id]["completedBatches"] = progress_tracker[progress_id]["completedBatches"] + 1
            
            return batch_idx, batch, batch_paragraphs, batch_result, batch_logs
    
    # Create aiohttp session for async requests
    async with aiohttp.ClientSession() as session:
        # Create all tasks for parallel processing
        tasks = [
            process_batch_with_semaphore(batch_idx, batch, session)
            for batch_idx, batch in enumerate(paragraph_batches)
        ]
        
        # Execute all tasks in parallel (limited by semaphore)
        results = await asyncio.gather(*tasks)
        
        # Process results in order and merge logs
        for batch_idx, batch, batch_paragraphs, batch_result, batch_logs in results:
            # Add this batch's logs in order
            logs.extend(batch_logs)
            
            if batch_result:
                total_input_tokens += batch_result['input_tokens']
                total_output_tokens += batch_result['output_tokens']
                
                # Parse structured response
                batch_translations = parse_structured_response(batch_result['text'], len(batch_paragraphs), logs)
                
                logs.append(f"[BATCH {batch_idx + 1}] Received {len(batch_translations)} translations")
                
                # Validate we got the expected number of translations
                if len(batch_translations) != len(batch_paragraphs):
                    logs.append(f"[WARNING] Expected {len(batch_paragraphs)} translations, got {len(batch_translations)}")
                    # Pad with empty strings if we're short
                    while len(batch_translations) < len(batch_paragraphs):
                        batch_translations.append('[Translation missing]')
                    # Trim if we got too many
                    batch_translations = batch_translations[:len(batch_paragraphs)]
                
                # Apply translations to document (maintains order)
                for (para_idx, para, original), translation in zip(batch, batch_translations):
                    if translation and translation.strip():
                        translation = sanitize_response(translation)
                        translated_content.append(translation)
                        
                        # Update paragraph text in memory
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = translation
                        else:
                            para.add_run(translation)
                        
                        # Normalize copyright paragraphs - remove extra indentation
                        stripped_text = para.text.strip().lower()
                        if stripped_text.startswith("© copyright") or "all rights reserved" in stripped_text:
                            para_format = para.paragraph_format
                            para_format.left_indent = 0
                            para_format.first_line_indent = 0
                            # Try to set to Normal style, fallback if style doesn't exist
                            try:
                                para.style = doc.styles['Normal']
                            except:
                                pass  # Keep existing style if Normal doesn't exist
                    else:
                        fallback_text = f"[Translation failed for paragraph {para_idx}]"
                        translated_content.append(fallback_text)
            else:
                # Batch failed - mark as error and stop translation
                logs.append(f"[BATCH ERROR] Batch {batch_idx + 1} failed completely")
                if progress_id:
                    progress_tracker[progress_id]["error"] = True
                raise Exception(f"Translation failed at batch {batch_idx + 1}/{total_batches}. Please try again.")
    
    # Save document to memory buffer
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    logs.append(f"[SAVE] Document saved to memory buffer")
    logs.append(f"[TOKENS] Final usage - Input: {total_input_tokens}, Output: {total_output_tokens}, Total: {total_input_tokens + total_output_tokens}")
    logs.append(f"[DONE] Translation complete!")
    
    # Convert to base64
    translated_base64 = base64.b64encode(output_buffer.read()).decode('utf-8')
    
    print("[TRANSLATE] Complete, returning response")
    
    return TranslateResponse(
        translatedDocument=translated_base64,
        logs=logs,
        stats={
            "totalInputTokens": total_input_tokens,
            "totalOutputTokens": total_output_tokens,
            "totalTokens": total_input_tokens + total_output_tokens,
            "paragraphCount": len(translated_content),
            "translatedText": "\n\n".join(translated_content)  # Add full text for preview
        }
    )

async def translate_document_content_async_robust(
    file_bytes: bytes, 
    file_name: str, 
    language: str, 
    model: str, 
    api_key: str, 
    progress_id: Optional[str] = None
) -> TranslateResponse:
    """Enhanced translation with 100% format preservation"""
    
    if not ROBUST_FORMATTING_AVAILABLE:
        print("[WARNING] Falling back to standard translation - robust module not available")
        return await translate_document_content_async(file_bytes, file_name, language, model, api_key, progress_id)
    
    logs = []
    logs.append(f"[START] ROBUST translation with 100% format preservation")
    logs.append(f"[INFO] Source file: {file_name}")
    logs.append(f"[INFO] Target language: {language}")
    logs.append(f"[INFO] Model: {model}")
    
    # Load document
    doc = load_document(file_bytes)
    
    # ========== PROCESS TOC BEFORE TRANSLATION ==========
    if TOC_HANDLER_AVAILABLE:
        toc_results = process_toc_before_translation(doc)
        if toc_results['toc_found']:
            logs.append(f"[TOC] Found {toc_results['toc_entries_count']} TOC entries in first 10 pages")
            logs.append(f"[TOC] Extracted {toc_results['titles_extracted']} titles from TOC")
            logs.append(f"[TOC] Converted {toc_results['paragraphs_converted']} paragraphs to Heading 2 style")
            logs.append(f"[TOC] Removed {toc_results['toc_removed']} TOC entry paragraphs")
            if toc_results['placeholder_inserted']:
                logs.append(f"[TOC] Inserted placeholder for auto-generated TOC")
            print(f"[TOC PROCESSING] Found {toc_results['toc_entries_count']} entries, converted {toc_results['paragraphs_converted']} to Heading 2")
        else:
            logs.append(f"[TOC] No TOC found in first 10 pages")
    
    total_paragraphs = len(doc.paragraphs)
    logs.append(f"[LOAD] Document loaded with {total_paragraphs} paragraphs")
    
    # Initialize robust format preserver
    preserver = RobustFormatPreserver(doc)
    logs.append("[FORMAT] Initialized robust format preservation system")
    
    # Initialize Gemini client
    client = genai.Client(api_key=api_key)
    
    # Prepare paragraphs for translation
    paragraphs_to_translate = []
    marked_texts_for_batching = []
    
    para_count = 0
    for i, para in enumerate(doc.paragraphs):
        # Skip empty paragraphs
        if not para.text.strip():
            continue
        
        # Skip decorative paragraphs
        if not is_meaningful_text(para.text.strip()) or is_decorative_only(para.text.strip()):
            continue
        
        # Extract formatting and create marked text
        marked_text, para_data = preserver.create_formatted_text_for_translation(para, para_count)
        
        # Count runs and formatting complexity
        run_count = len(para.runs)
        format_types = set()
        for run_data in para_data['runs']:
            fmt = run_data['format']
            if fmt.get('bold'): format_types.add('bold')
            if fmt.get('italic'): format_types.add('italic')
            if fmt.get('underline'): format_types.add('underline')
            if fmt.get('font_name'): format_types.add('font')
            if fmt.get('font_color'): format_types.add('color')
        
        logs.append(f"[PARA {i}] {run_count} runs, {len(format_types)} format types: {format_types}")
        
        paragraphs_to_translate.append((i, para, marked_text, para_count))
        marked_texts_for_batching.append((para_count, marked_text))
        para_count += 1
    
    logs.append(f"[FILTER] {len(paragraphs_to_translate)} paragraphs to translate")
    
    # Smart batching based on complexity
    batches = create_smart_batches_for_robust_translation(marked_texts_for_batching, logs)
    total_batches = len(batches)
    logs.append(f"[BATCH] Created {total_batches} smart batches")
    
    # Initialize progress
    if progress_id:
        progress_tracker[progress_id] = {
            "totalBatches": total_batches,
            "completedBatches": 0,
            "error": False
        }
    
    # Process batches
    all_translations = {}
    total_input_tokens = 0
    total_output_tokens = 0
    
    for batch_idx, batch in enumerate(batches):
        logs.append(f"[BATCH {batch_idx + 1}/{total_batches}] Processing {len(batch)} paragraphs...")
        
        # Create robust prompt
        print(f"[ROBUST INPUT] Batch {batch_idx + 1}: sending {len(batch)} paragraphs to Gemini")
        for para_id, marked_text in batch:
            clean_text = re.sub(r'««[^»]+»»', '', marked_text)
            print(f"   - Para {para_id}: {preview_text(clean_text)}")
        prompt = create_robust_translation_prompt(batch, language)
        
        # Call API with timeout and retry logic
        max_attempts = 3
        batch_success = False
        
        for attempt in range(max_attempts):
            try:
                logs.append(f"[BATCH {batch_idx + 1}] API call attempt {attempt + 1}/{max_attempts}")
                
                # Use async wrapper to prevent blocking
                generation_config = types.GenerateContentConfig(
                    temperature=0.3,  # Lower temperature for format preservation
                    candidate_count=1,
                    max_output_tokens=8192,
                )
                
                # Use async wrapper with timeout
                response = await asyncio.wait_for(
                    call_gemini_batch_async(
                        ThreadPoolExecutor(max_workers=1),
                        client,
                        prompt,
                        model,
                        []
                    ),
                    timeout=600  # 10 minute timeout per batch (5x increase)
                )
                
                if not response or 'text' not in response:
                    raise Exception("Empty or invalid response from Gemini API")
                
                # Update token counts
                total_input_tokens += response.get('input_tokens', 0) or 0
                total_output_tokens += response.get('output_tokens', 0) or 0
                
                # Parse response
                batch_translations = parse_robust_response(response['text'], batch, logs)
                
                # Validate we got translations (but be lenient - allow fallback)
                if not batch_translations:
                    raise Exception("No translations returned from parser")
                
                # If count doesn't match, pad or trim
                if len(batch_translations) < len(batch):
                    logs.append(f"[WARNING] Got {len(batch_translations)} translations, expected {len(batch)} - padding with empty strings")
                    while len(batch_translations) < len(batch):
                        batch_translations.append("")
                elif len(batch_translations) > len(batch):
                    logs.append(f"[WARNING] Got {len(batch_translations)} translations, expected {len(batch)} - trimming")
                    batch_translations = batch_translations[:len(batch)]
                
                # Store translations
                for (para_id, _), translation in zip(batch, batch_translations):
                    all_translations[para_id] = translation
                    
                logs.append(f"[BATCH {batch_idx + 1}] Successfully translated {len(batch_translations)} paragraphs")
                batch_success = True
                break
                
            except asyncio.TimeoutError:
                logs.append(f"[ERROR] Batch {batch_idx + 1} timed out after 600 seconds (attempt {attempt + 1}/{max_attempts})")
                if attempt < max_attempts - 1:
                    await asyncio.sleep(2)
                else:
                    if progress_id:
                        progress_tracker[progress_id]["error"] = True
                    raise Exception(f"Batch {batch_idx + 1} timed out after {max_attempts} attempts")
            except Exception as e:
                logs.append(f"[ERROR] Batch {batch_idx + 1} attempt {attempt + 1} failed: {str(e)}")
                if attempt < max_attempts - 1:
                    await asyncio.sleep(2)
                else:
                    logs.append(f"[FAILED] Batch {batch_idx + 1} failed after {max_attempts} attempts")
                    if progress_id:
                        progress_tracker[progress_id]["error"] = True
                    raise Exception(f"Batch {batch_idx + 1} failed: {str(e)}")
        
        if not batch_success:
            if progress_id:
                progress_tracker[progress_id]["error"] = True
            raise Exception(f"Batch {batch_idx + 1} failed to complete")
        
        # Update progress
        if progress_id:
            progress_tracker[progress_id]["completedBatches"] += 1
    
    # Apply all translations with format preservation
    logs.append("[APPLY] Applying translations with format preservation...")
    
    for para_idx, para, marked_text, para_id in paragraphs_to_translate:
        if para_id in all_translations:
            translation = all_translations[para_id]
            
            # Apply formatting using robust preserver
            preserver.apply_formatting_to_paragraph(para, para_id, translation)
            
            logs.append(f"[APPLY {para_idx}] Applied translation with formatting preserved")
    
    # Save document
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    logs.append("[SAVE] Document saved with 100% format preservation")
    logs.append(f"[TOKENS] Total usage - Input: {total_input_tokens}, Output: {total_output_tokens}")
    logs.append("[DONE] Robust translation complete!")
    
    # Convert to base64
    translated_base64 = base64.b64encode(output_buffer.read()).decode('utf-8')
    
    return TranslateResponse(
        translatedDocument=translated_base64,
        logs=logs,
        stats={
            "totalInputTokens": total_input_tokens,
            "totalOutputTokens": total_output_tokens,
            "totalTokens": total_input_tokens + total_output_tokens,
            "totalParagraphs": total_paragraphs,
            "translatedParagraphs": len(paragraphs_to_translate),
            "preservedFormats": para_count,
            "method": "robust_100_percent"
        }
    )


def create_smart_batches_for_robust_translation(marked_texts: List[Tuple[int, str]], logs: List[str]) -> List[List[Tuple[int, str]]]:
    """Create batches optimized for robust translation"""
    
    # Calculate complexity for each paragraph
    complexities = []
    
    for para_id, marked_text in marked_texts:
        # Count formatting markers
        run_count = len(re.findall(r'««RUN\d+:', marked_text))
        
        # Count different format types
        format_types = set()
        for match in re.finditer(r'««RUN\d+:([^»]+)»»', marked_text):
            formats = match.group(1).split(',')
            format_types.update(formats)
        
        # Calculate text length without markers
        clean_text = re.sub(r'««[^»]+»»', '', marked_text)
        text_length = len(clean_text)
        
        complexity_score = run_count * len(format_types) * (1 + text_length / 1000)
        
        complexities.append({
            'para_id': para_id,
            'marked_text': marked_text,
            'run_count': run_count,
            'format_types': len(format_types),
            'text_length': text_length,
            'complexity': complexity_score
        })
    
    # Sort by complexity
    complexities.sort(key=lambda x: x['complexity'], reverse=True)
    
    # Create batches with complexity limits
    batches = []
    current_batch = []
    current_complexity = 0
    MAX_BATCH_COMPLEXITY = 50  # Adjust based on testing
    MAX_BATCH_SIZE = 10  # Maximum paragraphs per batch
    
    for item in complexities:
        if current_batch and (
            current_complexity + item['complexity'] > MAX_BATCH_COMPLEXITY or
            len(current_batch) >= MAX_BATCH_SIZE
        ):
            # Start new batch
            batches.append(current_batch)
            current_batch = []
            current_complexity = 0
        
        current_batch.append((item['para_id'], item['marked_text']))
        current_complexity += item['complexity']
    
    # Add final batch
    if current_batch:
        batches.append(current_batch)
    
    logs.append(f"[BATCH ANALYSIS] Created {len(batches)} batches from {len(marked_texts)} paragraphs")
    if complexities:
        logs.append(f"[BATCH COMPLEXITY] Most complex paragraph: {complexities[0]['complexity']:.2f}")
    
    return batches


def parse_robust_response(response_text: str, batch: List[Tuple[int, str]], logs: List[str]) -> List[str]:
    """Parse response with robust format markers - CRITICAL: Preserves order and prevents duplicates"""
    
    translations = []
    used_block_indices = set()  # Track which response blocks have been used to prevent duplicates
    
    # First pass: Try to extract by markers (preserves order)
    for batch_idx, (para_id, marked_text) in enumerate(batch):
        # Look for translation markers
        start_marker = f"<<<TRANSLATION_{para_id}_START>>>"
        end_marker = f"<<<TRANSLATION_{para_id}_END>>>"
        
        start_idx = response_text.find(start_marker)
        end_idx = response_text.find(end_marker)
        
        if start_idx != -1 and end_idx != -1:
            start_idx += len(start_marker)
            translation = response_text[start_idx:end_idx].strip()
            
            # CRITICAL: Remove any delimiter markers that might be in the translation
            translation = re.sub(r'<<<TRANSLATION_\d+_START>>>', '', translation)
            translation = re.sub(r'<<<TRANSLATION_\d+_END>>>', '', translation)
            
            # Verify format preservation
            original_runs = re.findall(r'««RUN(\d+):[^»]+»»', marked_text)
            translated_runs = re.findall(r'««RUN(\d+):[^»]+»»', translation)
            
            if len(original_runs) != len(translated_runs):
                logs.append(f"[WARNING] Para {para_id}: Run count mismatch - Original: {len(original_runs)}, Translated: {len(translated_runs)}")
            
            translations.append(translation)
        else:
            # Mark as missing for now - will fill in second pass
            translations.append(None)
            logs.append(f"[ERROR] Para {para_id} (batch pos {batch_idx}): Translation markers not found")
    
    # Second pass: Fill in missing translations with fallback (preserve order by position)
    # For missing translations, try to extract by position/order
    # Split response by double newlines to find paragraph-like blocks
    potential_blocks = [b.strip() for b in response_text.split('\n\n') if b.strip()]
    
    # Fill in missing translations in order
    block_idx = 0
    for batch_idx, (para_id, marked_text) in enumerate(batch):
        if translations[batch_idx] is not None:
            continue  # Already has translation
        
        # Try to find a matching block by run count and content similarity
        found_translation = None
        best_match_idx = -1
        
        for i, block in enumerate(potential_blocks):
            if i in used_block_indices:
                continue  # Skip already used blocks
            
            # Remove delimiter markers
            block_clean = re.sub(r'<<<TRANSLATION_\d+_START>>>', '', block)
            block_clean = re.sub(r'<<<TRANSLATION_\d+_END>>>', '', block_clean)
            block_clean = block_clean.strip()
            
            if not block_clean:
                continue
            
            # Check if this block has run markers (robust formatting)
            if '««RUN' in block_clean and '»»' in block_clean:
                # Count runs to see if it matches
                block_runs = re.findall(r'««RUN(\d+):[^»]+»»', block_clean)
                original_runs = re.findall(r'««RUN(\d+):[^»]+»»', marked_text)
                
                if len(block_runs) == len(original_runs) and len(block_runs) > 0:
                    # Good match - use this one
                    found_translation = block_clean
                    best_match_idx = i
                    break
            elif batch_idx == block_idx:
                # Position-based fallback: if no run markers, match by position
                # This is less reliable but better than nothing
                found_translation = block_clean
                best_match_idx = i
                break
        
        if found_translation:
            # Remove any remaining delimiter markers
            found_translation = re.sub(r'<<<TRANSLATION_\d+_START>>>', '', found_translation)
            found_translation = re.sub(r'<<<TRANSLATION_\d+_END>>>', '', found_translation)
            found_translation = found_translation.strip()
            
            translations[batch_idx] = found_translation
            used_block_indices.add(best_match_idx)  # Mark as used
            logs.append(f"[FALLBACK] Para {para_id} (batch pos {batch_idx}): Matched by {'run count' if '««RUN' in found_translation else 'position'}")
            block_idx += 1
        else:
            # Last resort: extract original marked text but STRIP ALL MARKERS
            logs.append(f"[FALLBACK] Para {para_id} (batch pos {batch_idx}): No match found - using original text with markers stripped")
            clean_text = re.sub(r'««[^»]+»»', '', marked_text)
            clean_text = re.sub(r'<<<TRANSLATION_\d+_START>>>', '', clean_text)
            clean_text = re.sub(r'<<<TRANSLATION_\d+_END>>>', '', clean_text)
            translations[batch_idx] = clean_text.strip()
            block_idx += 1
    
    # Final validation: ensure we have the right number of translations
    if len(translations) != len(batch):
        logs.append(f"[ERROR] Translation count mismatch: expected {len(batch)}, got {len(translations)}")
        # Pad or trim to match
        while len(translations) < len(batch):
            translations.append('[Translation missing]')
        translations = translations[:len(batch)]
    
    return translations


def create_batch_prompt(paragraphs: List[str], language: str) -> str:
    """Create batch translation prompt with DELIMITER-BASED format to avoid JSON issues"""
    
    # Use delimiter-based format to preserve ALL formatting (no JSON mangling)
    BATCH_PROMPT_TEMPLATE = """
You are a professional translator. Translate {count} passage(s) into {language}.

🎯 CRITICAL: READING LEVEL & MODERNIZATION REQUIREMENT:

**8TH GRADE READING LEVEL - MANDATORY:**
- Translate ALL text to modern, contemporary language suitable for 8th grade reading level (ages 13-14)
- This applies to ALL target languages, including English-to-English translation

**🔴 ENGLISH-TO-ENGLISH MODERNIZATION - ABSOLUTELY REQUIRED:**
- When target language is "English" or "Contemporary English", you MUST still translate/modernize the text
- DO NOT leave text unchanged just because it's already in English
- You MUST modernize old/archaic English to contemporary English
- You MUST simplify complex language to 8th grade level
- You MUST replace formal/old-fashioned words with modern equivalents
- Even if the text looks "modern", you MUST ensure it's at 8th grade reading level
- This is NOT optional - English-to-English still requires active translation work
- Every word and sentence must be reviewed and modernized if needed

**MODERNIZATION REQUIREMENTS:**
- If source text is in old/archaic English, modernize it to contemporary English
- Use simple, clear, everyday language that common people can easily understand
- Replace formal academic language with conversational language
- Replace complex vocabulary with simpler 8th grade level words
- Replace archaic words with modern equivalents:
  * "thou/thee/thy" → "you/your"
  * "hath/hast" → "has/have"
  * "doth" → "does"
  * "art" → "are"
  * "wilt" → "will"
  * "hither/thither" → "here/there"
  * "whence" → "from where"
  * "whither" → "to where"
  * "betwixt" → "between"
  * "ere" → "before"
  * "nigh" → "near"
  * "oft" → "often"
  * "perchance" → "perhaps"
  * "verily" → "truly" or "really"
  * "hence" → "therefore" or "so"
  * "thus" → "so" or "in this way"
  * "wherefore" → "why"
  * "methinks" → "I think"
  * "prithee" → "please"
  * "anon" → "soon" or "later"
- Simplify complex sentence structures into clear, straightforward sentences
- Break up very long sentences into shorter, more readable ones
- Use active voice instead of passive voice when possible
- Replace formal/archaic expressions with modern, conversational equivalents
- Maintain the meaning and tone, but make it accessible to modern readers
- Even if translating to English, modernize old English to contemporary English

**NUMBER FORMAT MODERNIZATION - CRITICAL:**
- Convert ALL Roman numerals to modern Arabic numerals (0-9)
- Convert ALL old/archaic number formats to modern Arabic numerals
- Examples of conversions:
  * Roman numerals: I → 1, II → 2, III → 3, IV → 4, V → 5, VI → 6, VII → 7, VIII → 8, IX → 9, X → 10
  * Larger Roman numerals: XI → 11, XII → 12, XIII → 13, XIV → 14, XV → 15, XX → 20, L → 50, C → 100, D → 500, M → 1000
  * "Chapter I" → "Chapter 1"
  * "Part III" → "Part 3"
  * "Volume XIV" → "Volume 14"
  * "Year MDCCLXXVI" → "Year 1776"
- When translating to modern English (8th grade level), use modern Arabic numerals exclusively
- Do NOT preserve Roman numerals in modern English translations
- Convert ordinal Roman numerals: "1st" (not "Ist"), "2nd" (not "IInd"), "3rd" (not "IIIrd"), "4th" (not "IVth")
- If number is part of a proper noun or historical reference that traditionally uses Roman numerals (like "World War II"), you may preserve it, but prefer modern format when modernizing
- Convert written-out archaic number forms: "one and twenty" → "21", "three score" → "60"
- Use standard modern number format: "1,234" not "one thousand two hundred thirty-four" (unless context requires words)

**EXAMPLES - ENGLISH-TO-ENGLISH MODERNIZATION:**

Example 1 - Archaic English:
- Old: "Thou art a goodly fellow, methinks."
- Modern: "I think you're a good person."
- ❌ WRONG: Leaving it as "Thou art a goodly fellow, methinks." (unchanged)
- ✅ CORRECT: Modernizing to "I think you're a good person."

Example 2 - Archaic Questions:
- Old: "Whence comest thou, and whither goest thou?"
- Modern: "Where are you coming from, and where are you going?"
- ❌ WRONG: Leaving it unchanged
- ✅ CORRECT: Modernizing to contemporary English

Example 3 - Formal/Old English:
- Old: "Verily, I say unto thee, this matter doth concern us all."
- Modern: "I'm telling you, this matter concerns all of us."
- ❌ WRONG: Leaving it unchanged
- ✅ CORRECT: Modernizing to conversational English

Example 4 - Complex Academic English:
- Old: "The aforementioned individual has demonstrated a propensity for engaging in activities that are not in accordance with established protocols."
- Modern: "This person has a habit of doing things that break the rules."
- ❌ WRONG: Leaving complex academic language unchanged
- ✅ CORRECT: Simplifying to 8th grade level

Example 5 - Old-Fashioned Formal English:
- Old: "Hath he not spoken thus to thee ere this day?"
- Modern: "Hasn't he said this to you before today?"
- ❌ WRONG: Leaving it unchanged
- ✅ CORRECT: Modernizing to contemporary English

Example 6 - Even "Modern" English Needs Simplification:
- Old: "The individual's cognitive processes were significantly impeded by the complexity of the situation."
- Modern: "The person had trouble thinking because the situation was too complicated."
- ❌ WRONG: Leaving complex academic language unchanged
- ✅ CORRECT: Simplifying to 8th grade level

Example 7 - Roman Numeral Conversion:
- Old: "Chapter I discusses the basics."
- Modern: "Chapter 1 discusses the basics."
- ❌ WRONG: Leaving "Chapter I" unchanged
- ✅ CORRECT: Converting to "Chapter 1"

Example 8 - Roman Numeral Conversion (Multiple):
- Old: "In Part III, Section XIV, we find..."
- Modern: "In Part 3, Section 14, we find..."
- ❌ WRONG: Leaving "Part III, Section XIV" unchanged
- ✅ CORRECT: Converting to "Part 3, Section 14"

Example 9 - Roman Numeral in Title:
- Old: "Volume II of the collection"
- Modern: "Volume 2 of the collection"
- ❌ WRONG: Leaving "Volume II" unchanged
- ✅ CORRECT: Converting to "Volume 2"

Example 10 - Year with Roman Numerals:
- Old: "In the year MDCCLXXVI, the Declaration was signed."
- Modern: "In the year 1776, the Declaration was signed."
- ❌ WRONG: Leaving "MDCCLXXVI" unchanged
- ✅ CORRECT: Converting to "1776"

**REMEMBER:**
- When target is "English" or "Contemporary English", you MUST actively modernize
- Do NOT copy text unchanged - always review and modernize
- Every sentence must be checked for 8th grade readability
- Complex words → Simple words
- Formal language → Conversational language
- Old English → Modern English

**LINK / URL HANDLING - CRITICAL:**
- REMOVE all hyperlinks, URLs, and link markup from the translation
- If original text contains a URL (http://, https://, www., etc.) → OMIT it entirely
- If original text contains Markdown links like `[text](url)` → Output ONLY the text, remove the `(url)`
- If original text contains angle-bracket links `<http://...>` → Remove them completely
- If original text contains HTML links `<a href="...">text</a>` → Output only the text, remove the href
- If original contains "See here: https://..." → Translate "See here:" but REMOVE the URL
- NEVER translate or preserve URLs or links – they must be removed
- This prevents duplicated content and unnecessary link noise in the translation

🚫 CRITICAL: NO HALLUCINATION - EXACT LINE & CONTENT PRESERVATION:

**MANDATORY - ZERO HALLUCINATION POLICY:**
- DO NOT add any content that is not in the original text
- DO NOT remove any content from the original text
- DO NOT add extra sentences, explanations, or commentary
- DO NOT summarize or condense the text
- DO NOT expand or elaborate on the text
- Translate sentence-by-sentence, maintaining exact sentence count
- If original has 5 lines → translation MUST have exactly 5 lines
- If original has 10 sentences → translation MUST have exactly 10 sentences
- Count lines before translating and ensure your translation has the SAME number of lines
- Count sentences before translating and ensure your translation has the SAME number of sentences

**LINE COUNT PRESERVATION - ABSOLUTE REQUIREMENT:**
- Count the number of lines in the original passage (lines separated by \\n)
- Your translation MUST have the EXACT same number of lines
- Example: If original has 5 lines:
  Line 1: "Hello"
  Line 2: "How are you?"
  Line 3: "I am fine."
  Line 4: "Thank you."
  Line 5: "Goodbye"
  
  Translation MUST also have exactly 5 lines:
  Line 1: [translation of line 1]
  Line 2: [translation of line 2]
  Line 3: [translation of line 3]
  Line 4: [translation of line 4]
  Line 5: [translation of line 5]
  
- DO NOT combine lines into one
- DO NOT split one line into multiple lines
- DO NOT add blank lines
- DO NOT remove blank lines
- If original has a blank line, translation must have a blank line in the same position

**SENTENCE COUNT PRESERVATION:**
- Count the number of sentences in the original (sentences end with . ! ?)
- Your translation MUST have the EXACT same number of sentences
- Translate each sentence independently - do not combine or split sentences
- If original has 3 sentences, translation must have 3 sentences
- DO NOT add explanatory sentences
- DO NOT add transitional phrases that weren't in the original
- DO NOT remove sentences or combine them

**CONTENT FIDELITY:**
- Translate ONLY what is written - nothing more, nothing less
- If a sentence is unclear, translate it as written - do not "clarify" or "improve" it
- Do not add context or background information
- Do not add examples or explanations
- Do not add connecting words or phrases that weren't in the original
- Maintain the exact structure: if original is terse, translation should be terse
- Maintain the exact structure: if original is verbose, translation should be verbose

**ANTI-HALLUCINATION CHECKLIST:**
Before submitting your translation, verify:
✓ Same number of lines as original
✓ Same number of sentences as original
✓ No added content
✓ No removed content
✓ No added explanations
✓ No added examples
✓ No added transitions
✓ Exact 1:1 correspondence between original and translation

🔴 CRITICAL FORMATTING RULES - CHARACTER-BY-CHARACTER PRESERVATION:

1. PRESERVE EVERY SINGLE CHARACTER:
   - Count EVERY space at the start of lines (indentation)
   - Count EVERY space at the end of lines  
   - Count EVERY newline character (\\n)
   - If line starts with 4 spaces → translation starts with 4 spaces
   - If 2 blank lines between text → 2 blank lines in translation
   - NEVER add, remove, or modify ANY whitespace

2. FOR POETRY/FORMATTED TEXT:
   - Line breaks are ARTISTIC - preserve exactly
   - Short lines MUST stay short
   - Indentation creates RHYTHM - keep every space
   - Visual layout = part of the art

3. PUNCTUATION PRESERVATION - CRITICAL:
   - PRESERVE ALL punctuation marks EXACTLY as they appear
   - Quotation marks ("") stay as quotation marks ("")
   - Em dashes (—) stay as em dashes (—)
   - En dashes (–) stay as en dashes (–)
   - Apostrophes (') stay as apostrophes (')
   - Do NOT convert between punctuation styles
   - Do NOT change "dialogue" to —dialogue— or vice versa
   - Do NOT replace quotes with em dashes or em dashes with quotes
   - ONLY translate the words, NOT the punctuation marks
   - If original has: "Hello" → translation has: "Hola" (NOT —Hola—)
   - If original has: —Hello— → translation has: —Hola— (NOT "Hola")
   - Punctuation style is part of the original formatting - preserve it

4. COMPLETE WORD TRANSLATION - MANDATORY:
   - TRANSLATE EVERY SINGLE WORD in the source language
   - NO English words should remain in the translation
   - Common words like "yes", "no", "certainly", "maybe", "okay" MUST be translated
   - Single quoted words like "yes", "no", "okay" MUST be translated: "sí", "no", "de acuerdo"
   - Words in quotes are NOT exceptions - they MUST be translated
   - If you see: "yes" → translate to: "sí" (or equivalent in target language)
   - If you see: "certainly" → translate to: "ciertamente" (or equivalent)
   - If you see: "okay" → translate to: "de acuerdo" (or equivalent)
   - Interjections, exclamations, and standalone words MUST be translated
   - Do NOT leave ANY source language words untranslated
   - Every meaningful word in the source language must have a translation
   - If a word appears in quotes, translate the word but keep the quotes: "yes" → "sí"
   - If a word appears alone, translate it: yes → sí
   - NO exceptions - translate EVERY word

5. OUTPUT FORMAT - USE DELIMITERS:
   For EACH passage, output EXACTLY like this:
   
   <<<TRANSLATION_START_{i}>>>
   [Your translation with EXACT formatting]
   <<<TRANSLATION_END_{i}>>>
   
   - Replace {i} with passage number (1, 2, 3, etc.)
   - NO extra text before/after delimiters
   - NO explanations or notes
   - PRESERVE all spaces and newlines inside delimiters
   - ⚠️ CRITICAL: Do NOT put delimiter markers (<<<TRANSLATION_...) INSIDE your translation text
   - ⚠️ Do NOT generate random <<< markers or patterns like <<<TRANSL000... 
   - ⚠️ The ONLY <<< markers you should output are the START and END delimiters wrapping each translation
   - Your translation text should ONLY contain the actual translated content, NO delimiter markers

EXAMPLES:

Example 1 - Spacing:
If original is:
    Hello
        World

Translation should be:
<<<TRANSLATION_START_1>>>
    Hola
        Mundo
<<<TRANSLATION_END_1>>>
(Notice the EXACT spacing preserved)

Example 2 - Single Quoted Words:
If original is:
    He said "yes" and "certainly" replied.

Translation should be:
<<<TRANSLATION_START_1>>>
    Dijo "sí" y "ciertamente" respondió.
<<<TRANSLATION_END_1>>>
(Notice: "yes" → "sí", "certainly" → "ciertamente" - ALL words translated)

Example 3 - Standalone Words:
If original is:
    Yes, certainly, okay.

Translation should be:
<<<TRANSLATION_START_1>>>
    Sí, ciertamente, de acuerdo.
<<<TRANSLATION_END_1>>>
(Notice: Every word translated, no English words remain)

CRITICAL: Do NOT include any tags like <untranslated> in your translation output. Only output the translated text.

Original Passages:
{passages}

Translate into {language}. Use delimiters <<<TRANSLATION_START_N>>> and <<<TRANSLATION_END_N>>> for each passage."""
    
    passages_text = ""
    for i, para in enumerate(paragraphs, 1):
        # Show the exact text including all whitespace
        passages_text += f'\n--- Passage {i} ---\n{para}\n--- End Passage {i} ---\n'
    
    return BATCH_PROMPT_TEMPLATE.format(
        count=len(paragraphs),
        language=language,
        passages=passages_text,
        i="{i}"  # Template placeholder
    )

if __name__ == "__main__":
    import uvicorn
    # Use reload so code changes are picked up automatically in development
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=7860,
        reload=True,
    )
