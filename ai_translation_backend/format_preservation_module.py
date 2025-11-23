"""
Format Preservation Module for Document Translation
Implements intelligent formatting preservation strategies
"""

import re
from typing import List, Dict, Tuple, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run


class DocumentComplexityAnalyzer:
    """Analyzes document complexity to determine optimal translation strategy"""
    
    def __init__(self, doc: Document):
        self.doc = doc
        self.analysis = self._analyze()
    
    def _analyze(self) -> Dict:
        """Perform comprehensive document analysis"""
        stats = {
            'total_paragraphs': 0,
            'formatted_paragraphs': 0,
            'total_runs': 0,
            'inline_formatting_count': 0,
            'font_variations': set(),
            'has_tables': len(self.doc.tables) > 0,
            'has_images': False,  # Would need additional logic
            'avg_runs_per_para': 0,
            'formatting_density': 0,
            'complexity_tier': None
        }
        
        meaningful_paragraphs = 0
        
        for para in self.doc.paragraphs:
            if not para.text.strip():
                continue
                
            stats['total_paragraphs'] += 1
            meaningful_paragraphs += 1
            runs = list(para.runs)
            stats['total_runs'] += len(runs)
            
            # Check for inline formatting
            if len(runs) > 1:
                stats['formatted_paragraphs'] += 1
                
            for run in runs:
                # Count formatting types
                if run.bold:
                    stats['inline_formatting_count'] += 1
                if run.italic:
                    stats['inline_formatting_count'] += 1
                if run.underline:
                    stats['inline_formatting_count'] += 1
                if run.font.name:
                    stats['font_variations'].add(run.font.name)
        
        # Calculate metrics
        if meaningful_paragraphs > 0:
            stats['avg_runs_per_para'] = stats['total_runs'] / meaningful_paragraphs
        if stats['total_runs'] > 0:
            stats['formatting_density'] = stats['inline_formatting_count'] / stats['total_runs']
        
        # Determine complexity tier
        stats['complexity_tier'] = self._determine_tier(stats)
        
        return stats
    
    def _determine_tier(self, stats: Dict) -> str:
        """Determine document complexity tier"""
        avg_runs = stats['avg_runs_per_para']
        fmt_density = stats['formatting_density']
        font_diversity = len(stats['font_variations'])
        
        # Tier 3: Complex documents
        if avg_runs > 5 or fmt_density > 0.3 or font_diversity > 3:
            return "TIER_3_COMPLEX"
        # Tier 2: Moderate complexity
        elif avg_runs > 1.5 or fmt_density > 0.1 or font_diversity > 1:
            return "TIER_2_MODERATE"
        # Tier 1: Simple documents
        else:
            return "TIER_1_SIMPLE"
    
    def get_tier(self) -> str:
        """Get the complexity tier"""
        return self.analysis['complexity_tier']
    
    def get_analysis(self) -> Dict:
        """Get full analysis results"""
        return self.analysis


class FormatPreservationMap:
    """Maps and preserves formatting information from original to translated text"""
    
    def __init__(self):
        self.format_maps = []
    
    def extract_paragraph_formatting(self, paragraph: Paragraph) -> Dict:
        """Extract all formatting information from a paragraph"""
        format_info = {
            'runs': [],
            'full_text': paragraph.text,
            'style': paragraph.style.name if paragraph.style else None,
            'alignment': paragraph.alignment,
            'left_indent': paragraph.paragraph_format.left_indent,
            'right_indent': paragraph.paragraph_format.right_indent,
            'first_line_indent': paragraph.paragraph_format.first_line_indent,
            'space_before': paragraph.paragraph_format.space_before,
            'space_after': paragraph.paragraph_format.space_after,
            'line_spacing': paragraph.paragraph_format.line_spacing
        }
        
        char_position = 0
        for i, run in enumerate(paragraph.runs):
            run_info = {
                'index': i,
                'start': char_position,
                'end': char_position + len(run.text),
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'strike': run.font.strike,
                'subscript': run.font.subscript,
                'superscript': run.font.superscript,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color else None,
                'highlight': run.font.highlight_color
            }
            format_info['runs'].append(run_info)
            char_position += len(run.text)
        
        return format_info
    
    def mark_formatting_in_text(self, paragraph: Paragraph, tier: str) -> str:
        """Add formatting markers to text based on tier"""
        if tier == "TIER_1_SIMPLE":
            # No markers needed for simple tier
            return paragraph.text
        
        marked_text = ""
        
        for i, run in enumerate(paragraph.runs):
            text = run.text
            
            if tier == "TIER_2_MODERATE":
                # Use lightweight markers
                if run.bold:
                    text = f"««B»»{text}««/B»»"
                if run.italic:
                    text = f"««I»»{text}««/I»»"
                if run.underline:
                    text = f"««U»»{text}««/U»»"
            
            elif tier == "TIER_3_COMPLEX":
                # Use detailed markers with run index
                markers = []
                if run.bold: markers.append("B")
                if run.italic: markers.append("I")
                if run.underline: markers.append("U")
                if run.font.strike: markers.append("S")
                if run.font.subscript: markers.append("SUB")
                if run.font.superscript: markers.append("SUP")
                
                if markers:
                    marker_str = ",".join(markers)
                    text = f"««R{i}:{marker_str}»»{text}««/R{i}»»"
                else:
                    text = f"««R{i}»»{text}««/R{i}»»"
            
            marked_text += text
        
        return marked_text
    
    def parse_marked_translation(self, translated_text: str, tier: str) -> Tuple[str, List[Dict]]:
        """Parse formatting markers from translated text"""
        if tier == "TIER_1_SIMPLE":
            return translated_text, []
        
        format_segments = []
        
        if tier == "TIER_2_MODERATE":
            # Parse simple markers
            patterns = {
                'bold': r'««B»»(.*?)««/B»»',
                'italic': r'««I»»(.*?)««/I»»',
                'underline': r'««U»»(.*?)««/U»»'
            }
            
            for fmt_type, pattern in patterns.items():
                for match in re.finditer(pattern, translated_text):
                    format_segments.append({
                        'start': match.start(),
                        'end': match.end(),
                        'text': match.group(1),
                        'type': fmt_type,
                        'marker_text': match.group(0)
                    })
            
            # Remove all markers
            clean_text = re.sub(r'««[^»]+»»', '', translated_text)
        
        else:  # TIER_3_COMPLEX
            # Parse complex run markers
            pattern = r'««R(\d+)(?::([^»]+))?»»(.*?)««/R\1»»'
            
            for match in re.finditer(pattern, translated_text):
                run_index = int(match.group(1))
                formats = match.group(2).split(',') if match.group(2) else []
                text = match.group(3)
                
                format_segments.append({
                    'run_index': run_index,
                    'formats': formats,
                    'text': text,
                    'full_match': match.group(0),
                    'start': match.start(),
                    'end': match.end()
                })
            
            # Remove all markers
            clean_text = re.sub(r'««R\d+[^»]*»»|««/R\d+»»', '', translated_text)
        
        return clean_text, format_segments


class SmartBatchManager:
    """Manages intelligent batching based on document tier and content type"""
    
    # Batch size configuration by tier and content type
    BATCH_SIZES = {
        "TIER_1_SIMPLE": {
            "poetry": 50,
            "dialogue": 100,
            "prose": 300,
            "mixed": 200,
            "default": 150
        },
        "TIER_2_MODERATE": {
            "poetry": 10,
            "dialogue": 30,
            "prose": 50,
            "mixed": 40,
            "default": 30
        },
        "TIER_3_COMPLEX": {
            "poetry": 3,
            "dialogue": 5,
            "prose": 10,
            "mixed": 7,
            "default": 5
        }
    }
    
    @staticmethod
    def detect_content_type(text: str) -> str:
        """Detect the type of content for optimal batching"""
        # Check for poetry indicators
        if '\\' in text or text.count('\n') > 3:
            return "poetry"
        
        # Check for dialogue
        dialogue_indicators = ['"', '"', '"', '«', '»', '–', '—']
        dialogue_count = sum(text.count(indicator) for indicator in dialogue_indicators)
        if dialogue_count > 4:
            return "dialogue"
        
        # Check for prose
        if len(text) > 500 and '.' in text:
            return "prose"
        
        # Mixed or uncertain
        return "mixed"
    
    @classmethod
    def get_optimal_batch_size(cls, tier: str, content_type: str, 
                              is_long_document: bool = False) -> int:
        """Get optimal batch size based on tier and content"""
        base_size = cls.BATCH_SIZES[tier].get(content_type, cls.BATCH_SIZES[tier]["default"])
        
        # Further reduce for very long documents
        if is_long_document and tier != "TIER_1_SIMPLE":
            return max(1, base_size // 2)
        
        return base_size
    
    @classmethod
    def create_smart_batches(cls, paragraphs: List[Tuple[int, Paragraph, str]], 
                           tier: str, is_long_doc: bool) -> List[List[Tuple[int, Paragraph, str]]]:
        """Create optimized batches based on content analysis"""
        batches = []
        current_batch = []
        current_content_type = None
        current_max_size = None
        
        for item in paragraphs:
            idx, para, text = item
            
            # Detect content type for this paragraph
            content_type = cls.detect_content_type(text)
            
            # Get optimal size for this content
            optimal_size = cls.get_optimal_batch_size(tier, content_type, is_long_doc)
            
            # Check if we need to start a new batch
            should_new_batch = False
            
            if not current_batch:
                # First paragraph
                current_content_type = content_type
                current_max_size = optimal_size
            elif current_content_type != content_type and abs(optimal_size - current_max_size) > 20:
                # Content type changed significantly
                should_new_batch = True
            elif len(current_batch) >= current_max_size:
                # Current batch is full
                should_new_batch = True
            
            # Start new batch if needed
            if should_new_batch and current_batch:
                batches.append(current_batch)
                current_batch = []
                current_content_type = content_type
                current_max_size = optimal_size
            
            # Add to current batch
            current_batch.append(item)
            
            # Update max size to be most restrictive
            if current_max_size:
                current_max_size = min(current_max_size, optimal_size)
        
        # Add final batch
        if current_batch:
            batches.append(current_batch)
        
        return batches


class FormattingReconstructor:
    """Reconstructs document with preserved formatting after translation"""
    
    @staticmethod
    def apply_simple_translation(paragraph: Paragraph, translation: str):
        """Apply translation for simple tier (paragraph-level formatting only)"""
        # Clear existing runs
        for run in paragraph.runs:
            run.text = ""
        
        # Set translation in first run
        if paragraph.runs:
            paragraph.runs[0].text = translation
        else:
            paragraph.add_run(translation)
    
    @staticmethod
    def apply_moderate_translation(paragraph: Paragraph, translation: str, 
                                 format_info: Dict, format_segments: List[Dict]):
        """Apply translation for moderate tier with format preservation"""
        # Clear existing runs
        for run in paragraph.runs:
            run.text = ""
        
        if not format_segments:
            # No formatting found, apply as simple
            FormattingReconstructor.apply_simple_translation(paragraph, translation)
            return
        
        # Sort segments by position
        format_segments.sort(key=lambda x: x['start'])
        
        # Reconstruct with formatting
        current_pos = 0
        run_index = 0
        
        for segment in format_segments:
            # Add any text before this segment
            if segment['start'] > current_pos:
                plain_text = translation[current_pos:segment['start']]
                if plain_text:
                    if run_index < len(paragraph.runs):
                        paragraph.runs[run_index].text = plain_text
                    else:
                        paragraph.add_run(plain_text)
                    run_index += 1
            
            # Add formatted segment
            if run_index < len(paragraph.runs):
                run = paragraph.runs[run_index]
            else:
                run = paragraph.add_run()
            
            run.text = segment['text']
            
            # Apply formatting
            if segment['type'] == 'bold':
                run.bold = True
            elif segment['type'] == 'italic':
                run.italic = True
            elif segment['type'] == 'underline':
                run.underline = True
            
            current_pos = segment['end']
            run_index += 1
        
        # Add any remaining text
        if current_pos < len(translation):
            remaining_text = translation[current_pos:]
            if remaining_text:
                if run_index < len(paragraph.runs):
                    paragraph.runs[run_index].text = remaining_text
                else:
                    paragraph.add_run(remaining_text)
    
    @staticmethod
    def apply_complex_translation(paragraph: Paragraph, translation: str,
                                format_info: Dict, format_segments: List[Dict]):
        """Apply translation for complex tier with full format preservation"""
        # This would implement full run-by-run reconstruction
        # For now, fallback to moderate approach
        FormattingReconstructor.apply_moderate_translation(
            paragraph, translation, format_info, format_segments
        )


# Export main classes
__all__ = [
    'DocumentComplexityAnalyzer',
    'FormatPreservationMap',
    'SmartBatchManager',
    'FormattingReconstructor'
]
