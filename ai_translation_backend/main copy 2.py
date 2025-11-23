from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, RedirectResponse
from pydantic import BaseModel
import os
import io
import base64
import asyncio
import time
from typing import List, Optional
from dotenv import load_dotenv
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt
import re
from concurrent.futures import ThreadPoolExecutor
import aiohttp
load_dotenv()

# Allow OAuth scope changes (Google may return broader permissions than requested)
os.environ['OAUTHLIB_RELAX_TOKEN_SCOPE'] = '1'

app = FastAPI(title="Drive Document Translator API")

# CORS middleware - allow origins from environment or default to all
allowed_origins = os.getenv("ALLOWED_ORIGINS", "*").split(",")
if allowed_origins == ["*"]:
    # Allow all origins
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=False,  # Must be False when allow_origins is ["*"]
        allow_methods=["*"],
        allow_headers=["*"],
    )
else:
    # Specific origins with credentials support
    app.add_middleware(
        CORSMiddleware,
        allow_origins=[origin.strip() for origin in allowed_origins],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

# Configuration
BATCH_SIZE = 10000
MAX_RETRIES = 3
RETRY_DELAY = 2

# OAuth configuration
SCOPES = [
    'https://www.googleapis.com/auth/drive.file',  # Create and modify files that the app creates
    'https://www.googleapis.com/auth/drive.readonly',  # Read files
    'https://www.googleapis.com/auth/drive.metadata.readonly'  # Read metadata
]

# In-memory session storage (in production, use Redis or database)
sessions = {}

# Request/Response Models
class TranslateRequest(BaseModel):
    fileData: str
    fileName: str
    language: str
    model: str
    apiKey: str

class TranslateResponse(BaseModel):
    translatedDocument: str
    logs: List[str]
    stats: dict

class CreateFolderRequest(BaseModel):
    state: str
    folderName: str
    parentFolderId: Optional[str] = None

class UploadFileRequest(BaseModel):
    state: str
    folderId: str
    fileName: str
    fileData: str
    mimeType: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

@app.get("/")
def read_root():
    return {"status": "Drive Document Translator API is running"}

@app.get("/api/drive/auth")
async def drive_auth():
    """Initiate OAuth flow"""
    try:
        flow = Flow.from_client_config(
            {
                "web": {
                    "client_id": os.getenv("GOOGLE_CLIENT_ID"),
                    "client_secret": os.getenv("GOOGLE_CLIENT_SECRET"),
                    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                    "token_uri": "https://oauth2.googleapis.com/token",
                    "redirect_uris": [f"{os.getenv('BACKEND_URL', 'http://localhost:8000')}/api/drive/callback"]
                }
            },
            scopes=SCOPES,
            redirect_uri=f"{os.getenv('GOOGLE_REDIRECT_URI', 'http://localhost:8000')}/api/drive/callback"
        )
        
        authorization_url, state = flow.authorization_url(
            access_type='offline',
            include_granted_scopes='true',
            prompt='consent'
        )
        
        # Store state for verification
        sessions[state] = {"flow": flow}
        
        return {"authUrl": authorization_url, "state": state}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Auth failed: {str(e)}")

@app.get("/api/drive/callback")
async def drive_callback(code: str, state: str):
    """Handle OAuth callback"""
    try:
        print(f"[CALLBACK] Received callback with state: {state}")
        print(f"[CALLBACK] Available sessions: {list(sessions.keys())}")
        
        if state not in sessions:
            print(f"[CALLBACK ERROR] State not found in sessions")
            raise HTTPException(status_code=400, detail="Invalid state")
        
        print(f"[CALLBACK] Fetching token with code...")
        flow = sessions[state]["flow"]
        flow.fetch_token(code=code)
        
        credentials = flow.credentials
        print(f"[CALLBACK] Credentials obtained successfully")
        
        sessions[state]["credentials"] = {
            "token": credentials.token,
            "refresh_token": credentials.refresh_token,
            "token_uri": credentials.token_uri,
            "client_id": credentials.client_id,
            "client_secret": credentials.client_secret,
            "scopes": credentials.scopes
        }
        
        print(f"[CALLBACK] Redirecting to frontend with state: {state}")
        # Redirect to frontend with state
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        return RedirectResponse(url=f"{frontend_url}/?authenticated=true&state={state}")
    except Exception as e:
        print(f"[CALLBACK ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        return RedirectResponse(url=f"{frontend_url}/?error=auth_failed&message={str(e)}")

@app.get("/api/drive/folders")
async def get_folders(state: str):
    """List folders from Google Drive"""
    try:
        if state not in sessions or "credentials" not in sessions[state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[state]["credentials"]
        credentials = Credentials(**creds_dict)
        
        service = build('drive', 'v3', credentials=credentials)
        
        results = service.files().list(
            q="mimeType='application/vnd.google-apps.folder' and trashed = false",
            pageSize=100,
            fields="files(id, name)"
        ).execute()
        
        return {"folders": results.get('files', [])}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch folders: {str(e)}")

@app.get("/api/drive/files")
async def get_files(state: str, folderId: Optional[str] = None, driveLink: Optional[str] = None):
    """List files from a folder or Drive link"""
    try:
        if state not in sessions or "credentials" not in sessions[state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[state]["credentials"]
        credentials = Credentials(**creds_dict)
        service = build('drive', 'v3', credentials=credentials)
        
        # Extract ID from link if provided
        if driveLink:
            folderId = extract_drive_id(driveLink)
        
        if not folderId:
            raise HTTPException(status_code=400, detail="Folder ID or Drive link required")
        
        # Get metadata to check if folder or file
        metadata = service.files().get(
            fileId=folderId,
            fields="id, name, mimeType"
        ).execute()
        
        is_folder = metadata['mimeType'] == 'application/vnd.google-apps.folder'
        
        if not is_folder:
            return {
                "type": "file",
                "file": metadata
            }
        
        # List files in folder
        supported_types = [
            'application/vnd.google-apps.document',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/pdf'
        ]
        
        mime_conditions = ' or '.join([f"mimeType='{t}'" for t in supported_types])
        query = f"'{folderId}' in parents and trashed = false and ({mime_conditions})"
        
        results = service.files().list(
            q=query,
            pageSize=100,
            fields="files(id, name, mimeType, size, modifiedTime)"
        ).execute()
        
        return {
            "type": "folder",
            "folderName": metadata['name'],
            "files": results.get('files', [])
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch files: {str(e)}")

@app.get("/api/drive/download")
async def download_file(state: str, fileId: str, mimeType: str):
    """Download a file from Google Drive"""
    try:
        if state not in sessions or "credentials" not in sessions[state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[state]["credentials"]
        credentials = Credentials(**creds_dict)
        service = build('drive', 'v3', credentials=credentials)
        
        # Google Workspace types that need export
        workspace_types = {
            'application/vnd.google-apps.document': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        }
        
        if mimeType in workspace_types:
            request = service.files().export_media(
                fileId=fileId,
                mimeType=workspace_types[mimeType]
            )
        else:
            request = service.files().get_media(fileId=fileId)
        
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        
        while not done:
            status, done = downloader.next_chunk()
        
        fh.seek(0)
        file_data = fh.read()
        
        # Convert to base64
        base64_data = base64.b64encode(file_data).decode('utf-8')
        
        return {
            "data": base64_data,
            "size": len(file_data)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to download file: {str(e)}")

@app.post("/api/translate", response_model=TranslateResponse)
async def translate_document(request: TranslateRequest):
    """Translate a document with async batch processing using Gemini"""
    try:
        print(f"[TRANSLATE] Starting async translation for {request.fileName}")
        
        # Decode base64 file data
        file_bytes = base64.b64decode(request.fileData)
        print(f"[TRANSLATE] File decoded, size: {len(file_bytes)} bytes")
        
        # Translate document using async function
        result = await translate_document_content_async(
            file_bytes,
            request.fileName,
            request.language,
            request.model,
            request.apiKey
        )
        
        print(f"[TRANSLATE] Translation complete")
        return result
    except Exception as e:
        print(f"[TRANSLATE] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

@app.post("/api/translate/openrouter", response_model=TranslateResponse)
async def translate_document_openrouter(request: TranslateRequest):
    """Translate a document with async batch processing using OpenRouter"""
    try:
        print(f"[TRANSLATE OPENROUTER] Starting async translation for {request.fileName}")
        
        # Decode base64 file data
        file_bytes = base64.b64decode(request.fileData)
        print(f"[TRANSLATE OPENROUTER] File decoded, size: {len(file_bytes)} bytes")
        
        # Translate document using OpenRouter async function
        result = await translate_document_content_async_openrouter(
            file_bytes,
            request.fileName,
            request.language,
            request.model,
            request.apiKey
        )
        
        print(f"[TRANSLATE OPENROUTER] Translation complete")
        return result
    except Exception as e:
        print(f"[TRANSLATE OPENROUTER] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

@app.post("/api/drive/logout")
async def logout(state: str):
    """Clear session"""
    if state in sessions:
        del sessions[state]
    return {"success": True}

@app.post("/api/drive/create-folder")
async def create_folder(request: CreateFolderRequest):
    """Create a new folder in Google Drive"""
    try:
        if request.state not in sessions or "credentials" not in sessions[request.state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[request.state]["credentials"]
        credentials = Credentials(**creds_dict)
        service = build('drive', 'v3', credentials=credentials)
        
        file_metadata = {
            'name': request.folderName,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        
        # If parent folder specified, add to parents
        if request.parentFolderId:
            file_metadata['parents'] = [request.parentFolderId]
        
        folder = service.files().create(
            body=file_metadata,
            fields='id, name, webViewLink'
        ).execute()
        
        print(f"[DRIVE] Created folder: {folder.get('name')} (ID: {folder.get('id')})")
        
        return {
            "folderId": folder.get('id'),
            "folderName": folder.get('name'),
            "webViewLink": folder.get('webViewLink')
        }
    except Exception as e:
        print(f"[DRIVE] Error creating folder: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create folder: {str(e)}")

@app.post("/api/drive/upload")
async def upload_to_drive(request: UploadFileRequest):
    """Upload a file to Google Drive folder"""
    try:
        if request.state not in sessions or "credentials" not in sessions[request.state]:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        creds_dict = sessions[request.state]["credentials"]
        credentials = Credentials(**creds_dict)
        service = build('drive', 'v3', credentials=credentials)
        
        # Decode base64 file data
        file_bytes = base64.b64decode(request.fileData)
        
        print(f"[DRIVE] Uploading {request.fileName} to folder {request.folderId}, size: {len(file_bytes)} bytes")
        
        # Create file metadata
        file_metadata = {
            'name': request.fileName,
            'parents': [request.folderId]
        }
        
        # Upload file using MediaIoBaseUpload
        from googleapiclient.http import MediaIoBaseUpload
        
        media = MediaIoBaseUpload(
            io.BytesIO(file_bytes),
            mimetype=request.mimeType,
            resumable=True
        )
        
        file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, name, webViewLink'
        ).execute()
        
        print(f"[DRIVE] File uploaded successfully: {file.get('name')} (ID: {file.get('id')})")
        
        return {
            "fileId": file.get('id'),
            "fileName": file.get('name'),
            "webViewLink": file.get('webViewLink')
        }
    except Exception as e:
        print(f"[DRIVE] Error uploading file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to upload file: {str(e)}")

# Helper functions
def extract_drive_id(link: str) -> str:
    """Extract file or folder ID from Google Drive link"""
    import re
    
    patterns = [
        r'/file/d/([a-zA-Z0-9_-]+)',
        r'/folders/([a-zA-Z0-9_-]+)',
        r'[?&]id=([a-zA-Z0-9_-]+)',
        r'/document/d/([a-zA-Z0-9_-]+)',
        r'/spreadsheets/d/([a-zA-Z0-9_-]+)',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, link)
        if match:
            return match.group(1)
    
    return link.strip()

def is_meaningful_text(text):
    """Check if text contains meaningful content"""
    cleaned = re.sub(r'[\W_]+', '', text)
    return bool(cleaned.strip())

def is_decorative_only(text):
    """Check if text is decorative only (symbols, single letters, etc.)"""
    stripped = text.strip()
    return not stripped or re.fullmatch(r"[^\w\s]+", stripped) or re.fullmatch(r"[A-Z]", stripped)

def sanitize_response(text: str) -> str:
    """Remove any <think>...</think> tokens and trim the response."""
    if not text:
        return text
    text = re.sub(r'<think>.*?</think>', '', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'</?think>', '', text, flags=re.IGNORECASE)
    return text.strip()

def call_gemini_batch_api(client, prompt, model, logs=None):
    """
    Synchronous function to call Gemini API for batch processing with token tracking.
    This will be run in a thread executor.
    """
    for attempt in range(MAX_RETRIES):
        try:
            if logs is not None:
                logs.append(f"[BATCH API] Attempt {attempt + 1}/{MAX_RETRIES} - Model: {model}")
            
            # Generate content with JSON response format
            response = client.models.generate_content(
                model=model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.0,
                    response_mime_type="application/json"
                )
            )
            
            result_text = response.text.strip()
            
            # Get token usage from response
            usage = response.usage_metadata
            actual_input_tokens = usage.prompt_token_count
            output_tokens = usage.candidates_token_count
            total_tokens = usage.total_token_count
            
            if logs is not None:
                logs.append(f"[SUCCESS] Received response ({len(result_text)} chars)")
                logs.append(f"[TOKENS] Input: {actual_input_tokens}, Output: {output_tokens}, Total: {total_tokens}")
            
            return {
                'text': result_text,
                'input_tokens': actual_input_tokens,
                'output_tokens': output_tokens,
                'total_tokens': total_tokens
            }
            
        except Exception as e:
            if logs is not None:
                logs.append(f"[ERROR] Attempt {attempt + 1} failed: {str(e)}")
            
            if attempt < MAX_RETRIES - 1:
                time.sleep(RETRY_DELAY)
    
    if logs is not None:
        logs.append("[FAILED] All retry attempts exhausted")
    return None

async def call_gemini_batch_async(executor, client, prompt, model, logs=None):
    """Async wrapper for batch API call using ThreadPoolExecutor."""
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(executor, call_gemini_batch_api, client, prompt, model, logs)
    return result

async def call_openrouter_batch_api(session, prompt, model, openrouter_api_key, log_list=None):
    """
    Call OpenRouter API for batch processing with token tracking.
    """
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {openrouter_api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": os.getenv("OPENROUTER_REFERER", "https://yourdomain.com"),
        "X-Title": "EasyTranslate"
    }
    
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.0,
        "response_format": {"type": "json_object"}
    }
    
    for attempt in range(MAX_RETRIES):
        try:
            if log_list is not None:
                log_entry = f"[BATCH API] Attempt {attempt + 1}/{MAX_RETRIES} - Model: {model}"
                log_list.append(log_entry)
            
            async with session.post(url, headers=headers, json=payload) as resp:
                resp.raise_for_status()
                data = await resp.json()
                
                result = data["choices"][0]["message"]["content"].strip()
                
                # Get token usage from response
                usage = data.get("usage", {})
                input_tokens = usage.get("prompt_tokens", 0)
                output_tokens = usage.get("completion_tokens", 0)
                total_tokens = usage.get("total_tokens", input_tokens + output_tokens)
                
                if log_list is not None:
                    log_list.append(f"[SUCCESS] Received response ({len(result)} chars)")
                    log_list.append(f"[TOKENS] Input: {input_tokens}, Output: {output_tokens}, Total: {total_tokens}")
                
                return {
                    'text': result,
                    'input_tokens': input_tokens,
                    'output_tokens': output_tokens,
                    'total_tokens': total_tokens
                }
                
        except Exception as e:
            if log_list is not None:
                log_list.append(f"[ERROR] Attempt {attempt + 1} failed: {str(e)}")
            
            if attempt < MAX_RETRIES - 1:
                await asyncio.sleep(RETRY_DELAY)
    
    if log_list is not None:
        log_list.append("[FAILED] All retry attempts exhausted")
    return None

def parse_structured_response(response_text, expected_count, logs=None):
    """Parse structured JSON response into individual translations."""
    import json
    
    try:
        # Clean response text (remove markdown code blocks if present)
        cleaned_text = response_text.strip()
        if cleaned_text.startswith('```json'):
            cleaned_text = cleaned_text[7:]
        if cleaned_text.endswith('```'):
            cleaned_text = cleaned_text[:-3]
        cleaned_text = cleaned_text.strip()
        
        # Parse JSON
        parsed_response = json.loads(cleaned_text)
        
        if logs:
            logs.append(f"[JSON] Successfully parsed structured response")
        
        # Extract translations in order
        translations = []
        if 'translations' in parsed_response:
            # Sort by ID to ensure correct order
            translation_items = sorted(parsed_response['translations'], key=lambda x: x.get('id', 0))
            
            for item in translation_items:
                if 'translation' in item:
                    translations.append(item['translation'])
                    
            if logs:
                logs.append(f"[JSON] Extracted {len(translations)} translations from structured response")
        
        return translations
        
    except json.JSONDecodeError as e:
        if logs:
            logs.append(f"[JSON ERROR] Failed to parse JSON: {str(e)}")
            logs.append(f"[FALLBACK] Attempting fallback parsing...")
        
        # Fallback: split by double newlines
        lines = response_text.split('\n\n')
        alt_translations = [t.strip() for t in lines if t.strip() and not t.startswith('{')]
        return alt_translations[:expected_count]
    
    except Exception as e:
        if logs:
            logs.append(f"[PARSE ERROR] {str(e)}")
        return []

async def translate_document_content_async(file_bytes: bytes, file_name: str, language: str, model: str, api_key: str) -> TranslateResponse:
    """Translate document content using Gemini with async batch processing - matches Streamlit logic"""
    import json
    
    print(f"[TRANSLATOR] Initializing Gemini client")
    client = genai.Client(api_key=api_key)
    
    # Load document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = doc.paragraphs
    
    print(f"[TRANSLATOR] Document has {len(paragraphs)} total paragraphs")
    
    translated_content = []
    logs = []
    total_input_tokens = 0
    total_output_tokens = 0
    
    logs.append(f"[START] Batch translation started for language: {language}")
    logs.append(f"[INFO] Source file: {file_name}")
    logs.append(f"[INFO] Document has {len(paragraphs)} total paragraphs")
    logs.append(f"[INFO] Batch size: {BATCH_SIZE} paragraphs per request")
    logs.append(f"[INFO] Using Gemini model: {model}")
    logs.append(f"[INFO] Processing in memory - no files saved to disk")
    
    # Prepare paragraph batches with filtering logic
    paragraph_batches = []
    current_batch = []
    
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        original = para.text.strip()

        # Skip single uppercase letters followed by uppercase text
        if re.fullmatch(r"[A-Z]", original) and i + 1 < len(paragraphs) and paragraphs[i + 1].text.strip()[:1].isupper():
            p = para._element
            p.getparent().remove(p)
            paragraphs = doc.paragraphs
            logs.append(f"[SKIP] Removed single uppercase letter at para {i}")
            continue

        # Skip empty or decorative text
        if not original or not is_meaningful_text(original) or is_decorative_only(original):
            i += 1
            continue

        word_count = len(original.split())
        is_heading = para.style.name.lower().startswith("heading") or para.alignment == 1

        # Skip single-word non-heading paragraphs
        if word_count <= 1:
            if not original.isupper() and not is_heading:
                i += 1
                continue

        # Add to current batch
        current_batch.append((i, para, original))
        
        # If batch is full or we're at the end, save it
        if len(current_batch) >= BATCH_SIZE or i == len(paragraphs) - 1:
            if current_batch:
                paragraph_batches.append(current_batch)
                current_batch = []
        
        i += 1

    total_batches = len(paragraph_batches)
    logs.append(f"[QUEUE] Prepared {total_batches} batches for translation")
    logs.append(f"[PROCESSING] Starting parallel batch API requests (max 4 concurrent)...")
    
    # Create thread pool executor for async processing
    executor = ThreadPoolExecutor(max_workers=4)
    
    async def process_batch_gemini(batch_idx, batch):
        """Process a single batch for Gemini - returns logs separately"""
        batch_logs = []  # Separate logs for this batch
        batch_paragraphs = [item[2] for item in batch]  # Extract text
        
        print(f"[TRANSLATOR] Processing batch {batch_idx + 1}/{total_batches} ({len(batch_paragraphs)} paragraphs)")
        batch_logs.append(f"[BATCH {batch_idx + 1}/{total_batches}] Processing {len(batch_paragraphs)} paragraphs...")
        
        # Create batch prompt
        prompt = create_batch_prompt(batch_paragraphs, language)
        
        # Call API asynchronously
        batch_result = await call_gemini_batch_async(executor, client, prompt, model, batch_logs)
        
        return batch_idx, batch, batch_paragraphs, batch_result, batch_logs
    
    # Create all tasks for parallel processing
    tasks = [
        process_batch_gemini(batch_idx, batch)
        for batch_idx, batch in enumerate(paragraph_batches)
    ]
    
    # Execute all tasks in parallel (limited by ThreadPoolExecutor max_workers)
    results = await asyncio.gather(*tasks)
    
    # Process results in order and merge logs
    for batch_idx, batch, batch_paragraphs, batch_result, batch_logs in results:
        # Add this batch's logs in order
        logs.extend(batch_logs)
        
        if batch_result:
            total_input_tokens += batch_result['input_tokens']
            total_output_tokens += batch_result['output_tokens']
            
            # Parse structured response
            batch_translations = parse_structured_response(batch_result['text'], len(batch_paragraphs), logs)
            
            logs.append(f"[BATCH {batch_idx + 1}] Received {len(batch_translations)} translations")
            
            # Validate we got the expected number of translations
            if len(batch_translations) != len(batch_paragraphs):
                logs.append(f"[WARNING] Expected {len(batch_paragraphs)} translations, got {len(batch_translations)}")
                # Pad with empty strings if we're short
                while len(batch_translations) < len(batch_paragraphs):
                    batch_translations.append('[Translation missing]')
                # Trim if we got too many
                batch_translations = batch_translations[:len(batch_paragraphs)]
            
            # Apply translations to document (maintains order)
            for (para_idx, para, original), translation in zip(batch, batch_translations):
                if translation and translation.strip():
                    translation = sanitize_response(translation)
                    translated_content.append(translation)
                    
                    # Update paragraph text in memory
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = translation
                    else:
                        para.add_run(translation)
                else:
                    fallback_text = f"[Translation failed for paragraph {para_idx}]"
                    translated_content.append(fallback_text)
        else:
            # Batch failed completely, add fallback content
            logs.append(f"[BATCH ERROR] Batch {batch_idx + 1} failed completely, adding fallback content")
            for para_idx, para, original in batch:
                fallback_text = f"[Batch translation failed - Original: {original[:100]}...]"
                translated_content.append(fallback_text)
    
    # Shutdown executor
    executor.shutdown(wait=False)
    
    # Save document to memory buffer
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    logs.append(f"[SAVE] Document saved to memory buffer")
    logs.append(f"[TOKENS] Final usage - Input: {total_input_tokens}, Output: {total_output_tokens}, Total: {total_input_tokens + total_output_tokens}")
    logs.append(f"[DONE] Translation complete!")
    
    # Convert to base64
    translated_base64 = base64.b64encode(output_buffer.read()).decode('utf-8')
    
    print("[TRANSLATE] Complete, returning response")
    
    return TranslateResponse(
        translatedDocument=translated_base64,
        logs=logs,
        stats={
            "totalInputTokens": total_input_tokens,
            "totalOutputTokens": total_output_tokens,
            "totalTokens": total_input_tokens + total_output_tokens,
            "paragraphCount": len(translated_content),
            "translatedText": "\n\n".join(translated_content)  # Add full text for preview
        }
    )

async def translate_document_content_async_openrouter(file_bytes: bytes, file_name: str, language: str, model: str, api_key: str) -> TranslateResponse:
    """Translate document content using OpenRouter with async batch processing"""
    import json
    
    print(f"[TRANSLATOR] Initializing OpenRouter client")
    
    # Load document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = doc.paragraphs
    
    print(f"[TRANSLATOR] Document has {len(paragraphs)} total paragraphs")
    
    translated_content = []
    logs = []
    total_input_tokens = 0
    total_output_tokens = 0
    
    logs.append(f"[START] Batch translation started for language: {language}")
    logs.append(f"[INFO] Source file: {file_name}")
    logs.append(f"[INFO] Document has {len(paragraphs)} total paragraphs")
    logs.append(f"[INFO] Batch size: {BATCH_SIZE} paragraphs per request")
    logs.append(f"[INFO] Using OpenRouter model: {model}")
    logs.append(f"[INFO] Processing in memory - no files saved to disk")
    
    # Prepare paragraph batches with filtering logic
    paragraph_batches = []
    current_batch = []
    
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        original = para.text.strip()

        # Skip single uppercase letters followed by uppercase text
        if re.fullmatch(r"[A-Z]", original) and i + 1 < len(paragraphs) and paragraphs[i + 1].text.strip()[:1].isupper():
            p = para._element
            p.getparent().remove(p)
            paragraphs = doc.paragraphs
            logs.append(f"[SKIP] Removed single uppercase letter at para {i}")
            continue

        # Skip empty or decorative text
        if not original or not is_meaningful_text(original) or is_decorative_only(original):
            i += 1
            continue

        word_count = len(original.split())
        is_heading = para.style.name.lower().startswith("heading") or para.alignment == 1

        # Skip single-word non-heading paragraphs
        if word_count <= 1:
            if not original.isupper() and not is_heading:
                i += 1
                continue

        # Add to current batch
        current_batch.append((i, para, original))
        
        # If batch is full or we're at the end, save it
        if len(current_batch) >= BATCH_SIZE or i == len(paragraphs) - 1:
            if current_batch:
                paragraph_batches.append(current_batch)
                current_batch = []
        
        i += 1

    total_batches = len(paragraph_batches)
    logs.append(f"[QUEUE] Prepared {total_batches} batches for translation")
    logs.append(f"[PROCESSING] Starting parallel batch API requests (max 4 concurrent)...")
    
    # Create semaphore to limit concurrent requests
    semaphore = asyncio.Semaphore(4)
    
    async def process_batch_with_semaphore(batch_idx, batch, session):
        """Process a single batch with semaphore control - returns logs separately"""
        batch_logs = []  # Separate logs for this batch
        
        async with semaphore:
            batch_paragraphs = [item[2] for item in batch]  # Extract text
            
            print(f"[TRANSLATOR] Processing batch {batch_idx + 1}/{total_batches} ({len(batch_paragraphs)} paragraphs)")
            batch_logs.append(f"[BATCH {batch_idx + 1}/{total_batches}] Processing {len(batch_paragraphs)} paragraphs...")
            
            # Create batch prompt
            prompt = create_batch_prompt(batch_paragraphs, language)
            
            # Call API asynchronously
            batch_result = await call_openrouter_batch_api(session, prompt, model, api_key, batch_logs)
            
            return batch_idx, batch, batch_paragraphs, batch_result, batch_logs
    
    # Create aiohttp session for async requests
    async with aiohttp.ClientSession() as session:
        # Create all tasks for parallel processing
        tasks = [
            process_batch_with_semaphore(batch_idx, batch, session)
            for batch_idx, batch in enumerate(paragraph_batches)
        ]
        
        # Execute all tasks in parallel (limited by semaphore)
        results = await asyncio.gather(*tasks)
        
        # Process results in order and merge logs
        for batch_idx, batch, batch_paragraphs, batch_result, batch_logs in results:
            # Add this batch's logs in order
            logs.extend(batch_logs)
            
            if batch_result:
                total_input_tokens += batch_result['input_tokens']
                total_output_tokens += batch_result['output_tokens']
                
                # Parse structured response
                batch_translations = parse_structured_response(batch_result['text'], len(batch_paragraphs), logs)
                
                logs.append(f"[BATCH {batch_idx + 1}] Received {len(batch_translations)} translations")
                
                # Validate we got the expected number of translations
                if len(batch_translations) != len(batch_paragraphs):
                    logs.append(f"[WARNING] Expected {len(batch_paragraphs)} translations, got {len(batch_translations)}")
                    # Pad with empty strings if we're short
                    while len(batch_translations) < len(batch_paragraphs):
                        batch_translations.append('[Translation missing]')
                    # Trim if we got too many
                    batch_translations = batch_translations[:len(batch_paragraphs)]
                
                # Apply translations to document (maintains order)
                for (para_idx, para, original), translation in zip(batch, batch_translations):
                    if translation and translation.strip():
                        translation = sanitize_response(translation)
                        translated_content.append(translation)
                        
                        # Update paragraph text in memory
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = translation
                        else:
                            para.add_run(translation)
                    else:
                        fallback_text = f"[Translation failed for paragraph {para_idx}]"
                        translated_content.append(fallback_text)
            else:
                # Batch failed completely, add fallback content
                logs.append(f"[BATCH ERROR] Batch {batch_idx + 1} failed completely, adding fallback content")
                for para_idx, para, original in batch:
                    fallback_text = f"[Batch translation failed - Original: {original[:100]}...]"
                    translated_content.append(fallback_text)
    
    # Save document to memory buffer
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    logs.append(f"[SAVE] Document saved to memory buffer")
    logs.append(f"[TOKENS] Final usage - Input: {total_input_tokens}, Output: {total_output_tokens}, Total: {total_input_tokens + total_output_tokens}")
    logs.append(f"[DONE] Translation complete!")
    
    # Convert to base64
    translated_base64 = base64.b64encode(output_buffer.read()).decode('utf-8')
    
    print("[TRANSLATE] Complete, returning response")
    
    return TranslateResponse(
        translatedDocument=translated_base64,
        logs=logs,
        stats={
            "totalInputTokens": total_input_tokens,
            "totalOutputTokens": total_output_tokens,
            "totalTokens": total_input_tokens + total_output_tokens,
            "paragraphCount": len(translated_content),
            "translatedText": "\n\n".join(translated_content)  # Add full text for preview
        }
    )

def create_batch_prompt(paragraphs: List[str], language: str) -> str:
    """Create batch translation prompt"""
    BATCH_PROMPT_TEMPLATE = """
Here are {count} passages from a manuscript that we want to translate into {language} to make it easily readable for everyone. These are the original English versions of the passages:

Can you translate each passage into fluent and natural {language}? Please rewrite them and do not plagiarize. Only give the translated text for each passage. 
Make sure that the reading experience "flows", for example by not using the same words & sentence structures too often. Only translate, do not mention anything else. Please format the text properly without separating lines between the passages and exactly in the same structure so I can easily copy/paste it entirely into the manuscript of the book. You can remove numbers if that makes the reading experience better. You must translate/rewrite everything exactly and do not shorten it. Keep the quotes in their formatted way if applicable.

This is important for my career: Please try to translate everything sentence-by-sentence and do not make it shorter!
Important: Please Translate everything sentence-by-sentence in fluent {language} and do not make it shorter!

This is crucial: When you encounter general terminology such as "chapter," "the end,", "introduction" etc., or phrases that are widely used across books, do not vary or rephrase them—this is to ensure consistency. Always translate the text into the targetted language
Always use the most widely accepted or officially published book title, chapter name, or term if multiple synonyms are possible — e.g., use Meditations, not Reflections; Letters from a Stoic, not Stoic Letters. "Introduction", not "getting started"
Also, always choose the most commonly published or recognized title, name or term in the target language — e.g., De Gedaanteverwisseling, not Metamorfose in Dutch. "Arme Peter" ipv "poor Peter" or "Armzalige Peter"

IMPORTANT: Return your response as a valid JSON object with the following structure:
{{
  "translations": [
    {{
      "id": 1,
      "translation": "translated text for passage 1"
    }},
    {{
      "id": 2,
      "translation": "translated text for passage 2"
    }}
    // ... continue for all passages
  ]
}}

Original Passages:
{passages}
"""
    
    passages_text = ""
    for i, para in enumerate(paragraphs, 1):
        passages_text += f'\nPassage {i} (ID: {i}):\n"""\n{para}\n"""\n'
    
    return BATCH_PROMPT_TEMPLATE.format(
        count=len(paragraphs),
        language=language,
        passages=passages_text
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=7860)
