"""
Table of Contents Handler
Detects TOC entries, extracts titles, and converts matching paragraphs to Heading 2 format.
"""

import re
from typing import List, Tuple, Optional
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.shared import Pt


def detect_toc_in_first_pages(doc, max_pages: int = 10) -> Tuple[List[Paragraph], int]:
    """
    Detect Table of Contents entries in the first N pages of the document.
    
    Args:
        doc: The document to search
        max_pages: Maximum number of pages to search (default: 10)
        
    Returns:
        Tuple of (list of TOC paragraphs, index where TOC ends)
    """
    # Estimate paragraphs per page (rough estimate: ~50 paragraphs per page)
    max_paragraphs = max_pages * 50
    
    toc_paragraphs = []
    toc_started = False
    toc_ended = False
    toc_end_index = -1
    
    for i, para in enumerate(doc.paragraphs[:max_paragraphs]):
        text = para.text.strip()
        
        # Check if this is a TOC heading
        if not toc_started:
            text_lower = text.lower()
            if any(keyword in text_lower for keyword in [
                'table of contents', 'contents', 'table des matières',
                'tabla de contenidos', 'índice', 'contenido', 'índice de contenidos',
                'sommaire', 'inhaltsverzeichnis'
            ]):
                toc_started = True
                continue
        
        if not toc_started:
            continue
        
        # Check if TOC has ended (we hit a major heading or section)
        if toc_started and not toc_ended:
            # TOC ends when we hit:
            # 1. A paragraph that looks like a chapter/section heading (ALL CAPS, short)
            # 2. A paragraph with Heading style
            # 3. Empty paragraph followed by non-TOC content
            if text and len(text) > 0:
                # Check if this looks like a chapter heading (not a TOC entry)
                if (text.isupper() and 3 < len(text) < 100 and 
                    not re.search(r'\t.*\d+', text) and  # No tab + number pattern
                    not re.search(r'\s{3,}\d+\s*$', text)):  # No trailing page number
                    toc_ended = True
                    toc_end_index = i
                    break
                
                # Check if it has a Heading style
                try:
                    if para.style and para.style.name and para.style.name.startswith('Heading'):
                        toc_ended = True
                        toc_end_index = i
                        break
                except:
                    pass
        
        # Detect TOC entry patterns
        if is_toc_entry(para, text):
            toc_paragraphs.append((i, para, text))
    
    return toc_paragraphs, toc_end_index


def is_toc_entry(para: Paragraph, text: str) -> bool:
    """
    Determine if a paragraph is a TOC entry.
    
    TOC entry indicators:
    - Contains tab character followed by page number
    - Has pattern: Title + tab/spaces + number
    - Has hyperlink formatting
    - Contains TOC field codes
    """
    if not text or len(text.strip()) < 3:
        return False
    
    # Pattern 1: Tab followed by digits (classic TOC pattern)
    if '\t' in text and re.search(r'\t+\d+', text):
        return True
    
    # Pattern 2: Multiple spaces followed by trailing number
    if re.search(r'[A-Za-z].{5,}\s{3,}\d+\s*$', text):
        return True
    
    # Pattern 3: Dots/leaders followed by number
    if re.search(r'[A-Za-z].{5,}\.{2,}\d+', text):
        return True
    
    # Pattern 4: Check for hyperlink in XML
    try:
        hyperlinks = list(para._p.iter(qn('w:hyperlink')))
        if hyperlinks and re.search(r'\d', text):
            return True
    except:
        pass
    
    # Pattern 5: Check for TOC field codes
    if re.search(r'TOC\s+\\[a-z]', text, re.IGNORECASE):
        return True
    
    return False


def extract_toc_titles(toc_paragraphs: List[Tuple[int, Paragraph, str]]) -> List[str]:
    """
    Extract clean titles from TOC entries.
    
    Removes:
    - Page numbers
    - Tabs
    - Field codes
    - Duplicate text
    
    Returns list of clean titles.
    """
    titles = []
    
    for para_idx, para, text in toc_paragraphs:
        # Remove TOC field codes
        text = re.sub(r'TOC\s+\\[a-z]+(\s+"[^"]*")?', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\\[a-z]+(\s+"[^"]*")?', '', text, flags=re.IGNORECASE)
        
        # Remove page numbers (trailing digits)
        text = re.sub(r'\s*\d+\s*$', '', text)
        
        # Remove tabs and extra whitespace
        text = re.sub(r'\t+', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Remove duplicate patterns (e.g., "Title Title" -> "Title")
        words = text.split()
        if len(words) >= 2:
            mid = len(words) // 2
            first_half = ' '.join(words[:mid])
            second_half = ' '.join(words[mid:])
            if first_half == second_half:
                text = first_half
        
        # Clean up any remaining artifacts
        text = re.sub(r'\.{3,}', '', text)  # Remove leader dots
        text = text.strip()
        
        if text and len(text) > 2:
            titles.append(text)
    
    return titles


def find_matching_paragraphs(doc, titles: List[str], start_index: int = 0) -> List[Tuple[int, Paragraph]]:
    """
    Find paragraphs in the document that match TOC titles.
    
    Uses fuzzy matching to handle:
    - Case differences
    - Extra whitespace
    - Partial matches
    
    Args:
        doc: The document to search
        titles: List of titles to match
        start_index: Only search paragraphs starting from this index (to skip TOC and content before it)
    
    Returns list of (paragraph_index, paragraph) tuples.
    """
    matches = []
    
    for title in titles:
        title_clean = re.sub(r'\s+', ' ', title.strip())
        title_lower = title_clean.lower()
        
        best_match = None
        best_score = 0
        
        # Only search paragraphs AFTER the start_index (after TOC)
        for i, para in enumerate(doc.paragraphs[start_index:], start=start_index):
            para_text = para.text.strip()
            if not para_text:
                continue
            
            para_clean = re.sub(r'\s+', ' ', para_text)
            para_lower = para_clean.lower()
            
            # Exact match (case-insensitive)
            if para_lower == title_lower:
                best_match = (i, para)
                best_score = 100
                break
            
            # Check if title is contained in paragraph (for longer paragraphs)
            if title_lower in para_lower and len(title_clean) >= 10:
                # Prefer matches at the start of the paragraph
                if para_lower.startswith(title_lower):
                    score = 90
                else:
                    score = 70
                
                if score > best_score:
                    best_match = (i, para)
                    best_score = score
            
            # Check if paragraph starts with title (for shorter matches)
            if para_lower.startswith(title_lower[:min(50, len(title_lower))]):
                score = 80
                if score > best_score:
                    best_match = (i, para)
                    best_score = score
        
        if best_match and best_score >= 70:
            matches.append(best_match)
    
    return matches


def convert_to_heading_2(doc, matches: List[Tuple[int, Paragraph]]) -> int:
    """
    Convert matching paragraphs to Heading 2 style.
    
    Returns: count of paragraphs converted
    """
    converted_count = 0
    
    for para_idx, para in matches:
        try:
            # Set to Heading 2 style
            para.style = doc.styles['Heading 2']
            converted_count += 1
        except (KeyError, AttributeError):
            # If Heading 2 doesn't exist, try to create it or use Heading 1
            try:
                para.style = doc.styles['Heading 1']
                converted_count += 1
            except:
                # If no heading styles exist, skip
                pass
    
    return converted_count


def remove_toc_paragraphs(doc, toc_paragraphs: List[Tuple[int, Paragraph, str]]) -> int:
    """
    Remove TOC paragraphs from the document.
    
    Returns: count of paragraphs removed
    """
    # Sort by index in reverse order to avoid index shifting issues
    sorted_toc = sorted(toc_paragraphs, key=lambda x: x[0], reverse=True)
    
    removed_count = 0
    for para_idx, para, text in sorted_toc:
        try:
            # Get the paragraph element and remove it
            p_element = para._element
            p_parent = p_element.getparent()
            if p_parent is not None:
                p_parent.remove(p_element)
                removed_count += 1
        except:
            pass
    
    return removed_count


def insert_toc_placeholder(doc, toc_heading_index: Optional[int] = None) -> bool:
    """
    Insert a TOC placeholder field after the TOC heading.
    
    Args:
        doc: The document
        toc_heading_index: Index of the TOC heading paragraph (if known)
        
    Returns: True if placeholder was inserted, False otherwise
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    
    # Find TOC heading if not provided
    if toc_heading_index is None:
        toc_heading_index = -1
        for i, para in enumerate(doc.paragraphs):
            text_lower = para.text.strip().lower()
            if any(keyword in text_lower for keyword in [
                'table of contents', 'contents', 'table des matières',
                'tabla de contenidos', 'índice', 'contenido', 'índice de contenidos',
                'sommaire', 'inhaltsverzeichnis'
            ]):
                toc_heading_index = i
                break
    
    if toc_heading_index < 0:
        return False
    
    # Insert TOC field after the heading
    # Since programmatically created TOC fields may not always work in Word,
    # we'll add both a field attempt AND clear instructions
    
    if toc_heading_index + 1 < len(doc.paragraphs):
        toc_para = doc.paragraphs[toc_heading_index + 1].insert_paragraph_before()
    else:
        toc_para = doc.add_paragraph()
    
    # Try to create a TOC field
    try:
        # Create a single run that will contain the field
        field_run = toc_para.add_run()
        
        # Field begin
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        field_run._r.append(fldChar_begin)
        
        # Field instruction (TOC command)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        field_run._r.append(instrText)
        
        # Field separator
        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')
        field_run._r.append(fldChar_sep)
        
        # Result text (placeholder)
        result_run = toc_para.add_run('Press F9 to update Table of Contents')
        result_run.font.italic = True
        try:
            result_run.font.color.rgb = RGBColor(100, 100, 100)
        except:
            pass
        
        # Field end
        end_run = toc_para.add_run()
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        end_run._r.append(fldChar_end)
        
        print(f"[TOC FIELD] Inserted TOC field at paragraph index {toc_heading_index + 1}")
    except Exception as e:
        print(f"[TOC FIELD] Error creating field: {e}, using instruction paragraph instead")
        # Fallback: Add instruction paragraph
        toc_para.clear()
        instruction = toc_para.add_run('To generate Table of Contents: Go to References > Table of Contents > Automatic Table 1')
        instruction.font.italic = True
        try:
            instruction.font.color.rgb = RGBColor(100, 100, 100)
        except:
            pass
    
    # Always add a helpful instruction paragraph below
    if toc_heading_index + 2 < len(doc.paragraphs):
        instruction_para = doc.paragraphs[toc_heading_index + 2].insert_paragraph_before()
    else:
        instruction_para = doc.add_paragraph()
    
    instruction_text = instruction_para.add_run('(All headings have been converted to Heading 2 style. Word will automatically generate the TOC from these headings.)')
    instruction_text.font.size = Pt(9)
    instruction_text.font.italic = True
    try:
        instruction_text.font.color.rgb = RGBColor(150, 150, 150)
    except:
        pass
    
    return True


def process_toc_before_translation(doc) -> dict:
    """
    Main function to process TOC before translation.
    
    Steps:
    1. Detect TOC in first 10 pages
    2. Extract titles from TOC entries
    3. Find matching paragraphs and convert to Heading 2
    4. Remove TOC entries
    5. Insert placeholder
    
    Returns: dict with processing results
    """
    print("[TOC PROCESSING] Starting TOC processing...")
    results = {
        'toc_found': False,
        'toc_entries_count': 0,
        'titles_extracted': 0,
        'paragraphs_converted': 0,
        'toc_removed': 0,
        'placeholder_inserted': False,
        'toc_end_index': -1
    }
    
    # Step 1: Detect TOC
    print(f"[TOC PROCESSING] Searching first 10 pages ({len(doc.paragraphs)} total paragraphs)...")
    toc_paragraphs, toc_end_index = detect_toc_in_first_pages(doc, max_pages=10)
    print(f"[TOC PROCESSING] Found {len(toc_paragraphs)} potential TOC entries")
    
    if not toc_paragraphs:
        print("[TOC PROCESSING] No TOC entries found, returning early")
        return results
    
    results['toc_found'] = True
    results['toc_entries_count'] = len(toc_paragraphs)
    results['toc_end_index'] = toc_end_index
    print(f"[TOC PROCESSING] TOC found! {len(toc_paragraphs)} entries, ends at index {toc_end_index}")
    
    # Step 2: Extract titles
    titles = extract_toc_titles(toc_paragraphs)
    results['titles_extracted'] = len(titles)
    print(f"[TOC PROCESSING] Extracted {len(titles)} titles from TOC entries")
    
    if not titles:
        print("[TOC PROCESSING] No titles extracted, returning early")
        return results
    
    # Step 3: Find matching paragraphs and convert to Heading 2
    # CRITICAL: Only search paragraphs AFTER the TOC to avoid converting titles that appear before TOC
    # We want to preserve the formatting of titles that appear before the TOC (like the first book title)
    # Only convert titles that appear in the actual document content after the TOC
    
    # Determine where to start searching - after the TOC ends
    if toc_end_index > 0:
        # TOC end was detected - search from there
        search_start_index = toc_end_index
    elif toc_paragraphs:
        # TOC end wasn't detected, but we have TOC entries - start after the last TOC entry
        last_toc_idx = max(toc_paragraphs, key=lambda x: x[0])[0]
        search_start_index = last_toc_idx + 1
    else:
        # Fallback: start from beginning (shouldn't happen if TOC was found)
        search_start_index = 0
    
    print(f"[TOC PROCESSING] Finding matching paragraphs for {len(titles)} titles (searching from paragraph {search_start_index} onwards to skip TOC and content before it)...")
    matches = find_matching_paragraphs(doc, titles, start_index=search_start_index)
    print(f"[TOC PROCESSING] Found {len(matches)} matching paragraphs (all after TOC)")
    converted = convert_to_heading_2(doc, matches)
    results['paragraphs_converted'] = converted
    print(f"[TOC PROCESSING] Converted {converted} paragraphs to Heading 2")
    
    # Step 4: Remove TOC entries
    removed = remove_toc_paragraphs(doc, toc_paragraphs)
    results['toc_removed'] = removed
    print(f"[TOC PROCESSING] Removed {removed} TOC entry paragraphs")
    
    # Step 5: Insert placeholder
    # Find TOC heading index
    toc_heading_idx = None
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.strip().lower()
        if any(keyword in text_lower for keyword in [
            'table of contents', 'contents', 'table des matières',
            'tabla de contenidos', 'índice', 'contenido', 'índice de contenidos',
            'sommaire', 'inhaltsverzeichnis'
        ]):
            toc_heading_idx = i
            break
    
    if toc_heading_idx is not None:
        placeholder_inserted = insert_toc_placeholder(doc, toc_heading_idx)
        results['placeholder_inserted'] = placeholder_inserted
    
    return results
