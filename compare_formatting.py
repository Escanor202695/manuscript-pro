"""
Compare formatting between original and translated DOCX files
"""

from docx import Document
import os

def analyze_document_formatting(file_path):
    """Analyze formatting of a DOCX document"""
    doc = Document(file_path)
    
    print(f"\n{'='*60}")
    print(f"Analyzing: {os.path.basename(file_path)}")
    print(f"{'='*60}\n")
    
    para_count = 0
    
    for i, para in enumerate(doc.paragraphs):
        # Skip empty paragraphs
        if not para.text.strip():
            continue
            
        para_count += 1
        print(f"Paragraph {para_count} (Index {i}):")
        print(f"  Text: {para.text[:50]}..." if len(para.text) > 50 else f"  Text: {para.text}")
        print(f"  Style: {para.style.name if para.style else 'None'}")
        print(f"  Alignment: {para.alignment}")
        print(f"  Number of runs: {len(para.runs)}")
        
        # Analyze runs (text segments with formatting)
        for j, run in enumerate(para.runs):
            if run.text.strip():  # Only show non-empty runs
                print(f"\n  Run {j}:")
                print(f"    Text: '{run.text}'")
                print(f"    Bold: {run.bold}")
                print(f"    Italic: {run.italic}")
                print(f"    Underline: {run.underline}")
                if run.font.name:
                    print(f"    Font: {run.font.name}")
                if run.font.size:
                    print(f"    Size: {run.font.size}")
        
        print("\n" + "-"*40 + "\n")

def compare_documents():
    """Compare original and translated documents"""
    original_path = "/Users/sakibchowdhury/Desktop/code/Translation Manuscript/translation/Untitled document (2).docx"
    translated_path = "/Users/sakibchowdhury/Desktop/code/Translation Manuscript/translation/Untitled document (2)_spanish_translated_spanish_translated.docx"
    
    print("DOCUMENT FORMATTING COMPARISON")
    print("="*60)
    
    # Analyze original
    analyze_document_formatting(original_path)
    
    # Analyze translated
    analyze_document_formatting(translated_path)
    
    # Quick comparison
    print("\n" + "="*60)
    print("QUICK COMPARISON SUMMARY")
    print("="*60)
    
    doc_orig = Document(original_path)
    doc_trans = Document(translated_path)
    
    orig_paras = [p for p in doc_orig.paragraphs if p.text.strip()]
    trans_paras = [p for p in doc_trans.paragraphs if p.text.strip()]
    
    print(f"\nOriginal paragraphs: {len(orig_paras)}")
    print(f"Translated paragraphs: {len(trans_paras)}")
    
    # Count bold runs
    orig_bold_count = 0
    trans_bold_count = 0
    
    print("\nBOLD TEXT ANALYSIS:")
    print("-"*40)
    
    for i, para in enumerate(orig_paras):
        bold_runs = [run for run in para.runs if run.bold]
        if bold_runs:
            orig_bold_count += len(bold_runs)
            print(f"Original Para {i+1}: {len(bold_runs)} bold runs")
            for run in bold_runs:
                print(f"  - '{run.text}'")
    
    for i, para in enumerate(trans_paras):
        bold_runs = [run for run in para.runs if run.bold]
        if bold_runs:
            trans_bold_count += len(bold_runs)
            print(f"Translated Para {i+1}: {len(bold_runs)} bold runs")
            for run in bold_runs:
                print(f"  - '{run.text}'")
    
    print(f"\nTotal bold runs - Original: {orig_bold_count}, Translated: {trans_bold_count}")
    
    # Check formatting preservation
    print("\nFORMATTING PRESERVATION CHECK:")
    print("-"*40)
    
    for i in range(min(len(orig_paras), len(trans_paras))):
        orig_runs = len(orig_paras[i].runs)
        trans_runs = len(trans_paras[i].runs)
        
        if orig_runs != trans_runs:
            print(f"Para {i+1}: Run count mismatch - Original: {orig_runs}, Translated: {trans_runs}")
            
            # Show what happened
            if orig_runs > 1 and trans_runs == 1:
                print(f"  ⚠️  Multiple runs collapsed into single run (formatting lost!)")
                print(f"  Original runs:")
                for j, run in enumerate(orig_paras[i].runs):
                    if run.text.strip():
                        print(f"    Run {j}: '{run.text}' (Bold: {run.bold})")
                print(f"  Translated: '{trans_paras[i].runs[0].text[:50]}...' (Bold: {trans_paras[i].runs[0].bold})")

if __name__ == "__main__":
    compare_documents()
